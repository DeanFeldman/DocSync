"""Repeatable local benchmark for DocSync document-set creation.

Run from the repository root:

    python apps/api/scripts/benchmark_workspace_loading.py
    python apps/api/scripts/benchmark_workspace_loading.py --baseline-ms 12000

The default dataset matches the v1.4.1 Standard profile: 10 documents with
approximately 500 body blocks each. Results stay local and can be written to a
JSON evidence file with ``--output``.
"""

from __future__ import annotations

import argparse
import io
import json
import os
from pathlib import Path
import statistics
import sys
import tempfile
from time import perf_counter

from docx import Document
from fastapi.testclient import TestClient


API_DIR = Path(__file__).resolve().parents[1]


def make_document(document_number: int, block_count: int) -> bytes:
    stream = io.BytesIO()
    document = Document()
    document.add_heading(f"Benchmark document {document_number + 1}", level=1)
    for block_number in range(block_count - 1):
        document.add_paragraph(
            f"Standard shared clause {block_number + 1:04d} for performance testing."
        )
    document.save(stream)
    return stream.getvalue()


def parse_server_timing(value: str) -> dict[str, float]:
    timings: dict[str, float] = {}
    for item in value.split(","):
        name, *parameters = item.strip().split(";")
        for parameter in parameters:
            if parameter.startswith("dur="):
                timings[name] = float(parameter.removeprefix("dur="))
    return timings


def percentile_95(values: list[float]) -> float:
    ordered = sorted(values)
    index = max(0, min(len(ordered) - 1, round(0.95 * len(ordered) + 0.5) - 1))
    return ordered[index]


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--documents", type=int, default=10)
    parser.add_argument("--blocks", type=int, default=500)
    parser.add_argument("--runs", type=int, default=3)
    parser.add_argument("--baseline-ms", type=float)
    parser.add_argument("--output", type=Path)
    parser.add_argument(
        "--work-dir",
        type=Path,
        help="Optional parent directory for isolated benchmark data.",
    )
    parser.add_argument(
        "--api-dir",
        type=Path,
        default=API_DIR,
        help="Backend source directory to benchmark (defaults to apps/api).",
    )
    args = parser.parse_args()

    if args.documents < 2 or args.blocks < 1 or args.runs < 1:
        parser.error("Use at least 2 documents, 1 block, and 1 run.")

    payloads = [
        make_document(document_number, args.blocks)
        for document_number in range(args.documents)
    ]

    if args.work_dir:
        args.work_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(
        prefix="docsync-v141-benchmark-",
        dir=args.work_dir,
    ) as temp_dir:
        benchmark_root = Path(temp_dir)
        os.environ["DOCUMENTSYNC_DATA_DIR"] = str(benchmark_root / "data")
        os.environ["DOCUMENTSYNC_DATABASE_URL"] = (
            f"sqlite:///{benchmark_root / 'benchmark.db'}"
        )

        selected_api_dir = args.api_dir.resolve()
        if str(selected_api_dir) not in sys.path:
            sys.path.insert(0, str(selected_api_dir))
        from app.main import app

        results: list[dict[str, object]] = []
        with TestClient(app) as client:
            for run_number in range(args.runs):
                files = [
                    (
                        "files",
                        (
                            f"Benchmark-{document_number + 1:02d}.docx",
                            io.BytesIO(payload),
                            (
                                "application/vnd.openxmlformats-officedocument."
                                "wordprocessingml.document"
                            ),
                        ),
                    )
                    for document_number, payload in enumerate(payloads)
                ]
                started_at = perf_counter()
                response = client.post(
                    "/api/document-sets",
                    data={"name": f"Standard benchmark {run_number + 1}"},
                    files=files,
                )
                elapsed_ms = (perf_counter() - started_at) * 1000
                response.raise_for_status()
                results.append(
                    {
                        "run": run_number + 1,
                        "elapsed_ms": round(elapsed_ms, 2),
                        "server_timing_ms": parse_server_timing(
                            response.headers.get("server-timing", "")
                        ),
                    }
                )
        from app.database import engine

        engine.dispose()

    elapsed_values = [float(item["elapsed_ms"]) for item in results]
    median_ms = statistics.median(elapsed_values)
    summary: dict[str, object] = {
        "dataset": {
            "documents": args.documents,
            "blocks_per_document": args.blocks,
            "runs": args.runs,
        },
        "median_ms": round(median_ms, 2),
        "p95_ms": round(percentile_95(elapsed_values), 2),
        "runs": results,
    }
    if args.baseline_ms is not None:
        improvement = (args.baseline_ms - median_ms) / args.baseline_ms * 100
        summary["baseline_ms"] = args.baseline_ms
        summary["improvement_percent"] = round(improvement, 2)
        summary["meets_40_percent_target"] = improvement >= 40

    rendered = json.dumps(summary, indent=2)
    print(rendered)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(rendered + "\n", encoding="utf-8")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
