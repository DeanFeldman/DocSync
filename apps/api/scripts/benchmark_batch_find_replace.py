"""Measure exhaustive Find, replacement compilation, and batched DOCX passes."""

from __future__ import annotations

import argparse
from io import BytesIO
import json
from pathlib import Path
import sqlite3
import statistics
import sys
from time import perf_counter

from docx import Document


API_DIRECTORY = Path(__file__).resolve().parents[1]
if str(API_DIRECTORY) not in sys.path:
    sys.path.insert(0, str(API_DIRECTORY))

from app.text_inventory_service import (  # noqa: E402
    DocumentTextInventory,
    DocumentTextSegment,
    TextNodeSpan,
    TextReplacementPatch,
    apply_text_replacements,
    build_text_inventory,
    compile_text_replacement_plan,
    find_occurrence_ranges,
    occurrence_id,
)


NEEDLE = "DOCSYNC_BENCHMARK_NEEDLE"


def median_ms(values: list[float]) -> float:
    return round(statistics.median(values), 2)


def synthetic_segments(documents: int, segments_per_document: int):
    by_document: dict[str, list[str]] = {}
    for document_index in range(documents):
        values = []
        for segment_index in range(segments_per_document):
            text = (
                f"Clause {segment_index}: {NEEDLE} appears here."
                if segment_index % 100 == 0
                else f"Document {document_index} ordinary logical segment {segment_index}."
            )
            values.append(text)
        by_document[f"document-{document_index:03d}"] = values
    return by_document


def benchmark_find(documents: int, segments_per_document: int, runs: int) -> dict:
    values = synthetic_segments(documents, segments_per_document)
    connection = sqlite3.connect(":memory:")
    engine = "fts5_trigram"
    try:
        connection.execute(
            "CREATE VIRTUAL TABLE segment_fts USING fts5(document_id UNINDEXED, text, tokenize='trigram')"
        )
        connection.executemany(
            "INSERT INTO segment_fts(document_id, text) VALUES (?, ?)",
            (
                (document_id, text)
                for document_id, document_segments in values.items()
                for text in document_segments
            ),
        )
    except sqlite3.OperationalError:
        engine = "sqlite_like_fallback"
        connection.execute("CREATE TABLE segment_fts(document_id TEXT, text TEXT)")
        connection.executemany(
            "INSERT INTO segment_fts(document_id, text) VALUES (?, ?)",
            (
                (document_id, text)
                for document_id, document_segments in values.items()
                for text in document_segments
            ),
        )
        connection.execute("CREATE INDEX segment_fts_document ON segment_fts(document_id)")
    connection.commit()

    candidate_runs: list[float] = []
    scanner_runs: list[float] = []
    total_runs: list[float] = []
    result_count = 0
    candidate_count = 0
    for _ in range(runs):
        total_started = perf_counter()
        candidate_started = perf_counter()
        if engine == "fts5_trigram":
            candidates = {
                row[0]
                for row in connection.execute(
                    "SELECT DISTINCT document_id FROM segment_fts WHERE segment_fts MATCH ?",
                    (f'"{NEEDLE}"',),
                )
            }
        else:
            candidates = {
                row[0]
                for row in connection.execute(
                    "SELECT DISTINCT document_id FROM segment_fts WHERE text LIKE ?",
                    (f"%{NEEDLE}%",),
                )
            }
        candidate_runs.append((perf_counter() - candidate_started) * 1000)
        candidate_count = len(candidates)

        scanner_started = perf_counter()
        results = []
        for document_id, document_segments in sorted(
            values.items(),
            key=lambda item: (0 if item[0] in candidates else 1, item[0]),
        ):
            for segment_index, text in enumerate(document_segments):
                for start, end in find_occurrence_ranges(text, NEEDLE):
                    results.append(
                        {
                            "document_id": document_id,
                            "segment_index": segment_index,
                            "match_start": start,
                            "match_end": end,
                            "matched_text": text[start:end],
                        }
                    )
        scanner_runs.append((perf_counter() - scanner_started) * 1000)
        total_runs.append((perf_counter() - total_started) * 1000)
        result_count = len(results)
    connection.close()
    return {
        "documents": documents,
        "segments_per_document": segments_per_document,
        "total_segments": documents * segments_per_document,
        "candidate_engine": engine,
        "candidate_document_count": candidate_count,
        "result_count": result_count,
        "runs": runs,
        "candidate_retrieval_median_ms": median_ms(candidate_runs),
        "exact_occurrence_scan_median_ms": median_ms(scanner_runs),
        "total_result_build_median_ms": median_ms(total_runs),
        "candidate_runs_ms": [round(value, 2) for value in candidate_runs],
        "scanner_runs_ms": [round(value, 2) for value in scanner_runs],
        "total_runs_ms": [round(value, 2) for value in total_runs],
    }


def compilation_inventory(count: int) -> tuple[DocumentTextInventory, list[TextReplacementPatch]]:
    segments = []
    patches = []
    for index in range(count):
        text = f"Prefix {NEEDLE} suffix {index}"
        start = text.index(NEEDLE)
        end = start + len(NEEDLE)
        segment = DocumentTextSegment(
            segment_id=f"segment-{index}",
            document_id="benchmark-document",
            version_id="benchmark-version",
            part_path="word/document.xml",
            structure_type="body_paragraph",
            text=text,
            normalized_text=text.casefold(),
            node_path=(index,),
            location={"kind": "body", "paragraph_index": index},
            spans=[
                TextNodeSpan(
                    node_path=(index, 0),
                    logical_start=0,
                    logical_end=len(text),
                    text=text,
                    node_kind="text",
                    editable=True,
                )
            ],
        )
        segments.append(segment)
        patches.append(
            TextReplacementPatch(
                occurrence_id=f"occurrence-{index}",
                segment_id=segment.segment_id,
                part_path=segment.part_path,
                match_start=start,
                match_end=end,
                expected_text=NEEDLE,
                replacement_text="DOCSYNC_REPLACEMENT",
            )
        )
    return (
        DocumentTextInventory(
            document_id="benchmark-document",
            version_id="benchmark-version",
            package_sha256="benchmark",
            segments=segments,
        ),
        patches,
    )


def benchmark_compilation(runs: int) -> dict:
    result = {}
    for count in (100, 500, 1_000, 5_000):
        inventory, patches = compilation_inventory(count)
        durations = []
        node_count = 0
        for _ in range(runs):
            started = perf_counter()
            plan = compile_text_replacement_plan(inventory, patches)
            durations.append((perf_counter() - started) * 1000)
            node_count = len(plan.node_edits)
        result[str(count)] = {
            "median_ms": median_ms(durations),
            "runs_ms": [round(value, 2) for value in durations],
            "compiled_nodes": node_count,
        }
    return result


def generation_payload() -> bytes:
    document = Document()
    for index in range(100):
        if index < 10:
            document.add_paragraph(f"Clause {index}: DOCSYNC_TOKEN_{index:02d}.")
        else:
            document.add_paragraph(f"Unchanged paragraph {index}.")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def patches_for_tokens(payload: bytes, token_indexes: list[int]) -> tuple[DocumentTextInventory, list[TextReplacementPatch]]:
    inventory = build_text_inventory(
        payload,
        document_id="benchmark-document",
        version_id="benchmark-version",
    )
    patches = []
    for token_index in token_indexes:
        token = f"DOCSYNC_TOKEN_{token_index:02d}"
        segment = next(item for item in inventory.segments if token in item.text)
        start, end = find_occurrence_ranges(segment.text, token)[0]
        patches.append(
            TextReplacementPatch(
                occurrence_id=occurrence_id(segment, start, end),
                segment_id=segment.segment_id,
                part_path=segment.part_path,
                match_start=start,
                match_end=end,
                expected_text=token,
                replacement_text=f"REPLACED_TOKEN_{token_index:02d}",
            )
        )
    return inventory, patches


def benchmark_generation(documents: int, edits_per_document: int, runs: int) -> dict:
    source = generation_payload()
    batched_runs = []
    repeated_runs = []
    for _ in range(runs):
        payloads = [source for _index in range(documents)]
        started = perf_counter()
        for payload in payloads:
            inventory, patches = patches_for_tokens(
                payload,
                list(range(edits_per_document)),
            )
            apply_text_replacements(payload, inventory, patches)
        batched_runs.append((perf_counter() - started) * 1000)

        payloads = [source for _index in range(documents)]
        started = perf_counter()
        for payload in payloads:
            current = payload
            for token_index in range(edits_per_document):
                inventory, patches = patches_for_tokens(current, [token_index])
                current = apply_text_replacements(current, inventory, patches)
        repeated_runs.append((perf_counter() - started) * 1000)
    batched = median_ms(batched_runs)
    repeated = median_ms(repeated_runs)
    return {
        "documents": documents,
        "edits_per_document": edits_per_document,
        "batched_package_passes": documents,
        "repeated_package_passes": documents * edits_per_document,
        "batched_median_ms": batched,
        "repeated_median_ms": repeated,
        "speedup_x": round(repeated / batched, 2) if batched else None,
        "batched_runs_ms": [round(value, 2) for value in batched_runs],
        "repeated_runs_ms": [round(value, 2) for value in repeated_runs],
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--documents", type=int, default=50)
    parser.add_argument("--segments", type=int, default=1_000)
    parser.add_argument("--generation-documents", type=int, default=10)
    parser.add_argument("--edits", type=int, default=10)
    parser.add_argument("--runs", type=int, default=3)
    args = parser.parse_args()
    if min(args.documents, args.segments, args.generation_documents, args.edits, args.runs) < 1:
        parser.error("All benchmark sizes must be positive.")
    result = {
        "find": benchmark_find(args.documents, args.segments, args.runs),
        "replace_compilation": benchmark_compilation(args.runs),
        "batch_generation": benchmark_generation(
            args.generation_documents,
            args.edits,
            args.runs,
        ),
    }
    print(json.dumps(result, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
