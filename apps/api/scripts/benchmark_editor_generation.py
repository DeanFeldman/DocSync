from __future__ import annotations

import argparse
import io
import json
import os
from pathlib import Path
import statistics
import sys
from time import perf_counter

from docx import Document
from fastapi.testclient import TestClient


DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def docx_payload(shared: str, blocks: int, document_number: int) -> bytes:
    document = Document()
    document.add_heading(f"Generation benchmark document {document_number}", level=1)
    for index in range(blocks):
        document.add_paragraph(
            shared
            if index == blocks // 2
            else f"Document {document_number}, unchanged paragraph {index}."
        )
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def run_once(client: TestClient, run_number: int, documents: int, blocks: int) -> float:
    shared = f"Shared benchmark wording for run {run_number}."
    upload = client.post(
        "/api/document-sets",
        data={"name": f"Generation benchmark {run_number}"},
        files=[
            (
                "files",
                (
                    f"Benchmark-{index + 1}.docx",
                    io.BytesIO(docx_payload(shared, blocks, index + 1)),
                    DOCX_MEDIA_TYPE,
                ),
            )
            for index in range(documents)
        ],
    )
    upload.raise_for_status()
    workspace = upload.json()
    targets = []
    base_versions = {}
    for summary in workspace["documents"]:
        version_id = summary["version_id"]
        content = client.get(
            f"/api/document-versions/{version_id}/editor-content"
        )
        content.raise_for_status()
        block = next(item for item in content.json()["blocks"] if item["text"] == shared)
        base_versions[summary["id"]] = version_id
        targets.append(
            {
                "element_id": block["element_id"],
                "replacement_text": f"{shared} Updated",
            }
        )

    request = {
        "base_versions": base_versions,
        "source_element_id": targets[0]["element_id"],
        "edit_mode": "shared",
        "targets": targets,
    }
    started = perf_counter()
    generated = client.post(
        f"/api/document-sets/{workspace['id']}/editor-generate",
        json=request,
    )
    elapsed_ms = (perf_counter() - started) * 1000
    generated.raise_for_status()
    assert len(generated.json()["versions"]) == documents
    return elapsed_ms


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--data-dir", type=Path, required=True)
    parser.add_argument("--documents", type=int, default=3)
    parser.add_argument("--blocks", type=int, default=300)
    parser.add_argument("--runs", type=int, default=3)
    args = parser.parse_args()
    args.data_dir.mkdir(parents=True, exist_ok=False)

    os.environ["DOCUMENTSYNC_DATA_DIR"] = str(args.data_dir)
    os.environ["DOCUMENTSYNC_DATABASE_URL"] = (
        f"sqlite:///{(args.data_dir / 'benchmark.db').as_posix()}"
    )
    os.environ["DOCUMENTSYNC_SESSION_TOKEN"] = ""
    api_directory = Path(__file__).resolve().parents[1]
    if str(api_directory) not in sys.path:
        sys.path.insert(0, str(api_directory))

    from app.main import app

    with TestClient(app) as client:
        durations = [
            run_once(client, run_number + 1, args.documents, args.blocks)
            for run_number in range(args.runs)
        ]
    result = {
        "documents": args.documents,
        "blocks_per_document": args.blocks,
        "runs_ms": [round(value, 2) for value in durations],
        "median_ms": round(statistics.median(durations), 2),
        "maximum_ms": round(max(durations), 2),
    }
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
