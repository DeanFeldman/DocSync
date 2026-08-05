from __future__ import annotations

import argparse
import io
import json
import os
from pathlib import Path
import statistics
import sys
from time import perf_counter

from docx import Document
from fastapi.testclient import TestClient


DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def document_payload(blocks: int, label: str) -> bytes:
    document = Document()
    document.add_heading(label, level=1)
    for index in range(blocks):
        document.add_paragraph(f"{label} readable paragraph {index}.")
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def timed_get(client: TestClient, url: str) -> tuple[float, str]:
    started = perf_counter()
    response = client.get(url)
    duration_ms = (perf_counter() - started) * 1000
    response.raise_for_status()
    return duration_ms, response.headers.get("x-docsync-preview-cache", "unknown")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--data-dir", type=Path, required=True)
    parser.add_argument("--blocks", type=int, default=1_000)
    parser.add_argument("--cached-runs", type=int, default=5)
    args = parser.parse_args()
    args.data_dir.mkdir(parents=True, exist_ok=False)

    os.environ["DOCUMENTSYNC_DATA_DIR"] = str(args.data_dir)
    os.environ["DOCUMENTSYNC_DATABASE_URL"] = (
        f"sqlite:///{(args.data_dir / 'benchmark.db').as_posix()}"
    )
    os.environ["DOCUMENTSYNC_SESSION_TOKEN"] = ""
    api_directory = Path(__file__).resolve().parents[1]
    if str(api_directory) not in sys.path:
        sys.path.insert(0, str(api_directory))

    from app.main import app

    with TestClient(app) as client:
        upload = client.post(
            "/api/document-sets",
            data={"name": "Opening benchmark"},
            files=[
                (
                    "files",
                    (
                        "Large.docx",
                        io.BytesIO(document_payload(args.blocks, "Large")),
                        DOCX_MEDIA_TYPE,
                    ),
                ),
                (
                    "files",
                    (
                        "Companion.docx",
                        io.BytesIO(document_payload(1, "Companion")),
                        DOCX_MEDIA_TYPE,
                    ),
                ),
            ],
        )
        upload.raise_for_status()
        summary = next(
            item for item in upload.json()["documents"] if item["name"] == "Large.docx"
        )
        url = f"/api/document-versions/{summary['version_id']}/pages"
        uncached_ms, uncached_state = timed_get(client, url)
        cached = [timed_get(client, url) for _ in range(args.cached_runs)]

    print(
        json.dumps(
            {
                "blocks": args.blocks,
                "uncached_structured_preview_ms": round(uncached_ms, 2),
                "uncached_cache_header": uncached_state,
                "cached_runs_ms": [round(duration, 2) for duration, _state in cached],
                "cached_median_ms": round(
                    statistics.median(duration for duration, _state in cached),
                    2,
                ),
                "cached_cache_headers": [state for _duration, state in cached],
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
