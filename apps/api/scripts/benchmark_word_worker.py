"""Measure cold startup and warm sequential renders for the persistent Word worker."""

from __future__ import annotations

import argparse
import json
import os
from pathlib import Path
import statistics
import sys
from time import perf_counter

from docx import Document


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--data-dir", type=Path, required=True)
    parser.add_argument("--renders", type=int, default=3)
    args = parser.parse_args()
    if args.renders < 3:
        parser.error("Use at least three sequential renders.")
    args.data_dir.mkdir(parents=True, exist_ok=False)
    os.environ["DOCUMENTSYNC_DATA_DIR"] = str(args.data_dir)
    os.environ["DOCUMENTSYNC_WORD_WORKER_AUTOSTART"] = "0"

    api_directory = Path(__file__).resolve().parents[1]
    if str(api_directory) not in sys.path:
        sys.path.insert(0, str(api_directory))
    from app.word_render_service import (
        WORD_RENDER_WORKER,
        render_docx_with_word_worker,
        shutdown_word_render_worker,
        start_word_render_worker,
    )

    sources = []
    for index in range(args.renders):
        source = args.data_dir / f"word-worker-{index + 1}.docx"
        document = Document()
        document.add_heading(f"Persistent worker benchmark {index + 1}", level=1)
        for paragraph in range(40):
            document.add_paragraph(
                f"Sequential uncached render {index + 1}, paragraph {paragraph + 1}."
            )
        document.save(source)
        sources.append(source)

    startup_started = perf_counter()
    if not start_word_render_worker():
        print(json.dumps({"status": "not_measured", "reason": "Microsoft Word worker unavailable"}, indent=2))
        return 0
    startup_ms = (perf_counter() - startup_started) * 1000
    runs_ms = []
    worker_process_ids = []
    try:
        for index, source in enumerate(sources):
            output = args.data_dir / f"word-worker-{index + 1}.pdf"
            started_at = perf_counter()
            render_docx_with_word_worker(source, output)
            runs_ms.append(round((perf_counter() - started_at) * 1000, 2))
            process = WORD_RENDER_WORKER._process
            worker_process_ids.append(process.pid if process is not None else None)
    except Exception as exc:
        print(json.dumps({"status": "not_measured", "reason": str(exc)}, indent=2))
        return 0
    finally:
        shutdown_word_render_worker()

    warm_runs = runs_ms[1:]
    print(
        json.dumps(
            {
                "status": "measured",
                "worker_startup_ms": round(startup_ms, 2),
                "render_runs_ms": runs_ms,
                "cold_first_render_including_startup_ms": round(startup_ms + runs_ms[0], 2),
                "warm_worker_median_ms": round(statistics.median(warm_runs), 2),
                "worker_reused": len(set(worker_process_ids)) == 1,
            },
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
