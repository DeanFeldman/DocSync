from __future__ import annotations

from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timezone
from difflib import SequenceMatcher
from io import BytesIO
import hashlib
import json
import logging
import math
from pathlib import Path
from types import SimpleNamespace
import re
import shutil
import threading
from time import perf_counter
import unicodedata
from uuid import uuid4
import zipfile

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from fastapi import HTTPException, status
from sqlalchemy import delete, func, select
from sqlalchemy.orm import Session, selectinload

from .config import settings
from .models import (
    DocumentBlockRevision,
    DocumentElement,
    DocumentHead,
    DocumentRecord,
    DocumentSet,
    DocumentVersion,
    EditBatchOperation,
    EditorOperation,
    EditorOperationTarget,
    GeneratedVersion,
    LinkGroup,
    MatchDecision,
)
from .schemas import (
    CompareRequest,
    EditorEditRequest,
    EditorMatchDecision,
    EditorTarget,
    MatchDecisionBatchRequest,
    QuillDelta,
    VersionRestoreRequest,
)


logger = logging.getLogger(__name__)


BODY_STYLE_PATTERN = re.compile(r"^body_order:(\d+):(.*)$", re.DOTALL)
BLOCK_BODY_STYLE_PATTERN = re.compile(
    r"^body_block_order:(\d+):(\d+):(\d+):(.*)$",
    re.DOTALL,
)
TABLE_STYLE_PATTERN = re.compile(r"^table_cell:(\d+):(\d+):(\d+)$")
ORDERED_TABLE_STYLE_PATTERN = re.compile(
    r"^table_cell_order:(\d+):(\d+):(\d+):(\d+)$"
)
TABLE_PARAGRAPH_STYLE_PATTERN = re.compile(
    r"^table_paragraph:(\d+):(\d+):(\d+):(\d+)$"
)
ORDERED_TABLE_PARAGRAPH_STYLE_PATTERN = re.compile(
    r"^table_paragraph_order:(\d+):(\d+):(\d+):(\d+):(\d+):(\d+)$"
)
HEADER_FOOTER_STYLE_PATTERN = re.compile(
    r"^header_footer_order:(\d+):(\d+):(header|footer):(\d+):"
    r"(default|first|even):(\d+):([^:]+)$"
)
HEADER_FOOTER_PART_TYPES = (
    ("header", "default"),
    ("header", "first"),
    ("header", "even"),
    ("footer", "default"),
    ("footer", "first"),
    ("footer", "even"),
)
LIST_LEVEL_SUFFIX = re.compile(r"\s+(\d+)$")
TOKEN_PATTERN = re.compile(r"\w+|[^\w\s]+|\s+", re.UNICODE)
EDITOR_GENERATION_LOCK = threading.RLock()
EDITOR_QUEUE_LOCK = threading.RLock()
EDITOR_GENERATION_EXECUTOR = ThreadPoolExecutor(
    max_workers=4,
    thread_name_prefix="docsync-generation",
)
MATCH_ALGORITHM_VERSION = "nfkc-sequence-v1"
ALLOWED_INLINE_ATTRIBUTES = {"bold", "italic", "underline"}
ALLOWED_BLOCK_ATTRIBUTES = {"header", "list", "indent", "align"}
UNSUPPORTED_XML_TAGS = {
    qn("w:drawing"): "Drawing or floating object",
    qn("w:pict"): "Legacy picture or text box",
    qn("w:object"): "Embedded object",
    qn("w:fldSimple"): "Word field",
    qn("w:fldChar"): "Word field",
    qn("w:instrText"): "Word field instruction",
    qn("w:sdt"): "Content control",
    qn("w:hyperlink"): "Hyperlink content",
    qn("w:ins"): "Tracked insertion",
    qn("w:del"): "Tracked deletion",
    qn("w:commentReference"): "Comment reference",
    qn("w:footnoteReference"): "Footnote reference",
    qn("w:endnoteReference"): "Endnote reference",
}


def _queue_previews_safely(session: Session, version_ids: list[str]) -> list[dict]:
    """Preview failure must never roll back a committed immutable version."""

    try:
        from .preview_job_service import queue_generated_version_previews

        return queue_generated_version_previews(session, version_ids)
    except Exception:
        session.rollback()
        logger.exception(
            "docsync.post_generation_preview_queue.failed versions=%s",
            ",".join(version_ids),
        )
        return []


def _header_footer_type(kind: str, variant: str) -> str:
    prefix = {"default": "default", "first": "first_page", "even": "even_page"}[
        variant
    ]
    return f"{prefix}_{kind}"


def _header_footer_paragraph_nodes(part_element) -> list:
    paragraphs = []
    for node in part_element.iter(qn("w:p")):
        ancestor = node.getparent()
        nested_in_table = False
        while ancestor is not None and ancestor is not part_element:
            if ancestor.tag == qn("w:tbl"):
                nested_in_table = True
                break
            ancestor = ancestor.getparent()
        if not nested_in_table:
            paragraphs.append(node)
    return paragraphs


def new_id() -> str:
    return str(uuid4())


def utc_now() -> datetime:
    return datetime.now(timezone.utc)


def utc_isoformat(value: datetime) -> str:
    if value.tzinfo is None:
        value = value.replace(tzinfo=timezone.utc)
    return value.astimezone(timezone.utc).isoformat()


def normalise_editor_text(text: str) -> str:
    """Canonical text used by exact matching and persisted exact hashes."""

    return " ".join(unicodedata.normalize("NFKC", text).split()).casefold()


def exact_match_hash(
    element_type: str,
    text: str,
    context_identity: str | None = None,
) -> str:
    normalized = normalise_editor_text(text)
    identity = (
        f"{element_type}\0{context_identity}\0{normalized}"
        if context_identity
        else f"{element_type}\0{normalized}"
    )
    return hashlib.sha256(
        identity.encode("utf-8")
    ).hexdigest()


def _canonical_hash(value: object) -> str:
    return hashlib.sha256(
        json.dumps(
            value,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()


def _safe_storage_path(storage_area: str, storage_name: str) -> Path:
    if storage_area not in {"originals", "generated"}:
        raise HTTPException(status_code=500, detail="Document version storage is invalid.")
    root = (settings.data_dir / storage_area).resolve()
    path = (root / storage_name).resolve()
    if path != root and root not in path.parents:
        raise HTTPException(status_code=500, detail="Document version storage is invalid.")
    return path


def document_version_path(version: DocumentVersion) -> Path:
    path = _safe_storage_path(version.storage_area, version.storage_name)
    if not path.is_file():
        raise HTTPException(status_code=404, detail="Document version file is missing.")
    return path


def get_version_or_404(session: Session, version_id: str) -> DocumentVersion:
    version = session.scalar(
        select(DocumentVersion)
        .where(DocumentVersion.id == version_id)
        .options(
            selectinload(DocumentVersion.document),
            selectinload(DocumentVersion.blocks),
        )
    )
    if version is None:
        # A newly inserted logical document may not yet have passed through the
        # editor initializer. Keep this fallback for non-lifespan unit callers.
        document = session.get(DocumentRecord, version_id)
        if document is not None:
            initialise_original_version(session, document)
            session.flush()
            version = session.scalar(
                select(DocumentVersion)
                .where(DocumentVersion.id == version_id)
                .options(
                    selectinload(DocumentVersion.document),
                    selectinload(DocumentVersion.blocks),
                )
            )
    if version is None:
        raise HTTPException(status_code=404, detail="Document version not found.")
    return version


def resolve_document_identifier(
    session: Session,
    document_or_version_id: str,
) -> tuple[DocumentRecord, DocumentVersion]:
    document = session.get(DocumentRecord, document_or_version_id)
    if document is not None:
        version = current_version_for_document(session, document)
        return document, version
    version = get_version_or_404(session, document_or_version_id)
    return version.document, version


def current_version_for_document(
    session: Session,
    document: DocumentRecord,
) -> DocumentVersion:
    head = session.get(DocumentHead, document.id)
    if head is None:
        initialise_original_version(session, document)
        session.flush()
        head = session.get(DocumentHead, document.id)
    if head is None:
        raise HTTPException(status_code=500, detail="Document version head is missing.")
    return get_version_or_404(session, head.current_version_id)


def _display_style_name(style_name: str | None) -> str | None:
    if HEADER_FOOTER_STYLE_PATTERN.fullmatch(style_name or ""):
        return None
    block_match = BLOCK_BODY_STYLE_PATTERN.fullmatch(style_name or "")
    if block_match is not None:
        return block_match.group(4) or None
    match = BODY_STYLE_PATTERN.fullmatch(style_name or "")
    if match is None:
        if TABLE_STYLE_PATTERN.fullmatch(style_name or ""):
            return None
        if ORDERED_TABLE_STYLE_PATTERN.fullmatch(style_name or ""):
            return None
        if TABLE_PARAGRAPH_STYLE_PATTERN.fullmatch(style_name or ""):
            return None
        if ORDERED_TABLE_PARAGRAPH_STYLE_PATTERN.fullmatch(style_name or ""):
            return None
        return style_name
    return match.group(2) or None


def _element_type(style_name: str | None) -> str:
    value = style_name or ""
    header_footer = HEADER_FOOTER_STYLE_PATTERN.fullmatch(value)
    if header_footer is not None:
        return f"{header_footer.group(3)}_paragraph"
    if TABLE_PARAGRAPH_STYLE_PATTERN.fullmatch(
        value
    ) or ORDERED_TABLE_PARAGRAPH_STYLE_PATTERN.fullmatch(value):
        return "table_paragraph"
    if TABLE_STYLE_PATTERN.fullmatch(value) or ORDERED_TABLE_STYLE_PATTERN.fullmatch(
        value
    ):
        return "table_cell"
    style = (_display_style_name(value) or "").casefold()
    if style.startswith(("heading", "title")):
        return "heading"
    if style.startswith("list"):
        return "list_item"
    return "paragraph"


def _element_location(element: DocumentElement) -> tuple[int, dict]:
    style_name = element.style_name or ""
    header_footer = HEADER_FOOTER_STYLE_PATTERN.fullmatch(style_name)
    if header_footer is not None:
        (
            order,
            document_order,
            kind,
            source_section_index,
            variant,
            paragraph_index,
            relationship_id,
        ) = header_footer.groups()
        return int(order), {
            "kind": f"{kind}_paragraph",
            "section_index": int(source_section_index),
            "source_section_index": int(source_section_index),
            "header_footer_type": _header_footer_type(kind, variant),
            "paragraph_index": int(paragraph_index),
            "part_relationship_id": relationship_id,
            "document_order": int(document_order),
        }
    table_paragraph = ORDERED_TABLE_PARAGRAPH_STYLE_PATTERN.fullmatch(style_name)
    if table_paragraph is not None:
        (
            order,
            document_order,
            table_index,
            row_index,
            column_index,
            paragraph_index,
        ) = (int(value) for value in table_paragraph.groups())
        return order, {
            "kind": "table_paragraph",
            "document_order": document_order,
            "paragraph_index": paragraph_index,
            "table_index": table_index,
            "row_index": row_index,
            "column_index": column_index,
        }
    legacy_table_paragraph = TABLE_PARAGRAPH_STYLE_PATTERN.fullmatch(style_name)
    if legacy_table_paragraph is not None:
        table_index, row_index, column_index, paragraph_index = (
            int(value) for value in legacy_table_paragraph.groups()
        )
        return element.paragraph_index, {
            "kind": "table_paragraph",
            "document_order": element.paragraph_index,
            "paragraph_index": paragraph_index,
            "table_index": table_index,
            "row_index": row_index,
            "column_index": column_index,
        }
    table_match = ORDERED_TABLE_STYLE_PATTERN.fullmatch(style_name)
    if table_match is not None:
        order, table_index, row_index, column_index = (
            int(value) for value in table_match.groups()
        )
        return order, {
            "kind": "table_cell",
            "document_order": order,
            "paragraph_index": element.paragraph_index,
            "table_index": table_index,
            "row_index": row_index,
            "column_index": column_index,
        }
    legacy_table = TABLE_STYLE_PATTERN.fullmatch(style_name)
    if legacy_table is not None:
        table_index, row_index, column_index = (
            int(value) for value in legacy_table.groups()
        )
        return element.paragraph_index, {
            "kind": "table_cell",
            "document_order": element.paragraph_index,
            "paragraph_index": element.paragraph_index,
            "table_index": table_index,
            "row_index": row_index,
            "column_index": column_index,
        }
    body_block = BLOCK_BODY_STYLE_PATTERN.fullmatch(style_name)
    if body_block is not None:
        order, document_order, paragraph_index, _style = body_block.groups()
        return int(order), {
            "kind": "body",
            "document_order": int(document_order),
            "paragraph_index": int(paragraph_index),
        }
    body_match = BODY_STYLE_PATTERN.fullmatch(style_name)
    order = int(body_match.group(1)) if body_match is not None else element.paragraph_index
    return order, {
        "kind": "body",
        "document_order": order,
        "paragraph_index": element.paragraph_index,
    }


def _location_key(location: dict | None) -> str:
    value = dict(location or {})
    for derived_key in (
        "section_indexes",
        "linked_section_indexes",
        "linked_sections",
        "is_linked_to_previous",
        "affected_header_footer_types",
    ):
        value.pop(derived_key, None)
    # paragraph_index is synthetic for cells and is not needed for write-back.
    if value.get("kind") == "table_cell":
        value.pop("paragraph_index", None)
    return json.dumps(value, sort_keys=True, separators=(",", ":"))


def _header_footer_part_map(document: DocxDocument) -> dict[str, dict]:
    """Return physical header/footer parts and every section usage.

    A section without its own reference inherits the active relationship from
    the previous section. Grouping by relationship ID makes one physical
    paragraph one editable target even when several sections inherit it.
    """

    active_references: dict[tuple[str, str], tuple[int, str]] = {}
    result: dict[str, dict] = {}
    for section_index, section in enumerate(document.sections):
        for kind, variant in HEADER_FOOTER_PART_TYPES:
            reference_tag = qn(f"w:{kind}Reference")
            direct_reference = next(
                (
                    reference
                    for reference in section._sectPr.findall(reference_tag)
                    if reference.get(qn("w:type"), "default") == variant
                ),
                None,
            )
            key = (kind, variant)
            if direct_reference is not None:
                relationship_id = direct_reference.get(qn("r:id"))
                if not relationship_id:
                    continue
                active_references[key] = (section_index, relationship_id)
            active = active_references.get(key)
            if active is None:
                continue
            source_section_index, relationship_id = active
            try:
                part = document.part.related_parts[relationship_id]
            except KeyError:
                continue
            entry = result.setdefault(
                relationship_id,
                {
                    "part": part,
                    "paragraphs": [
                        Paragraph(node, SimpleNamespace(part=part))
                        for node in _header_footer_paragraph_nodes(part.element)
                    ],
                    "usages": [],
                },
            )
            entry["usages"].append(
                {
                    "kind": kind,
                    "header_footer_type": _header_footer_type(kind, variant),
                    "section_index": section_index,
                    "source_section_index": source_section_index,
                    "is_linked_to_previous": direct_reference is None,
                }
            )
    return result


def _enrich_header_footer_location(
    location: dict,
    part_map: dict[str, dict],
) -> dict:
    result = dict(location)
    relationship_id = str(location.get("part_relationship_id") or "")
    entry = part_map.get(relationship_id)
    if entry is None:
        return result
    kind = str(location.get("kind") or "").removesuffix("_paragraph")
    usages = [
        usage
        for usage in entry["usages"]
        if usage["kind"] == kind
    ]
    section_indexes = sorted({int(usage["section_index"]) for usage in usages})
    linked_sections = sorted(
        {
            int(usage["section_index"])
            for usage in usages
            if usage["is_linked_to_previous"]
        }
    )
    result.update(
        {
            "section_indexes": section_indexes,
            "linked_section_indexes": linked_sections,
            "linked_sections": section_indexes,
            "is_linked_to_previous": bool(linked_sections),
            "affected_header_footer_types": sorted(
                {str(usage["header_footer_type"]) for usage in usages}
            ),
        }
    )
    return result


def _match_context_identity(
    element_type: str,
    location: dict | None,
) -> str:
    location = location or {}
    if element_type in {"header_paragraph", "footer_paragraph"}:
        return (
            f"{location.get('header_footer_type', '')}:"
            f"section:{location.get('section_index', '')}"
        )
    return ""


def _compatible_match_context(
    source: DocumentBlockRevision,
    candidate: DocumentBlockRevision,
) -> bool:
    if source.element_type != candidate.element_type:
        return False
    return _match_context_identity(
        source.element_type,
        source.location_json,
    ) == _match_context_identity(
        candidate.element_type,
        candidate.location_json,
    )


def _paragraph_for_element(
    document: DocxDocument,
    element: DocumentElement,
    *,
    paragraphs: list[Paragraph] | None = None,
    tables: list[object] | None = None,
    header_footer_parts: dict[str, dict] | None = None,
) -> tuple[Paragraph | None, object | None]:
    _ordinal, location = _element_location(element)
    if location["kind"] in {"header_paragraph", "footer_paragraph"}:
        parts = (
            header_footer_parts
            if header_footer_parts is not None
            else _header_footer_part_map(document)
        )
        entry = parts.get(str(location.get("part_relationship_id") or ""))
        if entry is None:
            return None, None
        try:
            return entry["paragraphs"][int(location["paragraph_index"])], None
        except (IndexError, KeyError, TypeError, ValueError):
            return None, None
    if location["kind"] == "body":
        index = int(location["paragraph_index"])
        body_paragraphs = (
            paragraphs if paragraphs is not None else document.paragraphs
        )
        if 0 <= index < len(body_paragraphs):
            return body_paragraphs[index], None
        return None, None
    try:
        document_tables = tables if tables is not None else document.tables
        cell = (
            document_tables[int(location["table_index"])]
            .rows[int(location["row_index"])]
            .cells[int(location["column_index"])]
        )
    except IndexError:
        return None, None
    if location["kind"] == "table_paragraph":
        try:
            return cell.paragraphs[int(location["paragraph_index"])], cell
        except (IndexError, KeyError, TypeError, ValueError):
            return None, cell
    non_empty = [paragraph for paragraph in cell.paragraphs if paragraph.text.strip()]
    return (non_empty[0] if non_empty else cell.paragraphs[0]), cell


def _alignment_name(paragraph: Paragraph | None) -> str | None:
    if paragraph is None or paragraph.alignment is None:
        return None
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
    }
    return mapping.get(paragraph.alignment)


def _run_attributes(run) -> dict:
    attributes: dict[str, object] = {}
    if run.bold:
        attributes["bold"] = True
    if run.italic:
        attributes["italic"] = True
    if run.underline:
        attributes["underline"] = True
    return attributes


def _paragraph_style_name(
    paragraph: Paragraph,
    cache: dict[str | None, str | None] | None = None,
) -> str | None:
    """Resolve a paragraph style once per underlying Word style ID.

    python-docx searches the complete style collection whenever an unstyled
    paragraph asks for its default style. Large document migrations contain
    thousands of those paragraphs, so cache both explicit style IDs and the
    default (``None``) lookup for the lifetime of one parsed document.
    """

    properties = paragraph._p.pPr
    style_element = properties.pStyle if properties is not None else None
    style_id = str(style_element.val) if style_element is not None else None
    if cache is not None and style_id in cache:
        return cache[style_id]
    style = paragraph.style
    name = style.name if style is not None else None
    if cache is not None:
        cache[style_id] = name
    return name


def _numbering_type(
    document: DocxDocument,
    paragraph: Paragraph,
    style_name: str | None = None,
) -> tuple[str | None, int | None]:
    resolved_style_name = style_name or ""
    paragraph_properties = paragraph._p.pPr
    numbering_properties = (
        paragraph_properties.numPr if paragraph_properties is not None else None
    )
    level: int | None = None
    number_id: int | None = None
    if numbering_properties is not None:
        if numbering_properties.ilvl is not None:
            level = int(numbering_properties.ilvl.val)
        if numbering_properties.numId is not None:
            number_id = int(numbering_properties.numId.val)

    list_type: str | None = None
    if number_id is not None:
        try:
            numbering = document.part.numbering_part.element
            abstract_id: int | None = None
            for num in numbering.findall(qn("w:num")):
                if int(num.get(qn("w:numId"))) != number_id:
                    continue
                value = num.find(qn("w:abstractNumId"))
                if value is not None:
                    abstract_id = int(value.get(qn("w:val")))
                break
            if abstract_id is not None:
                for abstract in numbering.findall(qn("w:abstractNum")):
                    if int(abstract.get(qn("w:abstractNumId"))) != abstract_id:
                        continue
                    requested_level = level or 0
                    selected_level = None
                    for item in abstract.findall(qn("w:lvl")):
                        if int(item.get(qn("w:ilvl"))) == requested_level:
                            selected_level = item
                            break
                    if selected_level is None:
                        selected_level = abstract.find(qn("w:lvl"))
                    if selected_level is not None:
                        number_format = selected_level.find(qn("w:numFmt"))
                        value = (
                            number_format.get(qn("w:val"))
                            if number_format is not None
                            else ""
                        )
                        list_type = "bullet" if value == "bullet" else "ordered"
                    break
        except (AttributeError, KeyError, TypeError, ValueError):
            list_type = None

    folded_style = resolved_style_name.casefold()
    if list_type is None and folded_style.startswith("list"):
        list_type = "ordered" if "number" in folded_style else "bullet"
    if list_type is not None and level is None:
        match = LIST_LEVEL_SUFFIX.search(resolved_style_name)
        level = max(int(match.group(1)) - 1, 0) if match else 0
    return list_type, level


def _paragraph_block_attributes(
    element_type: str,
    paragraph: Paragraph | None,
    list_type: str | None,
    list_level: int | None,
    alignment: str | None,
    style_name: str | None = None,
) -> dict:
    attributes: dict[str, object] = {}
    if element_type == "heading" and paragraph is not None:
        match = re.fullmatch(
            r"Heading\s+(\d+)",
            style_name or "",
            re.IGNORECASE,
        )
        attributes["header"] = (
            min(max(int(match.group(1)), 1), 6) if match is not None else 1
        )
    if list_type is not None:
        attributes["list"] = list_type
        if list_level:
            attributes["indent"] = list_level
    if alignment is not None and alignment != "left":
        attributes["align"] = alignment
    return attributes


def _paragraph_delta(
    text: str,
    paragraph: Paragraph | None,
    block_attributes: dict,
    formatted_runs: list[dict] | None = None,
) -> dict:
    operations: list[dict] = []
    if paragraph is not None and paragraph.runs:
        covered = ""
        runs = formatted_runs
        if runs is None:
            runs = [
                {"text": run.text, **_run_attributes(run)}
                for run in paragraph.runs
                if run.text
            ]
        for run in runs:
            run_text = str(run.get("text") or "")
            if not run_text:
                continue
            operation = {"insert": run_text}
            attributes = {
                key: value for key, value in run.items() if key != "text"
            }
            if attributes:
                operation["attributes"] = attributes
            operations.append(operation)
            covered += run_text
        # Hyperlinks and some field content are not exposed as normal runs. The
        # read-only diagnostic handles those; keep the visible text complete.
        if covered != text and text:
            operations = [{"insert": text}]
    elif text:
        operations.append({"insert": text})
    newline: dict[str, object] = {"insert": "\n"}
    if block_attributes:
        newline["attributes"] = block_attributes
    operations.append(newline)
    return {"ops": operations}


def _cell_uses_merged_structure(cell: object | None) -> bool:
    if cell is None:
        return False
    try:
        properties = cell._tc.tcPr
        grid_span = properties.gridSpan
        if grid_span is not None and int(grid_span.val or 1) > 1:
            return True
        if properties.vMerge is not None:
            return True
    except (AttributeError, TypeError, ValueError):
        return True
    return False


def _unsupported_reason(
    paragraph: Paragraph | None,
    cell: object | None,
    *,
    allow_multiple_paragraphs: bool = False,
    header_footer_kind: str | None = None,
) -> str | None:
    if paragraph is not None:
        for node in paragraph._p.iter():
            reason = UNSUPPORTED_XML_TAGS.get(node.tag)
            if reason is not None:
                if header_footer_kind is not None:
                    return (
                        f"This {header_footer_kind} paragraph contains "
                        f"{reason.lower()} that DocSync cannot safely rewrite. "
                        "The content will remain unchanged."
                    )
                return f"{reason} is read-only in Edit mode."
        ancestor = paragraph._p.getparent()
        while ancestor is not None:
            reason = UNSUPPORTED_XML_TAGS.get(ancestor.tag)
            if reason is not None:
                if header_footer_kind is not None:
                    return (
                        f"This {header_footer_kind} paragraph is inside "
                        f"{reason.lower()} that DocSync cannot safely rewrite. "
                        "The content will remain unchanged."
                    )
                return f"{reason} is read-only in Edit mode."
            ancestor = ancestor.getparent()
    if cell is not None:
        if getattr(cell, "tables", None):
            return (
                "This table cell uses a nested structure that DocSync cannot "
                "safely edit. The content will remain unchanged in the "
                "generated document."
            )
        if _cell_uses_merged_structure(cell) and not allow_multiple_paragraphs:
            return (
                "This table cell uses a merged structure that DocSync cannot "
                "safely edit. The content will remain unchanged in the "
                "generated document."
            )
        try:
            for node in cell._tc.iter():
                reason = UNSUPPORTED_XML_TAGS.get(node.tag)
                if reason is not None:
                    return (
                        f"This table cell contains {reason.lower()} that DocSync "
                        "cannot safely edit. The content will remain unchanged in "
                        "the generated document."
                    )
        except AttributeError:
            return (
                "This table cell cannot be mapped safely. The content will remain "
                "unchanged in the generated document."
            )
        non_empty = [
            item for item in getattr(cell, "paragraphs", []) if item.text.strip()
        ]
        if not allow_multiple_paragraphs and len(non_empty) > 1:
            return (
                "Table cells containing multiple paragraphs are read-only in Edit mode."
            )
    return None


def _revision_values(
    document: DocxDocument,
    element: DocumentElement,
    *,
    shared_state: str = "shared",
    paragraphs: list[Paragraph] | None = None,
    tables: list[object] | None = None,
    style_name_cache: dict[str | None, str | None] | None = None,
    header_footer_parts: dict[str, dict] | None = None,
) -> dict:
    ordinal, location = _element_location(element)
    element_type = _element_type(element.style_name)
    if element_type in {"header_paragraph", "footer_paragraph"}:
        parts = (
            header_footer_parts
            if header_footer_parts is not None
            else _header_footer_part_map(document)
        )
        location = _enrich_header_footer_location(location, parts)
    else:
        parts = header_footer_parts
    style_name = _display_style_name(element.style_name)
    paragraph, cell = _paragraph_for_element(
        document,
        element,
        paragraphs=paragraphs,
        tables=tables,
        header_footer_parts=parts,
    )
    if element_type in {
        "table_paragraph",
        "header_paragraph",
        "footer_paragraph",
    } and paragraph is not None:
        style_name = _paragraph_style_name(paragraph, style_name_cache)
    list_type, list_level = (
        _numbering_type(document, paragraph, style_name)
        if paragraph is not None
        else (None, None)
    )
    if element_type in {"paragraph", "heading"} and list_type is not None:
        element_type = "list_item"
    alignment = _alignment_name(paragraph)
    block_attributes = _paragraph_block_attributes(
        element_type,
        paragraph,
        list_type,
        list_level,
        alignment,
        style_name,
    )
    reason = _unsupported_reason(
        paragraph,
        cell,
        allow_multiple_paragraphs=element_type == "table_paragraph",
        header_footer_kind=(
            "header"
            if element_type == "header_paragraph"
            else "footer"
            if element_type == "footer_paragraph"
            else None
        ),
    )
    if (
        reason is None
        and paragraph is None
        and element_type in {"header_paragraph", "footer_paragraph"}
    ):
        reason = (
            "This header or footer paragraph location could not be validated "
            "against the immutable DOCX version. The content remains unchanged."
        )
    formatted_runs = [
        {
            "text": run.text,
            **_run_attributes(run),
        }
        for run in (paragraph.runs if paragraph is not None else [])
        if run.text
    ]
    formatting = {"style_name": style_name, "runs": formatted_runs}
    delta = _paragraph_delta(
        element.text,
        paragraph,
        block_attributes,
        formatted_runs,
    )
    structure = {
        "element_type": element_type,
        "context": (
            "table"
            if element_type == "table_paragraph"
            else "header_footer"
            if element_type in {"header_paragraph", "footer_paragraph"}
            else "body"
        ),
        "matching_context": _match_context_identity(element_type, location),
        "style_name": style_name,
        "list_type": list_type,
        "list_level": list_level,
        "alignment": alignment,
    }
    normalized = normalise_editor_text(element.text)
    return {
        "id": new_id(),
        "element_id": element.id,
        "document_id": element.document_id,
        "ordinal": ordinal,
        "element_type": element_type,
        "text": element.text,
        "normalized_text": normalized,
        "exact_match_hash": exact_match_hash(
            element_type,
            element.text,
            _match_context_identity(element_type, location),
        ),
        "structure_hash": _canonical_hash(structure),
        "delta_json": delta,
        "formatting_json": formatting,
        "list_type": list_type,
        "list_level": list_level,
        "alignment": alignment,
        "location_json": location,
        "shared_state": "detached" if not normalized else shared_state,
        "supported": reason is None,
        "unsupported_reason": reason,
    }


def _load_docx(path: Path) -> DocxDocument:
    try:
        return Document(path)
    except Exception as exc:
        raise HTTPException(
            status_code=422,
            detail=f"{path.name}: the DOCX could not be opened for structured editing.",
        ) from exc


def _create_revisions_for_existing_elements(
    session: Session,
    version: DocumentVersion,
    document: DocumentRecord,
    source_path: Path,
) -> list[DocumentBlockRevision]:
    docx = _load_docx(source_path)
    paragraphs = list(docx.paragraphs)
    tables = list(docx.tables)
    header_footer_parts = _header_footer_part_map(docx)
    style_name_cache: dict[str | None, str | None] = {}
    elements = list(
        session.scalars(
            select(DocumentElement)
            .where(DocumentElement.document_id == document.id)
            .order_by(DocumentElement.paragraph_index, DocumentElement.id)
        )
    )
    revisions = [
        DocumentBlockRevision(
            version_id=version.id,
            **_revision_values(
                docx,
                element,
                paragraphs=paragraphs,
                tables=tables,
                style_name_cache=style_name_cache,
                header_footer_parts=header_footer_parts,
            ),
        )
        for element in elements
    ]
    session.add_all(revisions)
    return revisions


def initialise_original_version(
    session: Session,
    document: DocumentRecord,
    *,
    parsed_document: DocxDocument | None = None,
    elements: list[DocumentElement] | None = None,
) -> DocumentVersion:
    if parsed_document is not None and elements is not None:
        paragraphs = list(parsed_document.paragraphs)
        tables = list(parsed_document.tables)
        style_name_cache: dict[str | None, str | None] = {}
        header_footer_parts = _header_footer_part_map(parsed_document)
        version = DocumentVersion(
            id=document.id,
            document=document,
            parent_version_id=None,
            generation_id=None,
            editor_operation_id=None,
            version_number=1,
            storage_area="originals",
            storage_name=document.stored_name,
            download_name=document.original_name,
            checksum_sha256=document.checksum_sha256,
            created_at=document.created_at,
        )
        head = DocumentHead(
            document=document,
            current_version=version,
            revision=1,
        )
        revisions = [
            DocumentBlockRevision(
                version=version,
                **_revision_values(
                    parsed_document,
                    element,
                    paragraphs=paragraphs,
                    tables=tables,
                    style_name_cache=style_name_cache,
                    header_footer_parts=header_footer_parts,
                ),
            )
            for element in elements
        ]
        session.add_all([version, head, *revisions])
        return version

    version = session.get(DocumentVersion, document.id)
    if version is None:
        version = DocumentVersion(
            id=document.id,
            document_id=document.id,
            parent_version_id=None,
            generation_id=None,
            editor_operation_id=None,
            version_number=1,
            storage_area="originals",
            storage_name=document.stored_name,
            download_name=document.original_name,
            checksum_sha256=document.checksum_sha256,
            created_at=document.created_at,
        )
        session.add(version)
        session.flush()
    head = session.get(DocumentHead, document.id)
    if head is None:
        head = DocumentHead(
            document_id=document.id,
            current_version_id=version.id,
            revision=max(version.version_number, 1),
        )
        session.add(head)
        session.flush()
    revision_count = int(
        session.scalar(
            select(func.count(DocumentBlockRevision.id)).where(
                DocumentBlockRevision.version_id == version.id
            )
        )
        or 0
    )
    if revision_count == 0:
        _create_revisions_for_existing_elements(
            session,
            version,
            document,
            _safe_storage_path("originals", document.stored_name),
        )
        session.flush()
    return version


def prepare_original_version_rows(
    document_row: dict,
    parsed_document: DocxDocument,
    element_rows: list[dict],
) -> tuple[dict, dict, list[dict]]:
    """Prepare immutable initial-version mappings without ORM identity tracking."""

    paragraphs = list(parsed_document.paragraphs)
    tables = list(parsed_document.tables)
    style_name_cache: dict[str | None, str | None] = {}
    header_footer_parts = _header_footer_part_map(parsed_document)
    created_at = document_row["created_at"]
    version_row = {
        "id": document_row["id"],
        "document_id": document_row["id"],
        "parent_version_id": None,
        "generation_id": None,
        "editor_operation_id": None,
        "version_number": 1,
        "storage_area": "originals",
        "storage_name": document_row["stored_name"],
        "download_name": document_row["original_name"],
        "checksum_sha256": document_row["checksum_sha256"],
        "created_at": created_at,
    }
    head_row = {
        "document_id": document_row["id"],
        "current_version_id": document_row["id"],
        "revision": 1,
        "updated_at": created_at,
    }
    revision_rows = []
    for element_row in element_rows:
        element = SimpleNamespace(**element_row)
        revision_rows.append(
            {
                "version_id": document_row["id"],
                "created_at": created_at,
                **_revision_values(
                    parsed_document,
                    element,
                    paragraphs=paragraphs,
                    tables=tables,
                    style_name_cache=style_name_cache,
                    header_footer_parts=header_footer_parts,
                ),
            }
        )
    return version_row, head_row, revision_rows


def _serialize_revision(
    revision: DocumentBlockRevision,
    *,
    document_name: str | None = None,
) -> dict:
    location = dict(revision.location_json or {})
    payload = {
        "id": revision.element_id,
        "element_id": revision.element_id,
        "document_id": revision.document_id,
        "document_name": document_name,
        "version_id": revision.version_id,
        "element_type": revision.element_type,
        "type": revision.element_type,
        "order": revision.ordinal,
        "ordinal": revision.ordinal,
        "paragraph_index": location.get("paragraph_index", revision.ordinal),
        "page_number": max(1, revision.ordinal // 18 + 1),
        "text": revision.text,
        "normalized_text": revision.normalized_text,
        "exact_match_hash": revision.exact_match_hash,
        "structure_hash": revision.structure_hash,
        "delta": revision.delta_json or {"ops": [{"insert": revision.text}, {"insert": "\n"}]},
        "formatting": revision.formatting_json or {},
        "style_name": (revision.formatting_json or {}).get("style_name"),
        "list_type": revision.list_type,
        "list_level": revision.list_level,
        "indent": revision.list_level,
        "alignment": revision.alignment,
        "location": location,
        "shared_state": revision.shared_state,
        "detached_from_shared": revision.shared_state == "detached",
        "supported": revision.supported,
        "read_only": not revision.supported,
        "unsupported_reason": revision.unsupported_reason,
    }
    payload.update(
        {
            key: location[key]
            for key in (
                "document_order",
                "table_index",
                "row_index",
                "column_index",
                "paragraph_index",
                "kind",
                "section_index",
                "source_section_index",
                "header_footer_type",
                "part_relationship_id",
                "is_linked_to_previous",
                "section_indexes",
                "linked_section_indexes",
                "linked_sections",
                "affected_header_footer_types",
            )
            if key in location
        }
    )
    return payload


def _document_diagnostics(path: Path) -> list[dict]:
    document = _load_docx(path)
    diagnostics: list[dict] = []
    seen: set[tuple[str, str]] = set()
    for node in document.element.body.iter():
        reason = UNSUPPORTED_XML_TAGS.get(node.tag)
        if reason is None:
            continue
        key = (node.tag, reason)
        if key in seen:
            continue
        seen.add(key)
        diagnostics.append(
            {
                "kind": "unsupported_ooxml",
                "element_type": "unsupported",
                "label": reason,
                "reason": f"{reason} is preserved in Layout mode and is read-only.",
                "read_only": True,
            }
        )
    return diagnostics


def serialize_editor_content(session: Session, version_id: str) -> dict:
    version = get_version_or_404(session, version_id)
    head = session.get(DocumentHead, version.document_id)
    blocks = list(
        session.scalars(
            select(DocumentBlockRevision)
            .where(DocumentBlockRevision.version_id == version.id)
            .order_by(DocumentBlockRevision.ordinal, DocumentBlockRevision.id)
        )
    )
    diagnostics = [
        {
            "id": block.element_id,
            "kind": "unsupported_block",
            "element_type": block.element_type,
            "reason": block.unsupported_reason
            or "This Word block is preserved but read-only in Edit mode.",
            "location": _location_key(block.location_json),
            "read_only": True,
        }
        for block in blocks
        if not block.supported
    ]
    return {
        "document_id": version.document_id,
        "document_set_id": version.document.document_set_id,
        "document_name": version.document.original_name,
        "version_id": version.id,
        "version_number": version.version_number,
        "parent_version_id": version.parent_version_id,
        "current_version": bool(head and head.current_version_id == version.id),
        "is_current": bool(head and head.current_version_id == version.id),
        "created_at": utc_isoformat(version.created_at),
        "blocks": [
            _serialize_revision(block, document_name=version.document.original_name)
            for block in blocks
        ],
        "unsupported": diagnostics,
        "diagnostics": diagnostics,
        "unsupported_count": len(diagnostics),
        "notice": (
            "Unsupported Word objects remain preserved in Layout mode and are "
            "read-only in Edit mode."
            if diagnostics or any(not block.supported for block in blocks)
            else None
        ),
    }


def serialize_version_document_view(session: Session, version_id: str) -> dict:
    """Compatibility structured-view payload backed by immutable block revisions."""

    content = serialize_editor_content(session, version_id)
    pages: list[dict] = []
    current: list[dict] = []
    used_units = 0
    for block in content["blocks"]:
        units = max(1, math.ceil(len(block["text"]) / 120))
        if block["element_type"] == "heading":
            units += 1
        if current and used_units + units > 18:
            pages.append({"page_number": len(pages) + 1, "elements": current})
            current = []
            used_units = 0
        item = {
            **block,
            "page_number": len(pages) + 1,
        }
        current.append(item)
        used_units += units
    if current or not pages:
        pages.append({"page_number": len(pages) + 1, "elements": current})
    return {
        "document_id": content["document_id"],
        "version_id": content["version_id"],
        "document_set_id": content["document_set_id"],
        "document_name": content["document_name"],
        "render_status": "ready",
        "render_mode": "structured",
        "pagination": "estimated",
        "page_count": len(pages),
        "notice": (
            "Structured browser preview backed by an immutable document version. "
            "Unsupported Word objects remain available in Layout mode and are read-only."
        ),
        "pages": pages,
        "layout_regions": [],
        "unsupported": content["unsupported"],
    }


def serialize_document_versions(session: Session, document_id: str) -> dict:
    document = session.get(DocumentRecord, document_id)
    if document is None:
        raise HTTPException(status_code=404, detail="Document not found.")
    initialise_original_version(session, document)
    session.flush()
    head = session.get(DocumentHead, document.id)
    versions = list(
        session.scalars(
            select(DocumentVersion)
            .where(DocumentVersion.document_id == document.id)
            .options(selectinload(DocumentVersion.editor_operation))
            .order_by(DocumentVersion.version_number.desc())
        )
    )
    return {
        "document_id": document.id,
        "document_name": document.original_name,
        "current_version_id": head.current_version_id if head else document.id,
        "versions": [
            {
                "id": version.id,
                "version_id": version.id,
                "version_number": version.version_number,
                "document_id": version.document_id,
                "parent_version_id": version.parent_version_id,
                "created_at": utc_isoformat(version.created_at),
                "checksum_sha256": version.checksum_sha256,
                "status": "completed",
                "is_current": bool(head and head.current_version_id == version.id),
                "generation_id": version.generation_id
                or version.editor_operation_id,
                "operation_type": (
                    version.editor_operation.operation_type
                    if version.editor_operation is not None
                    else None
                ),
                "restored_from_version_id": (
                    (version.editor_operation.preview_json or {}).get(
                        "restored_from_version_id"
                    )
                    if version.editor_operation is not None
                    and version.editor_operation.operation_type == "version_restore"
                    and isinstance(version.editor_operation.preview_json, dict)
                    else None
                ),
                "restored_from_version_number": (
                    (version.editor_operation.preview_json or {}).get(
                        "restored_from_version_number"
                    )
                    if version.editor_operation is not None
                    and version.editor_operation.operation_type == "version_restore"
                    and isinstance(version.editor_operation.preview_json, dict)
                    else None
                ),
                "download_url": f"/api/document-versions/{version.id}/download",
                "editor_content_url": (
                    f"/api/document-versions/{version.id}/editor-content"
                ),
            }
            for version in versions
        ],
    }


def current_version_summary_map(
    session: Session,
    document_ids: list[str],
) -> dict[str, dict]:
    if not document_ids:
        return {}
    rows = session.execute(
        select(DocumentHead, DocumentVersion)
        .join(
            DocumentVersion,
            DocumentVersion.id == DocumentHead.current_version_id,
        )
        .where(DocumentHead.document_id.in_(document_ids))
    ).all()
    return {
        head.document_id: {
            "version_id": version.id,
            "version_number": version.version_number,
            "parent_version_id": version.parent_version_id,
            "checksum_sha256": version.checksum_sha256,
        }
        for head, version in rows
    }


def _get_current_revision_or_404(
    session: Session,
    element_id: str,
) -> tuple[DocumentBlockRevision, DocumentRecord, DocumentHead]:
    element = session.get(DocumentElement, element_id)
    if element is None:
        raise HTTPException(status_code=404, detail="Document element not found.")
    document = session.get(DocumentRecord, element.document_id)
    if document is None:
        raise HTTPException(status_code=404, detail="Document not found.")
    head = session.get(DocumentHead, document.id)
    if head is None:
        initialise_original_version(session, document)
        session.flush()
        head = session.get(DocumentHead, document.id)
    revision = session.scalar(
        select(DocumentBlockRevision).where(
            DocumentBlockRevision.version_id == head.current_version_id,
            DocumentBlockRevision.element_id == element_id,
        )
    )
    if revision is None:
        raise HTTPException(
            status_code=409,
            detail="The selected block is no longer part of the current document version.",
        )
    return revision, document, head


def _latest_decisions(
    session: Session,
    source: DocumentBlockRevision,
) -> dict[str, MatchDecision]:
    rows = list(
        session.scalars(
            select(MatchDecision)
            .where(
                MatchDecision.source_version_id == source.version_id,
                MatchDecision.source_element_id == source.element_id,
            )
            .order_by(MatchDecision.updated_at.desc(), MatchDecision.id.desc())
        )
    )
    return {row.candidate_element_id: row for row in rows}


def _position_similarity(
    source: DocumentBlockRevision,
    candidate: DocumentBlockRevision,
    maximum_ordinal: int,
) -> float:
    return max(
        0.0,
        1.0 - abs(source.ordinal - candidate.ordinal) / max(maximum_ordinal, 1),
    )


def _neighbor_context(
    session: Session,
    revision: DocumentBlockRevision,
) -> str:
    neighbors = list(
        session.scalars(
            select(DocumentBlockRevision.normalized_text)
            .where(
                DocumentBlockRevision.version_id == revision.version_id,
                DocumentBlockRevision.ordinal.in_(
                    [max(revision.ordinal - 1, 0), revision.ordinal + 1]
                ),
            )
            .order_by(DocumentBlockRevision.ordinal)
        )
    )
    return " \u241f ".join(neighbors)


def _token_jaccard(left: str, right: str) -> float:
    left_tokens = set(left.split())
    right_tokens = set(right.split())
    if not left_tokens and not right_tokens:
        return 1.0
    union = left_tokens | right_tokens
    return len(left_tokens & right_tokens) / len(union) if union else 0.0


def similarity_score(
    source: DocumentBlockRevision,
    candidate: DocumentBlockRevision,
    *,
    maximum_ordinal: int,
    session: Session | None = None,
) -> float:
    sequence = SequenceMatcher(
        None,
        source.normalized_text,
        candidate.normalized_text,
        autojunk=False,
    ).ratio()
    tokens = _token_jaccard(source.normalized_text, candidate.normalized_text)
    position = _position_similarity(source, candidate, maximum_ordinal)
    if session is None:
        return round(0.8 * sequence + 0.15 * tokens + 0.05 * position, 6)
    source_context = _neighbor_context(session, source)
    candidate_context = _neighbor_context(session, candidate)
    neighbor = (
        SequenceMatcher(
            None,
            source_context,
            candidate_context,
            autojunk=False,
        ).ratio()
        if source_context or candidate_context
        else 1.0
    )
    return round(
        0.72 * sequence + 0.14 * tokens + 0.04 * position + 0.10 * neighbor,
        6,
    )


def difference_spans(source_text: str, candidate_text: str) -> list[dict]:
    source_tokens = TOKEN_PATTERN.findall(source_text)
    candidate_tokens = TOKEN_PATTERN.findall(candidate_text)
    matcher = SequenceMatcher(None, source_tokens, candidate_tokens, autojunk=False)
    spans: list[dict] = []
    for operation, left_start, left_end, right_start, right_end in matcher.get_opcodes():
        source_value = "".join(source_tokens[left_start:left_end])
        candidate_value = "".join(candidate_tokens[right_start:right_end])
        kind = {
            "equal": "equal",
            "delete": "delete",
            "insert": "insert",
            "replace": "changed",
        }[operation]
        spans.append(
            {
                "kind": kind,
                "text": candidate_value if candidate_value else source_value,
                "source_text": source_value,
                "candidate_text": candidate_value,
                "source_start": left_start,
                "source_end": left_end,
                "candidate_start": right_start,
                "candidate_end": right_end,
            }
        )
    return spans


def _current_candidate_revisions(
    session: Session,
    source: DocumentBlockRevision,
    document_set_id: str,
) -> list[tuple[DocumentBlockRevision, DocumentRecord]]:
    rows = list(
        session.execute(
            select(DocumentBlockRevision, DocumentRecord)
            .join(
                DocumentHead,
                DocumentHead.current_version_id == DocumentBlockRevision.version_id,
            )
            .join(
                DocumentRecord,
                DocumentRecord.id == DocumentBlockRevision.document_id,
            )
            .where(
                DocumentRecord.document_set_id == document_set_id,
                DocumentBlockRevision.element_type == source.element_type,
                DocumentBlockRevision.element_id != source.element_id,
                DocumentBlockRevision.supported.is_(True),
            )
            .order_by(
                DocumentRecord.original_name,
                DocumentBlockRevision.ordinal,
            )
        )
    )
    return [
        (candidate, record)
        for candidate, record in rows
        if _compatible_match_context(source, candidate)
    ]


def get_similar_matches(
    session: Session,
    element_id: str,
    *,
    threshold: float | None = None,
    limit: int | None = None,
) -> dict:
    source, document, _head = _get_current_revision_or_404(session, element_id)
    effective_threshold = (
        settings.near_match_threshold if threshold is None else threshold
    )
    if not 0 <= effective_threshold <= 1:
        raise HTTPException(status_code=422, detail="Similarity threshold must be 0–1.")
    effective_limit = min(
        limit or settings.near_match_candidate_limit,
        settings.near_match_candidate_limit,
    )
    decisions = _latest_decisions(session, source)
    candidates = _current_candidate_revisions(
        session,
        source,
        document.document_set_id,
    )
    maximum_ordinal = max(
        [source.ordinal, *(candidate.ordinal for candidate, _record in candidates)],
        default=1,
    )
    matches: list[dict] = []
    for candidate, candidate_document in candidates[
        : settings.near_match_candidate_limit * 4
    ]:
        if candidate.exact_match_hash == source.exact_match_hash:
            continue
        decision = decisions.get(candidate.element_id)
        if (
            candidate.shared_state == "detached"
            and (decision is None or decision.decision != "confirmed")
        ):
            continue
        score = similarity_score(
            source,
            candidate,
            maximum_ordinal=maximum_ordinal,
            session=session,
        )
        if score < effective_threshold:
            continue
        spans = difference_spans(source.text, candidate.text)
        item = {
            **_serialize_revision(
                candidate,
                document_name=candidate_document.original_name,
            ),
            "match_type": "near",
            "similarity_score": score,
            "score": score,
            "difference_spans": spans,
            "diff_spans": spans,
            "decision": decision.decision if decision is not None else "pending",
        }
        matches.append(item)
    matches.sort(
        key=lambda item: (
            -item["similarity_score"],
            item["document_name"].casefold(),
            item["ordinal"],
        )
    )
    matches = matches[:effective_limit]
    return {
        "source_element_id": source.element_id,
        "source": _serialize_revision(source, document_name=document.original_name),
        "threshold": effective_threshold,
        "algorithm_version": MATCH_ALGORITHM_VERSION,
        "matches": matches,
        "similar_matches": matches,
        "count": len(matches),
    }


def get_editor_matches(session: Session, element_id: str) -> dict:
    # Preserve the legacy exact-link-group shape while enriching it with
    # version/hash metadata used by the editor.
    from .document_service import get_element_matches_or_404

    legacy = get_element_matches_or_404(session, element_id)
    source, document, _head = _get_current_revision_or_404(session, element_id)
    rows = _current_candidate_revisions(session, source, document.document_set_id)
    exact = [
        {
            **_serialize_revision(candidate, document_name=record.original_name),
            "match_type": "exact",
            "similarity_score": 1.0,
            "score": 1.0,
            "difference_spans": difference_spans(source.text, candidate.text),
        }
        for candidate, record in rows
        if candidate.exact_match_hash == source.exact_match_hash
        and source.shared_state == "shared"
        and candidate.shared_state == "shared"
    ]
    serialized_source = _serialize_revision(
        source,
        document_name=document.original_name,
    )
    legacy["source"].update(serialized_source)
    if legacy["link_group"] is not None:
        if exact:
            safe_members = [serialized_source, *exact]
            legacy["link_group"]["members"] = safe_members
            legacy["link_group"]["member_count"] = len(safe_members)
            legacy["link_group"]["document_count"] = len(
                {member["document_id"] for member in safe_members}
            )
        else:
            # Existing workspaces may contain a legacy text/style group that
            # predates revision-aware structural matching. Do not expose that
            # incompatible group as an editable exact match.
            legacy["link_group"] = None
    legacy["exact_matches"] = exact
    legacy["matches"] = exact
    legacy["exact_match_count"] = len(exact)
    return legacy


def compare_elements(
    session: Session,
    element_id: str,
    request: CompareRequest,
) -> dict:
    source, document, _head = _get_current_revision_or_404(session, element_id)
    items: list[dict] = []
    if request.candidate_element_ids is not None:
        for candidate_id in request.candidate_element_ids:
            candidate, candidate_document, _candidate_head = (
                _get_current_revision_or_404(session, candidate_id)
            )
            if candidate_document.document_set_id != document.document_set_id:
                raise HTTPException(
                    status_code=422,
                    detail="Comparison targets must belong to the same document set.",
                )
            if not _compatible_match_context(source, candidate):
                raise HTTPException(
                    status_code=422,
                    detail="Comparison targets must use a compatible block type.",
                )
            score = similarity_score(
                source,
                candidate,
                maximum_ordinal=max(source.ordinal, candidate.ordinal, 1),
                session=session,
            )
            spans = difference_spans(source.text, candidate.text)
            items.append(
                {
                    **_serialize_revision(
                        candidate,
                        document_name=candidate_document.original_name,
                    ),
                    "match_type": (
                        "exact"
                        if candidate.exact_match_hash == source.exact_match_hash
                        else "near"
                    ),
                    "similarity_score": score,
                    "score": score,
                    "difference_spans": spans,
                    "diff_spans": spans,
                }
            )
    else:
        exact_payload = get_editor_matches(session, element_id)
        items.extend(exact_payload["exact_matches"])
        if request.include_near_matches:
            near = get_similar_matches(
                session,
                element_id,
                threshold=request.threshold,
                limit=request.limit,
            )
            items.extend(near["matches"])
    items.sort(
        key=lambda item: (
            0 if item["match_type"] == "exact" else 1,
            -item["similarity_score"],
            (item.get("document_name") or "").casefold(),
        )
    )
    items = items[: request.limit]
    return {
        "source_element_id": source.element_id,
        "source": _serialize_revision(source, document_name=document.original_name),
        "items": items,
        "matches": items,
        "shared_spans": [
            span
            for item in items
            for span in item["difference_spans"]
            if span["kind"] == "equal"
        ],
    }


def _upsert_decision(
    session: Session,
    *,
    document_set_id: str,
    source: DocumentBlockRevision,
    candidate: DocumentBlockRevision,
    decision: str,
) -> MatchDecision:
    score = similarity_score(
        source,
        candidate,
        maximum_ordinal=max(source.ordinal, candidate.ordinal, 1),
        session=session,
    )
    spans = difference_spans(source.text, candidate.text)
    row = session.scalar(
        select(MatchDecision).where(
            MatchDecision.source_version_id == source.version_id,
            MatchDecision.source_element_id == source.element_id,
            MatchDecision.candidate_version_id == candidate.version_id,
            MatchDecision.candidate_element_id == candidate.element_id,
        )
    )
    if row is None:
        row = MatchDecision(
            id=new_id(),
            document_set_id=document_set_id,
            source_element_id=source.element_id,
            candidate_element_id=candidate.element_id,
            source_version_id=source.version_id,
            candidate_version_id=candidate.version_id,
            decision=decision,
            similarity_score=score,
            algorithm_version=MATCH_ALGORITHM_VERSION,
            difference_json=spans,
        )
        session.add(row)
    else:
        row.decision = decision
        row.similarity_score = score
        row.algorithm_version = MATCH_ALGORITHM_VERSION
        row.difference_json = spans
        row.updated_at = utc_now()
    return row


def save_match_decisions(
    session: Session,
    element_id: str,
    request: MatchDecisionBatchRequest,
) -> dict:
    source, document, _head = _get_current_revision_or_404(session, element_id)
    rows: list[MatchDecision] = []
    for item in request.decisions:
        candidate, candidate_document, _candidate_head = _get_current_revision_or_404(
            session,
            item.candidate_element_id,
        )
        if candidate_document.document_set_id != document.document_set_id:
            raise HTTPException(
                status_code=422,
                detail="Match decisions must stay within one document set.",
            )
        if not _compatible_match_context(source, candidate):
            raise HTTPException(
                status_code=422,
                detail="Match decisions require compatible block types.",
            )
        candidate_score = similarity_score(
            source,
            candidate,
            maximum_ordinal=max(source.ordinal, candidate.ordinal, 1),
            session=session,
        )
        if (
            item.status == "confirmed"
            and candidate.exact_match_hash != source.exact_match_hash
            and candidate_score < settings.near_match_threshold
        ):
            raise HTTPException(
                status_code=422,
                detail=(
                    "This candidate is below the configured near-match review "
                    "threshold and cannot be confirmed."
                ),
            )
        rows.append(
            _upsert_decision(
                session,
                document_set_id=document.document_set_id,
                source=source,
                candidate=candidate,
                decision=item.status,
            )
        )
    session.commit()
    return {
        "saved": True,
        "source_element_id": source.element_id,
        "decisions": [
            {
                "candidate_element_id": row.candidate_element_id,
                "element_id": row.candidate_element_id,
                "status": row.decision,
                "decision": row.decision,
                "similarity_score": row.similarity_score,
            }
            for row in rows
        ],
    }


def _delta_visible_text(delta: QuillDelta) -> str:
    operations = delta.ops
    for operation in operations:
        if operation.retain is not None or operation.delete is not None:
            raise HTTPException(
                status_code=422,
                detail=(
                    "Structural Delta retain/delete operations are unsupported. "
                    "Submit the complete contents of one existing block."
                ),
            )
    inserted = "".join(operation.insert or "" for operation in operations)
    normalized = inserted.replace("\r\n", "\n").replace("\r", "\n")
    if normalized.endswith("\n"):
        normalized = normalized[:-1]
    return normalized


def _validate_delta(target: EditorTarget) -> dict | None:
    if target.delta is None:
        return None
    visible = _delta_visible_text(target.delta)
    if visible != target.replacement_text:
        raise HTTPException(
            status_code=422,
            detail=(
                "Delta text must match replacement_text for the targeted block."
            ),
        )
    operations = target.delta.model_dump(exclude_none=True)["ops"]
    for operation in operations:
        attributes = set((operation.get("attributes") or {}).keys())
        allowed = (
            ALLOWED_BLOCK_ATTRIBUTES
            if operation.get("insert") == "\n"
            else ALLOWED_INLINE_ATTRIBUTES
        )
        unsupported = attributes - allowed
        if unsupported:
            raise HTTPException(
                status_code=422,
                detail=(
                    "Unsupported Delta formatting attribute(s): "
                    + ", ".join(sorted(unsupported))
                    + "."
                ),
            )
    return {"ops": operations}


def _delta_with_replacement(
    revision: DocumentBlockRevision,
    replacement_text: str,
) -> dict:
    current = revision.delta_json or {"ops": [{"insert": revision.text}, {"insert": "\n"}]}
    operations = [dict(item) for item in current.get("ops", [])]
    newline = (
        operations[-1]
        if operations and operations[-1].get("insert") == "\n"
        else {"insert": "\n"}
    )
    first_attributes = None
    for operation in operations:
        if operation.get("insert") != "\n":
            first_attributes = operation.get("attributes")
            break
    first: dict[str, object] = {"insert": replacement_text}
    if first_attributes:
        first["attributes"] = first_attributes
    return {"ops": [first, newline]}


def _stored_confirmation(
    session: Session,
    source: DocumentBlockRevision,
    candidate: DocumentBlockRevision,
) -> bool:
    decision = session.scalar(
        select(MatchDecision.decision).where(
            MatchDecision.source_version_id == source.version_id,
            MatchDecision.source_element_id == source.element_id,
            MatchDecision.candidate_version_id == candidate.version_id,
            MatchDecision.candidate_element_id == candidate.element_id,
        )
    )
    return decision == "confirmed"


def _validate_editor_request(
    session: Session,
    document_set_id: str,
    request: EditorEditRequest,
) -> tuple[
    DocumentBlockRevision,
    DocumentRecord,
    dict[str, tuple[DocumentRecord, DocumentHead, DocumentVersion]],
    list[tuple[EditorTarget, DocumentBlockRevision, DocumentRecord, dict]],
]:
    document_set = session.get(DocumentSet, document_set_id)
    if document_set is None:
        raise HTTPException(status_code=404, detail="Document set not found.")

    base_context: dict[
        str, tuple[DocumentRecord, DocumentHead, DocumentVersion]
    ] = {}
    for document_id, expected_version_id in request.base_versions.items():
        document = session.get(DocumentRecord, document_id)
        if document is None or document.document_set_id != document_set_id:
            raise HTTPException(
                status_code=422,
                detail="Every base version must belong to the selected document set.",
            )
        head = session.get(DocumentHead, document.id)
        if head is None:
            initialise_original_version(session, document)
            session.flush()
            head = session.get(DocumentHead, document.id)
        if head.current_version_id != expected_version_id:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=(
                    f"{document.original_name} changed after this edit was opened. "
                    "Reload the current version and review the edit again."
                ),
            )
        base_context[document.id] = (
            document,
            head,
            get_version_or_404(session, head.current_version_id),
        )

    source, source_document, _source_head = _get_current_revision_or_404(
        session,
        request.source_element_id,
    )
    if source_document.document_set_id != document_set_id:
        raise HTTPException(
            status_code=422,
            detail="The source block does not belong to this document set.",
        )
    if source_document.id not in base_context:
        raise HTTPException(
            status_code=422,
            detail="base_versions must include the source document.",
        )
    if not source.supported:
        raise HTTPException(
            status_code=422,
            detail=source.unsupported_reason
            or "The selected source block is read-only.",
        )
    if request.source_element_id not in {
        target.element_id for target in request.targets
    }:
        raise HTTPException(
            status_code=422,
            detail="The source block must remain included in editor targets.",
        )

    request_decisions = {
        decision.candidate_element_id: decision.status
        for decision in request.match_decisions
    }
    validated_targets: list[
        tuple[EditorTarget, DocumentBlockRevision, DocumentRecord, dict]
    ] = []
    for target in request.targets:
        revision, document, _head = _get_current_revision_or_404(
            session,
            target.element_id,
        )
        if document.document_set_id != document_set_id:
            raise HTTPException(
                status_code=422,
                detail="Every target must belong to the selected document set.",
            )
        if document.id not in base_context:
            raise HTTPException(
                status_code=422,
                detail=f"base_versions is missing {document.original_name}.",
            )
        if revision.version_id != base_context[document.id][1].current_version_id:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=f"{document.original_name} no longer contains this target version.",
            )
        if not revision.supported:
            raise HTTPException(
                status_code=422,
                detail=revision.unsupported_reason
                or f"{document.original_name} contains a read-only target.",
            )
        if not _compatible_match_context(source, revision):
            raise HTTPException(
                status_code=422,
                detail="Editor targets must use a compatible block type.",
            )
        is_source = revision.element_id == source.element_id
        is_exact = (
            revision.exact_match_hash == source.exact_match_hash
            and revision.shared_state == "shared"
            and source.shared_state == "shared"
        )
        explicitly_confirmed = (
            request_decisions.get(revision.element_id) == "confirmed"
        )
        persisted_confirmation = _stored_confirmation(session, source, revision)
        near_score = (
            similarity_score(
                source,
                revision,
                maximum_ordinal=max(source.ordinal, revision.ordinal, 1),
                session=session,
            )
            if not is_source and not is_exact
            else 1.0
        )
        if (
            (explicitly_confirmed or persisted_confirmation)
            and not is_source
            and not is_exact
            and near_score < settings.near_match_threshold
        ):
            raise HTTPException(
                status_code=422,
                detail=(
                    f"{document.original_name}: the selected block is below the "
                    "configured near-match threshold."
                ),
            )
        if not (
            is_source
            or is_exact
            or explicitly_confirmed
            or persisted_confirmation
        ):
            raise HTTPException(
                status_code=422,
                detail=(
                    f"{document.original_name}: near-match targets must be explicitly "
                    "confirmed before they can be edited."
                ),
            )
        if request_decisions.get(revision.element_id) in {"ignored", "removed"}:
            raise HTTPException(
                status_code=422,
                detail="Ignored or removed match candidates cannot be editor targets.",
            )
        delta = _validate_delta(target) or _delta_with_replacement(
            revision,
            target.replacement_text,
        )
        if revision.element_type in {
            "table_paragraph",
            "header_paragraph",
            "footer_paragraph",
        }:
            operations = list(delta.get("ops") or [])
            newline_attributes = (
                operations[-1].get("attributes") or {}
                if operations and operations[-1].get("insert") == "\n"
                else {}
            )
            if "header" in newline_attributes:
                raise HTTPException(
                    status_code=422,
                    detail=(
                        "Heading levels cannot be applied inside a table, header, "
                        "or footer paragraph. "
                        "Use bold, lists, indentation, or alignment instead."
                    ),
                )
        validated_targets.append((target, revision, document, delta))

    return source, source_document, base_context, validated_targets


def _clear_paragraph_runs(paragraph: Paragraph) -> list:
    runs = list(paragraph.runs)
    for run in runs:
        run.text = ""
    return runs


def _apply_inline_delta(paragraph: Paragraph, delta: dict) -> dict:
    operations = list(delta.get("ops") or [])
    newline_attributes: dict = {}
    if operations and operations[-1].get("insert") == "\n":
        newline_attributes = dict(operations.pop().get("attributes") or {})
    existing_runs = _clear_paragraph_runs(paragraph)
    run_index = 0
    for operation in operations:
        text = operation.get("insert")
        if not isinstance(text, str) or not text:
            continue
        run = (
            existing_runs[run_index]
            if run_index < len(existing_runs)
            else paragraph.add_run()
        )
        run_index += 1
        run.text = text
        attributes = operation.get("attributes") or {}
        run.bold = bool(attributes.get("bold", False))
        run.italic = bool(attributes.get("italic", False))
        run.underline = bool(attributes.get("underline", False))
    return newline_attributes


def _set_alignment(paragraph: Paragraph, value: object) -> None:
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if value in mapping:
        paragraph.alignment = mapping[value]


def _set_list_format(
    document: DocxDocument,
    paragraph: Paragraph,
    list_type: object,
    level: object,
    *,
    existing_list_type: str | None = None,
    style_names: set[str] | None = None,
) -> None:
    if list_type not in {"ordered", "bullet"}:
        return

    level_value = max(0, int(level or 0))

    # Preserve an existing custom or localised Word list style.
    if existing_list_type == list_type:
        properties = paragraph._p.get_or_add_pPr()
        num_properties = properties.numPr

        if num_properties is not None:
            level_node = num_properties.get_or_add_ilvl()
            level_node.val = level_value

        return

    base_style = "List Number" if list_type == "ordered" else "List Bullet"
    requested_style = (
        f"{base_style} {level_value + 1}"
        if level_value
        else base_style
    )

    style_names = style_names or {style.name for style in document.styles}

    selected_style = next(
        (
            style_name
            for style_name in (requested_style, base_style)
            if style_name in style_names
        ),
        None,
    )

    if selected_style is None:
        raise HTTPException(
            status_code=422,
            detail=(
                "This document does not contain a compatible Word list style. "
                "The existing custom list formatting was preserved, but DocSync "
                "cannot change this paragraph to a different list type."
            ),
        )

    paragraph.style = selected_style

    properties = paragraph._p.get_or_add_pPr()
    num_properties = properties.numPr

    if num_properties is not None:
        level_node = num_properties.get_or_add_ilvl()
        level_node.val = level_value
        

def _clear_numbering(paragraph: Paragraph) -> None:
    properties = paragraph._p.pPr
    if properties is not None and properties.numPr is not None:
        properties.remove(properties.numPr)


def _apply_target_to_docx(
    document: DocxDocument,
    revision: DocumentBlockRevision,
    target: EditorTarget,
    delta: dict,
    *,
    paragraphs: list[Paragraph] | None = None,
    tables: list | None = None,
    header_footer_parts: dict[str, dict] | None = None,
    style_names: set[str] | None = None,
) -> None:
    location = revision.location_json or {}
    paragraph: Paragraph
    if location.get("kind") in {"header_paragraph", "footer_paragraph"}:
        parts = header_footer_parts or _header_footer_part_map(document)
        relationship_id = str(location.get("part_relationship_id") or "")
        entry = parts.get(relationship_id)
        if entry is None:
            raise HTTPException(
                status_code=409,
                detail=(
                    "The target header or footer part is no longer present in "
                    "the selected immutable version."
                ),
            )
        expected_kind = str(location.get("kind")).removesuffix("_paragraph")
        expected_type = str(location.get("header_footer_type") or "")
        expected_source = int(location.get("source_section_index", -1))
        if not any(
            usage["kind"] == expected_kind
            and usage["header_footer_type"] == expected_type
            and int(usage["source_section_index"]) == expected_source
            for usage in entry["usages"]
        ):
            raise HTTPException(
                status_code=409,
                detail=(
                    "The target section and header/footer relationship could not "
                    "be validated against the selected immutable version."
                ),
            )
        try:
            paragraph = entry["paragraphs"][int(location["paragraph_index"])]
        except (IndexError, KeyError, TypeError, ValueError) as exc:
            raise HTTPException(
                status_code=409,
                detail="The target header or footer paragraph is no longer present.",
            ) from exc
        reason = _unsupported_reason(paragraph, None, header_footer_kind=expected_kind)
        if reason is not None:
            raise HTTPException(status_code=422, detail=reason)
    elif location.get("kind") in {"table_cell", "table_paragraph"}:
        try:
            cell = (
                (tables or list(document.tables))[int(location["table_index"])]
                .rows[int(location["row_index"])]
                .cells[int(location["column_index"])]
            )
        except (IndexError, KeyError, TypeError, ValueError) as exc:
            raise HTTPException(
                status_code=409,
                detail="The target table-paragraph location is no longer valid.",
            ) from exc
        if location.get("kind") == "table_paragraph":
            try:
                paragraph = cell.paragraphs[int(location["paragraph_index"])]
            except (IndexError, KeyError, TypeError, ValueError) as exc:
                raise HTTPException(
                    status_code=409,
                    detail="The target table paragraph is no longer present.",
                ) from exc
        else:
            non_empty = [item for item in cell.paragraphs if item.text.strip()]
            paragraph = non_empty[0] if non_empty else cell.paragraphs[0]
        reason = _unsupported_reason(
            paragraph,
            cell,
            allow_multiple_paragraphs=location.get("kind") == "table_paragraph",
        )
        if reason is not None:
            raise HTTPException(
                status_code=422,
                detail=reason,
            )
    else:
        try:
            paragraph = (paragraphs or list(document.paragraphs))[
                int(location["paragraph_index"])
            ]
        except (IndexError, KeyError, TypeError, ValueError) as exc:
            raise HTTPException(
                status_code=409,
                detail="The target paragraph location is no longer valid.",
            ) from exc
    existing_style_name = (
        paragraph.style.name
        if paragraph.style is not None
        else None
    )
    existing_list_type, _existing_level = _numbering_type(
        document,
        paragraph,
        existing_style_name,
    )
    newline_attributes = _apply_inline_delta(paragraph, delta)
    if target.delta is not None:
        # A submitted Quill document is authoritative for supported block
        # formatting. Missing attributes therefore represent clear-format.
        if "header" not in newline_attributes and revision.element_type == "heading":
            if "Normal" in (style_names or {style.name for style in document.styles}):
                paragraph.style = "Normal"
        if "list" not in newline_attributes and existing_list_type is not None:
            _clear_numbering(paragraph)
            if "Normal" in (style_names or {style.name for style in document.styles}):
                paragraph.style = "Normal"
        if "align" not in newline_attributes:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if "header" in newline_attributes:
        _clear_numbering(paragraph)
        level = int(newline_attributes["header"])
        style_name = f"Heading {min(max(level, 1), 9)}"
        if style_name in (style_names or {style.name for style in document.styles}):
            paragraph.style = style_name
    if "list" in newline_attributes:
        if existing_list_type != newline_attributes.get("list"):
            _clear_numbering(paragraph)
        _set_list_format(
            document,
            paragraph,
            newline_attributes.get("list"),
            newline_attributes.get("indent", 0),
            existing_list_type=existing_list_type,
            style_names=style_names,
        )
    if "align" in newline_attributes:
        _set_alignment(paragraph, newline_attributes["align"])


def _apply_document_targets(
    source_path: Path,
    targets: list[tuple[EditorTarget, DocumentBlockRevision, DocumentRecord, dict]],
) -> tuple[DocxDocument, bytes]:
    total_started = perf_counter()
    stage_started = perf_counter()
    document = _load_docx(source_path)
    source_read_ms = (perf_counter() - stage_started) * 1000
    paragraphs = list(document.paragraphs)
    tables = list(document.tables)
    header_footer_parts = _header_footer_part_map(document)
    style_names = {style.name for style in document.styles}
    stage_started = perf_counter()
    for target, revision, _record, delta in sorted(
        targets,
        key=lambda item: item[1].ordinal,
    ):
        _apply_target_to_docx(
            document,
            revision,
            target,
            delta,
            paragraphs=paragraphs,
            tables=tables,
            header_footer_parts=header_footer_parts,
            style_names=style_names,
        )
    replacement_ms = (perf_counter() - stage_started) * 1000
    stage_started = perf_counter()
    output = BytesIO()
    document.save(output)
    payload = output.getvalue()
    package_write_ms = (perf_counter() - stage_started) * 1000
    stage_started = perf_counter()
    try:
        validated_document = Document(BytesIO(payload))
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="The edited DOCX failed validation and was not saved.",
        ) from exc
    validation_ms = (perf_counter() - stage_started) * 1000
    logger.info(
        "docsync.generation_document_timing source=%s targets=%s "
        "source_read_ms=%.2f replacement_ms=%.2f package_write_ms=%.2f "
        "validation_ms=%.2f total_ms=%.2f",
        source_path.name,
        len(targets),
        source_read_ms,
        replacement_ms,
        package_write_ms,
        validation_ms,
        (perf_counter() - total_started) * 1000,
    )
    return validated_document, payload


def _preview_payload(
    document_set_id: str,
    request: EditorEditRequest,
    validated_targets: list[
        tuple[EditorTarget, DocumentBlockRevision, DocumentRecord, dict]
    ],
) -> dict:
    documents: dict[str, dict] = {}
    for target, revision, document, delta in validated_targets:
        location = dict(revision.location_json or {})
        item = documents.setdefault(
            document.id,
            {
                "document_id": document.id,
                "document_name": document.original_name,
                "base_version_id": revision.version_id,
                "version_id": revision.version_id,
                "changes": [],
            },
        )
        item["changes"].append(
            {
                "element_id": revision.element_id,
                "paragraph_index": location.get(
                    "paragraph_index", revision.ordinal
                ),
                "element_type": revision.element_type,
                "document_order": location.get("document_order", revision.ordinal),
                "location": location,
                **{
                    key: location[key]
                    for key in (
                        "table_index",
                        "row_index",
                        "column_index",
                        "section_index",
                        "source_section_index",
                        "header_footer_type",
                        "part_relationship_id",
                        "is_linked_to_previous",
                        "section_indexes",
                        "linked_section_indexes",
                        "linked_sections",
                    )
                    if key in location
                },
                "before": revision.text,
                "after": target.replacement_text,
                "before_delta": revision.delta_json,
                "after_delta": delta,
                "delta": delta,
                "shared_state": (
                    "detached" if request.edit_mode == "override" else revision.shared_state
                ),
            }
        )
    result_documents = sorted(
        documents.values(),
        key=lambda item: item["document_name"].casefold(),
    )
    return {
        "operation_id": None,
        "document_set_id": document_set_id,
        "source_element_id": request.source_element_id,
        "edit_mode": request.edit_mode,
        "base_versions": dict(request.base_versions),
        "affected_document_count": len(result_documents),
        "affected_location_count": sum(
            len(document["changes"]) for document in result_documents
        ),
        "documents": result_documents,
        "status": "previewed",
        "writes_performed": False,
    }


def preview_editor_edit(
    session: Session,
    document_set_id: str,
    request: EditorEditRequest,
) -> dict:
    _source, _source_document, base_context, validated_targets = (
        _validate_editor_request(session, document_set_id, request)
    )
    by_document: dict[
        str, list[tuple[EditorTarget, DocumentBlockRevision, DocumentRecord, dict]]
    ] = defaultdict(list)
    for item in validated_targets:
        by_document[item[2].id].append(item)
    # Perform the full targeted OOXML round-trip in memory. Preview remains
    # side-effect free while surfacing unsupported/write-back failures early.
    for document_id, targets in by_document.items():
        version = base_context[document_id][2]
        _apply_document_targets(document_version_path(version), targets)
    return _preview_payload(document_set_id, request, validated_targets)


def _sha256(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def _revisions_by_location(
    session: Session,
    version_id: str,
) -> dict[str, DocumentBlockRevision]:
    revisions = session.scalars(
        select(DocumentBlockRevision).where(
            DocumentBlockRevision.version_id == version_id
        )
    )
    return {
        _location_key(revision.location_json): revision
        for revision in revisions
    }


def _replace_current_elements_and_create_revisions(
    session: Session,
    *,
    document: DocumentRecord,
    version: DocumentVersion,
    source_path: Path,
    base_version_id: str,
    override_locations: set[str] | None = None,
    reconnect_locations: set[str] | None = None,
    prepared_docx: DocxDocument | None = None,
    changed_match_hashes: set[str] | None = None,
) -> dict[str, str]:
    from .document_service import _extract_paragraphs

    docx = prepared_docx or _load_docx(source_path)
    extracted = _extract_paragraphs(docx)
    paragraphs = list(docx.paragraphs)
    tables = list(docx.tables)
    header_footer_parts = _header_footer_part_map(docx)
    style_name_cache: dict[str | None, str | None] = {}
    previous_revisions = _revisions_by_location(session, base_version_id)
    override_locations = override_locations or set()
    reconnect_locations = reconnect_locations or set()

    current_elements = list(
        session.scalars(
            select(DocumentElement).where(DocumentElement.document_id == document.id)
        )
    )
    existing_by_location = {
        (
            element.paragraph_index,
            _location_key(_element_location(element)[1]),
        ): element
        for element in current_elements
    }
    prepared_elements: list[
        tuple[int, str, str | None, str, DocumentElement | None]
    ] = []
    retained_element_ids: set[str] = set()
    for paragraph_index, text, style_name in extracted:
        probe = SimpleNamespace(
            paragraph_index=paragraph_index,
            style_name=style_name,
        )
        key = _location_key(_element_location(probe)[1])
        existing = existing_by_location.get((paragraph_index, key))
        if existing is not None:
            retained_element_ids.add(existing.id)
        prepared_elements.append((paragraph_index, text, style_name, key, existing))

    for element in current_elements:
        if element.id not in retained_element_ids:
            session.delete(element)
    session.flush()

    elements: list[tuple[DocumentElement, str]] = []
    new_elements: list[DocumentElement] = []
    for paragraph_index, text, style_name, key, existing in prepared_elements:
        element = existing or DocumentElement(
            id=new_id(),
            document_id=document.id,
            paragraph_index=paragraph_index,
            text=text,
            normalized_text=normalise_editor_text(text),
            style_name=style_name,
        )
        if existing is not None:
            element.text = text
            element.normalized_text = normalise_editor_text(text)
            element.style_name = style_name
        else:
            new_elements.append(element)
        elements.append((element, key))
    location_to_element: dict[str, str] = {}
    revisions: list[DocumentBlockRevision] = []
    current_locations: set[str] = set()
    for element, key in elements:
        current_locations.add(key)
        previous_revision = previous_revisions.get(key)
        shared_state = (
            previous_revision.shared_state
            if previous_revision is not None
            else "shared"
        )
        if key in override_locations:
            shared_state = "detached"
        elif key in reconnect_locations:
            shared_state = "shared"
        values = _revision_values(
            docx,
            element,
            shared_state=shared_state,
            paragraphs=paragraphs,
            tables=tables,
            style_name_cache=style_name_cache,
            header_footer_parts=header_footer_parts,
        )
        revision = DocumentBlockRevision(version_id=version.id, **values)
        revisions.append(revision)
        location_to_element[key] = element.id
        if (
            changed_match_hashes is not None
            and previous_revision is not None
            and previous_revision.exact_match_hash != revision.exact_match_hash
        ):
            if previous_revision.exact_match_hash:
                changed_match_hashes.add(previous_revision.exact_match_hash)
            if revision.exact_match_hash:
                changed_match_hashes.add(revision.exact_match_hash)
        elif changed_match_hashes is not None and previous_revision is None:
            if revision.exact_match_hash:
                changed_match_hashes.add(revision.exact_match_hash)
    if changed_match_hashes is not None:
        for key, previous_revision in previous_revisions.items():
            if key not in current_locations and previous_revision.exact_match_hash:
                changed_match_hashes.add(previous_revision.exact_match_hash)
    session.add_all([*new_elements, *revisions])
    session.flush()
    return location_to_element


def _persist_request_decisions(
    session: Session,
    *,
    document_set_id: str,
    source: DocumentBlockRevision,
    decisions: list[EditorMatchDecision],
) -> None:
    for item in decisions:
        candidate, candidate_document, _head = _get_current_revision_or_404(
            session,
            item.candidate_element_id,
        )
        if candidate_document.document_set_id != document_set_id:
            raise HTTPException(
                status_code=422,
                detail="Match decisions must stay within one document set.",
            )
        _upsert_decision(
            session,
            document_set_id=document_set_id,
            source=source,
            candidate=candidate,
            decision=item.status,
        )


def queue_editor_generation(
    session: Session,
    document_set_id: str,
    request: EditorEditRequest,
    *,
    retry_of_operation_id: str | None = None,
) -> dict:
    queue_started = perf_counter()
    with EDITOR_QUEUE_LOCK:
        # Reject an overlapping second click before doing the comparatively
        # expensive target/delta validation. base_versions contains exactly
        # the documents selected by the editor request.
        requested_documents = set(request.base_versions)
        pending_scan_started = perf_counter()
        pending_operations = list(
            session.scalars(
                select(EditorOperation)
                .where(
                    EditorOperation.document_set_id == document_set_id,
                    EditorOperation.status.in_(("queued", "processing")),
                )
                .order_by(EditorOperation.created_at.desc())
            )
        )
        pending_scan_ms = (perf_counter() - pending_scan_started) * 1000
        for pending in pending_operations:
            pending_envelope = pending.preview_json or {}
            pending_document_ids = set(
                pending_envelope.get("affected_document_ids", [])
                if isinstance(pending_envelope, dict)
                else []
            )
            if not pending_document_ids and isinstance(pending_envelope, dict):
                pending_request = pending_envelope.get("request", {})
                if isinstance(pending_request, dict):
                    pending_document_ids = set(
                        (pending_request.get("base_versions") or {}).keys()
                    )
            overlap = requested_documents & pending_document_ids
            if overlap:
                overlapping_names = list(
                    session.scalars(
                        select(DocumentRecord.original_name)
                        .where(DocumentRecord.id.in_(overlap))
                        .order_by(DocumentRecord.original_name)
                    )
                )
                logger.info(
                    "docsync.generation_queue_timing document_set_id=%s "
                    "status=overlap pending_scan_ms=%.2f total_ms=%.2f",
                    document_set_id,
                    pending_scan_ms,
                    (perf_counter() - queue_started) * 1000,
                )
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail=(
                        f"{', '.join(overlapping_names) or 'A selected document'} "
                        "already has a background update in progress. Wait for it "
                        "to finish, then reload and review this change against the "
                        "latest version."
                    ),
                )

        validation_started = perf_counter()
        _source, _source_document, _base_context, validated_targets = (
            _validate_editor_request(session, document_set_id, request)
        )
        validation_ms = (perf_counter() - validation_started) * 1000
        affected_documents = sorted(
            {
                item[2].id: item[2].original_name
                for item in validated_targets
            }.items(),
            key=lambda item: item[1].casefold(),
        )
        affected_document_ids = [item[0] for item in affected_documents]
        affected_document_names = [item[1] for item in affected_documents]
        preview = _preview_payload(
            document_set_id,
            request,
            validated_targets,
        )
        operation = EditorOperation(
            id=new_id(),
            document_set_id=document_set_id,
            operation_type=request.edit_mode,
            status="queued",
            stage="queued",
            source_element_id=request.source_element_id,
            replacement_text=(
                validated_targets[0][0].replacement_text
                if len(validated_targets) == 1
                else None
            ),
            preview_json={
                "request": request.model_dump(mode="json"),
                "preview": preview,
                "affected_document_ids": affected_document_ids,
                "affected_document_names": affected_document_names,
                "retry_of_operation_id": retry_of_operation_id,
            },
        )
        session.add(operation)
        commit_started = perf_counter()
        session.commit()
        commit_ms = (perf_counter() - commit_started) * 1000
        total_ms = (perf_counter() - queue_started) * 1000
        logger.info(
            "docsync.generation_queue_timing operation_id=%s document_set_id=%s "
            "status=queued pending_scan_ms=%.2f validation_ms=%.2f "
            "commit_ms=%.2f total_ms=%.2f",
            operation.id,
            document_set_id,
            pending_scan_ms,
            validation_ms,
            commit_ms,
            total_ms,
        )
        return {
            "id": operation.id,
            "job_id": operation.id,
            "operation_id": operation.id,
            "generation_id": operation.id,
            "status": "queued",
            "stage": "queued",
            "submitted_at": utc_isoformat(operation.created_at),
            "completed_at": None,
            "affected_document_ids": affected_document_ids,
            "affected_documents": [
                {"id": document_id, "name": document_name}
                for document_id, document_name in zip(
                    affected_document_ids,
                    affected_document_names,
                )
            ],
            "result_version_ids": [],
            "error": None,
            "edit_mode": request.edit_mode,
            "affected_document_count": preview["affected_document_count"],
            "affected_location_count": preview["affected_location_count"],
            "status_url": f"/api/editor-operations/{operation.id}",
            "timings": {
                "pending_scan_ms": round(pending_scan_ms, 2),
                "validation_ms": round(validation_ms, 2),
                "commit_ms": round(commit_ms, 2),
                "total_ms": round(total_ms, 2),
            },
        }


def _serialize_generated_document_updates(
    session: Session,
    versions: list[DocumentVersion],
) -> list[dict]:
    """Return only the changed workspace rows needed by the active UI."""

    if not versions:
        return []
    document_ids = {version.document_id for version in versions}
    element_counts = {
        document_id: int(count)
        for document_id, count in session.execute(
            select(DocumentElement.document_id, func.count(DocumentElement.id))
            .where(DocumentElement.document_id.in_(document_ids))
            .group_by(DocumentElement.document_id)
        ).all()
    }
    return [
        {
            "id": version.document_id,
            "version_id": version.id,
            "current_version_id": version.id,
            "version_number": version.version_number,
            "parent_version_id": version.parent_version_id,
            "name": version.document.original_name,
            "checksum_sha256": version.document.checksum_sha256,
            "current_checksum_sha256": version.checksum_sha256,
            "element_count": element_counts.get(version.document_id, 0),
            "view_url": f"/api/document-versions/{version.id}/pages",
            "download_url": f"/api/documents/{version.document_id}/download",
        }
        for version in sorted(
            versions,
            key=lambda item: item.document.original_name.casefold(),
        )
    ]


def serialize_editor_generation_status(
    session: Session,
    operation_id: str,
) -> dict:
    operation = session.get(EditorOperation, operation_id)
    if operation is None:
        raise HTTPException(status_code=404, detail="Editor operation not found.")

    preview_envelope = operation.preview_json or {}
    preview = (
        preview_envelope.get("preview", {})
        if isinstance(preview_envelope, dict)
        else {}
    )
    affected_document_ids = list(
        preview_envelope.get("affected_document_ids", [])
        if isinstance(preview_envelope, dict)
        else []
    )
    affected_document_names = list(
        preview_envelope.get("affected_document_names", [])
        if isinstance(preview_envelope, dict)
        else []
    )
    if not affected_document_ids and isinstance(preview, dict):
        affected_document_ids = [
            str(item.get("document_id"))
            for item in preview.get("documents", [])
            if item.get("document_id")
        ]
        affected_document_names = [
            str(item.get("document_name", "Document"))
            for item in preview.get("documents", [])
            if item.get("document_id")
        ]
    payload = {
        "id": operation.id,
        "job_id": operation.id,
        "operation_id": operation.id,
        "generation_id": operation.id,
        "status": operation.status,
        "stage": operation.stage,
        "submitted_at": utc_isoformat(operation.created_at),
        "updated_at": utc_isoformat(operation.updated_at),
        "completed_at": (
            utc_isoformat(operation.completed_at)
            if operation.completed_at is not None
            else None
        ),
        "affected_document_ids": affected_document_ids,
        "affected_documents": [
            {"id": document_id, "name": document_name}
            for document_id, document_name in zip(
                affected_document_ids,
                affected_document_names,
            )
        ],
        "result_version_ids": [version.id for version in operation.versions],
        "error": operation.error_detail,
        "edit_mode": operation.operation_type,
        "status_url": f"/api/editor-operations/{operation.id}",
        "timings": (
            preview_envelope.get("timings", {})
            if isinstance(preview_envelope, dict)
            else {}
        ),
        "progress": (
            preview_envelope.get("progress")
            if isinstance(preview_envelope, dict)
            else None
        ),
    }
    if operation.status in {"failed", "interrupted"}:
        payload["error_detail"] = (
            operation.error_detail
            or "The background document update failed."
        )
        return payload
    if operation.status != "completed":
        payload.update(
            {
                "affected_document_count": preview.get(
                    "affected_document_count", 0
                ),
                "affected_location_count": preview.get(
                    "affected_location_count", 0
                ),
            }
        )
        return payload

    versions = sorted(
        operation.versions,
        key=lambda item: item.document.original_name.casefold(),
    )
    serialized_versions = [
        {
            "id": version.id,
            "document_id": version.document_id,
            "document_name": version.document.original_name,
            "version_id": version.id,
            "version_number": version.version_number,
            "parent_version_id": version.parent_version_id,
            "checksum_sha256": version.checksum_sha256,
            "created_at": utc_isoformat(version.created_at),
            "status": "completed",
            "is_current": True,
            "generation_id": operation.id,
            "download_url": f"/api/document-versions/{version.id}/download",
            "editor_content_url": (
                f"/api/document-versions/{version.id}/editor-content"
            ),
        }
        for version in versions
    ]
    payload.update(
        {
            "versions": serialized_versions,
            "files": [
                {
                    "source_document_id": item["document_id"],
                    "version_id": item["version_id"],
                    "name": item["document_name"],
                    "download_url": item["download_url"],
                }
                for item in serialized_versions
            ],
            "download_url": (
                f"/api/editor-operations/{operation.id}/download"
                if operation.status == "completed"
                else None
            ),
            "document_updates": _serialize_generated_document_updates(
                session,
                versions,
            ),
        }
    )
    return payload


def process_queued_editor_generation(operation_id: str) -> None:
    from .database import SessionLocal

    with SessionLocal() as session:
        operation = session.get(EditorOperation, operation_id)
        if operation is None or operation.status != "queued":
            return
        envelope = operation.preview_json or {}
        request_payload = (
            envelope.get("request")
            if isinstance(envelope, dict)
            else None
        )
        if not isinstance(request_payload, dict):
            operation.status = "failed"
            operation.stage = "failed"
            operation.error_detail = "The queued editor request is missing."
            operation.completed_at = utc_now()
            session.commit()
            return

        operation.status = "processing"
        operation.stage = "preparing_documents"
        session.commit()
        try:
            request = EditorEditRequest.model_validate(request_payload)
            generate_editor_versions(
                session,
                operation.document_set_id,
                request,
                queued_operation_id=operation.id,
            )
        except Exception as exc:
            session.rollback()
            failed = session.get(EditorOperation, operation_id)
            if failed is not None:
                failed.status = "failed"
                failed.stage = "failed"
                failed.error_detail = (
                    str(exc.detail)
                    if isinstance(exc, HTTPException)
                    else str(exc) or "The background document update failed."
                )
                failed.completed_at = utc_now()
                session.commit()


def fail_interrupted_editor_generations(session: Session) -> int:
    interrupted = list(
        session.scalars(
            select(EditorOperation).where(
                EditorOperation.status.in_(("queued", "processing"))
            )
        )
    )
    for operation in interrupted:
        operation.status = "interrupted"
        operation.stage = "interrupted"
        operation.error_detail = (
            "DocSync restarted before this background update finished. "
            "No document versions were changed. Review the current document "
            "and retry the update when ready."
        )
        operation.completed_at = utc_now()
        generated_root = settings.data_dir / "generated" / operation.document_set_id
        shutil.rmtree(generated_root / f".{operation.id}.staging", ignore_errors=True)
        shutil.rmtree(generated_root / operation.id, ignore_errors=True)
    if interrupted:
        session.commit()
    return len(interrupted)


def submit_editor_generation(operation_id: str) -> None:
    """Start a durable queued operation without extending the request lifetime."""

    EDITOR_GENERATION_EXECUTOR.submit(process_queued_editor_generation, operation_id)


def _commit_editor_generation_stage(
    session: Session,
    operation_id: str | None,
    stage: str,
) -> None:
    if operation_id is None:
        return
    operation = session.get(EditorOperation, operation_id)
    if operation is None:
        raise HTTPException(status_code=404, detail="Generation job not found.")
    operation.stage = stage
    session.commit()


def retry_editor_generation(session: Session, operation_id: str) -> dict:
    operation = session.get(EditorOperation, operation_id)
    if operation is None:
        raise HTTPException(status_code=404, detail="Generation job not found.")
    if operation.status not in {"failed", "interrupted"}:
        raise HTTPException(
            status_code=409,
            detail="Only failed or interrupted generation jobs can be retried.",
        )
    envelope = operation.preview_json or {}
    request_payload = envelope.get("request") if isinstance(envelope, dict) else None
    if not isinstance(request_payload, dict):
        raise HTTPException(
            status_code=409,
            detail="This generation job does not contain a safe retry request.",
        )
    request = EditorEditRequest.model_validate(request_payload)
    return queue_editor_generation(
        session,
        operation.document_set_id,
        request,
        retry_of_operation_id=operation.id,
    )


def list_editor_generation_jobs(
    session: Session,
    document_set_id: str,
    *,
    limit: int = 50,
) -> list[dict]:
    if session.get(DocumentSet, document_set_id) is None:
        raise HTTPException(status_code=404, detail="Document set not found.")
    operation_ids = list(
        session.scalars(
            select(EditorOperation.id)
            .where(
                EditorOperation.document_set_id == document_set_id,
                EditorOperation.operation_type != "version_restore",
            )
            .order_by(EditorOperation.created_at.desc())
            .limit(limit)
        )
    )
    return [
        serialize_editor_generation_status(session, operation_id)
        for operation_id in operation_ids
    ]


def list_recoverable_editor_generation_jobs(
    session: Session,
    *,
    limit: int = 50,
) -> list[dict]:
    """Return jobs that still need global UI attention after startup."""

    operation_ids = list(
        session.scalars(
            select(EditorOperation.id)
            .where(
                EditorOperation.status.in_(("queued", "processing", "interrupted")),
                EditorOperation.operation_type != "version_restore",
            )
            .order_by(EditorOperation.created_at.desc())
            .limit(limit)
        )
    )
    return [
        serialize_editor_generation_status(session, operation_id)
        for operation_id in operation_ids
    ]


def generate_editor_versions(
    session: Session,
    document_set_id: str,
    request: EditorEditRequest,
    *,
    queued_operation_id: str | None = None,
) -> dict:
    generation_started = perf_counter()
    timings: dict[str, float] = {}
    operation_id = queued_operation_id or new_id()
    with EDITOR_GENERATION_LOCK:
        committed = False
        staging_directory: Path | None = None
        final_directory: Path | None = None
        try:
            stage_started = perf_counter()
            source, _source_document, base_context, validated_targets = (
                _validate_editor_request(session, document_set_id, request)
            )
            timings["database_lookup_and_validation_ms"] = (
                perf_counter() - stage_started
            ) * 1000
            preview = _preview_payload(
                document_set_id,
                request,
                validated_targets,
            )
            generated_root = settings.data_dir / "generated" / document_set_id
            generated_root.mkdir(parents=True, exist_ok=True)
            staging_directory = generated_root / f".{operation_id}.staging"
            final_directory = generated_root / operation_id
            if staging_directory.exists() or final_directory.exists():
                raise HTTPException(
                    status_code=409,
                    detail="An editor generation storage collision occurred. Try again.",
                )
            staging_directory.mkdir(parents=False, exist_ok=False)
            _commit_editor_generation_stage(
                session,
                queued_operation_id,
                "applying_changes",
            )

            by_document: dict[
                str,
                list[
                    tuple[
                        EditorTarget,
                        DocumentBlockRevision,
                        DocumentRecord,
                        dict,
                    ]
                ],
            ] = defaultdict(list)
            for item in validated_targets:
                by_document[item[2].id].append(item)

            staged: dict[str, dict] = {}
            from .document_service import safe_download_name

            timings["docx_generation_ms"] = 0.0
            timings["file_write_ms"] = 0.0
            for document_id, targets in by_document.items():
                document, head, base_version = base_context[document_id]
                stage_started = perf_counter()
                prepared_docx, payload = _apply_document_targets(
                    document_version_path(base_version),
                    targets,
                )
                timings["docx_generation_ms"] += (
                    perf_counter() - stage_started
                ) * 1000
                next_number = base_version.version_number + 1
                output_name = safe_download_name(
                    f"{Path(document.original_name).stem}-v{next_number}.docx"
                )
                output_path = staging_directory / output_name
                stage_started = perf_counter()
                output_path.write_bytes(payload)
                timings["file_write_ms"] += (
                    perf_counter() - stage_started
                ) * 1000
                staged[document_id] = {
                    "document": document,
                    "head": head,
                    "base_version": base_version,
                    "targets": targets,
                    "version_id": new_id(),
                    "version_number": next_number,
                    "output_name": output_name,
                    "staging_path": output_path,
                    "checksum": _sha256(payload),
                    "prepared_docx": prepared_docx,
                }

            _commit_editor_generation_stage(
                session,
                queued_operation_id,
                "validating_generated_files",
            )

            zip_path = staging_directory / "current-documents.zip"
            stage_started = perf_counter()
            with zipfile.ZipFile(
                zip_path,
                mode="w",
                # DOCX files are ZIP packages already; recompressing them wastes
                # CPU and provides negligible size reduction.
                compression=zipfile.ZIP_STORED,
            ) as archive:
                documents = list(
                    session.scalars(
                        select(DocumentRecord)
                        .where(DocumentRecord.document_set_id == document_set_id)
                        .order_by(DocumentRecord.original_name)
                    )
                )
                for document in documents:
                    staged_item = staged.get(document.id)
                    if staged_item is not None:
                        path = staged_item["staging_path"]
                    else:
                        path = document_version_path(
                            current_version_for_document(session, document)
                        )
                    archive.write(
                        path,
                        arcname=safe_download_name(document.original_name),
                    )
            timings["archive_write_ms"] = (perf_counter() - stage_started) * 1000

            _commit_editor_generation_stage(
                session,
                queued_operation_id,
                "saving_new_versions",
            )

            staging_directory.replace(final_directory)
            staging_directory = None
            _commit_editor_generation_stage(
                session,
                queued_operation_id,
                "refreshing_workspace",
            )

            database_update_started = perf_counter()
            operation = (
                session.get(EditorOperation, operation_id)
                if queued_operation_id is not None
                else None
            )
            if queued_operation_id is not None and operation is None:
                raise HTTPException(
                    status_code=404,
                    detail="The queued editor operation no longer exists.",
                )
            if operation is None:
                operation = EditorOperation(
                    id=operation_id,
                    document_set_id=document_set_id,
                    operation_type=request.edit_mode,
                    status="completed",
                    stage="completed",
                    source_element_id=request.source_element_id,
                    replacement_text=(
                        validated_targets[0][0].replacement_text
                        if len(validated_targets) == 1
                        else None
                    ),
                    preview_json={"preview": preview},
                    completed_at=utc_now(),
                )
                session.add(operation)
            else:
                operation.status = "completed"
                operation.stage = "completed"
                queued_envelope = (
                    dict(operation.preview_json)
                    if isinstance(operation.preview_json, dict)
                    else {}
                )
                operation.preview_json = {**queued_envelope, "preview": preview}
                operation.error_detail = None
                operation.completed_at = utc_now()
            session.flush()
            _persist_request_decisions(
                session,
                document_set_id=document_set_id,
                source=source,
                decisions=request.match_decisions,
            )

            response_versions: list[dict] = []
            changed_match_hashes: set[str] = set()
            timings["block_parsing_and_database_ms"] = 0.0
            for document_id, item in staged.items():
                document: DocumentRecord = item["document"]
                head: DocumentHead = item["head"]
                base_version: DocumentVersion = item["base_version"]
                version = DocumentVersion(
                    id=item["version_id"],
                    document_id=document.id,
                    parent_version_id=base_version.id,
                    generation_id=None,
                    editor_operation_id=operation.id,
                    version_number=item["version_number"],
                    storage_area="generated",
                    storage_name=(
                        f"{document_set_id}/{operation_id}/{item['output_name']}"
                    ),
                    download_name=item["output_name"],
                    checksum_sha256=item["checksum"],
                )
                session.add(version)
                session.flush()

                override_locations = {
                    _location_key(revision.location_json)
                    for _target, revision, _record, _delta in item["targets"]
                    if request.edit_mode == "override"
                }
                reconnect_locations = {
                    _location_key(revision.location_json)
                    for _target, revision, _record, _delta in item["targets"]
                    if request.edit_mode != "override"
                    and revision.shared_state == "detached"
                }
                block_stage_started = perf_counter()
                location_to_element = (
                    _replace_current_elements_and_create_revisions(
                        session,
                        document=document,
                        version=version,
                        source_path=final_directory / item["output_name"],
                        base_version_id=base_version.id,
                        override_locations=override_locations,
                        reconnect_locations=reconnect_locations,
                        prepared_docx=item["prepared_docx"],
                        changed_match_hashes=changed_match_hashes,
                    )
                )

                for ordinal, (
                    target,
                    revision,
                    _target_document,
                    delta,
                ) in enumerate(item["targets"]):
                    session.add(
                        EditorOperationTarget(
                            id=new_id(),
                            operation_id=operation.id,
                            document_id=document.id,
                            element_id=revision.element_id,
                            base_version_id=base_version.id,
                            result_version_id=version.id,
                            expected_head_revision=head.revision,
                            ordinal=ordinal,
                            before_text=revision.text,
                            after_text=target.replacement_text,
                            before_delta_json=revision.delta_json,
                            after_delta_json=delta,
                        )
                    )

                head.current_version_id = version.id
                head.revision += 1
                head.updated_at = utc_now()
                response_versions.append(
                    {
                        "id": version.id,
                        "document_id": document.id,
                        "document_name": document.original_name,
                        "version_id": version.id,
                        "version_number": version.version_number,
                        "parent_version_id": version.parent_version_id,
                        "checksum_sha256": version.checksum_sha256,
                        "created_at": utc_isoformat(version.created_at),
                        "status": "completed",
                        "is_current": True,
                        "generation_id": operation.id,
                        "download_url": (
                            f"/api/document-versions/{version.id}/download"
                        ),
                        "editor_content_url": (
                            f"/api/document-versions/{version.id}/editor-content"
                        ),
                        "element_ids_by_location": location_to_element,
                    }
                )
                timings["block_parsing_and_database_ms"] += (
                    perf_counter() - block_stage_started
                ) * 1000

            session.flush()
            from .document_service import (
                _rebuild_exact_link_groups_for_hashes,
                rendered_pdf_path,
            )

            stage_started = perf_counter()
            _rebuild_exact_link_groups_for_hashes(
                session,
                document_set_id,
                changed_match_hashes,
            )
            timings["matching_and_synchronisation_ms"] = (
                perf_counter() - stage_started
            ) * 1000
            for item in staged.values():
                # Remove legacy document-ID render cache. True-version renders
                # use their own cache name.
                rendered_pdf_path(item["document"]).unlink(missing_ok=True)
            timings["database_update_ms"] = (
                perf_counter() - database_update_started
            ) * 1000
            if isinstance(operation.preview_json, dict):
                # Assign a new JSON value so SQLAlchemy persists the timing
                # envelope without requiring a mutable JSON extension.
                operation.preview_json = {
                    **operation.preview_json,
                    "timings": {
                        key: round(value, 2) for key, value in timings.items()
                    },
                }
            commit_started = perf_counter()
            session.commit()
            timings["database_commit_ms"] = (perf_counter() - commit_started) * 1000
            committed = True

            stage_started = perf_counter()
            preview_jobs = _queue_previews_safely(
                session,
                [item["version_id"] for item in response_versions],
            )
            timings["deferred_preview_queue_ms"] = (
                perf_counter() - stage_started
            ) * 1000

            stage_started = perf_counter()
            generated_versions = [
                session.get(DocumentVersion, item["version_id"])
                for item in response_versions
            ]
            document_updates = _serialize_generated_document_updates(
                session,
                [version for version in generated_versions if version is not None],
            )
            refreshed = None
            if queued_operation_id is None:
                # Preserve the synchronous compatibility response. The active
                # UI uses the async endpoint and only needs changed rows.
                from .document_service import (
                    get_document_set_or_404,
                    serialize_document_set,
                )

                refreshed = serialize_document_set(
                    get_document_set_or_404(session, document_set_id)
                )
            timings["response_serialization_ms"] = (
                perf_counter() - stage_started
            ) * 1000
            timings["total_ms"] = (perf_counter() - generation_started) * 1000
            result = {
                "operation_id": operation.id,
                "generation_id": operation.id,
                "status": "completed",
                "edit_mode": request.edit_mode,
                "versions": sorted(
                    response_versions,
                    key=lambda item: item["document_name"].casefold(),
                ),
                "files": [
                    {
                        "source_document_id": item["document_id"],
                        "version_id": item["version_id"],
                        "name": item["document_name"],
                        "download_url": item["download_url"],
                    }
                    for item in response_versions
                ],
                "download_url": (
                    f"/api/editor-operations/{operation.id}/download"
                ),
                "preview_jobs": preview_jobs,
                "document_updates": document_updates,
                "timings": {
                    key: round(value, 2) for key, value in timings.items()
                },
            }
            if refreshed is not None:
                result["document_set"] = refreshed
            return result
        except Exception:
            if not committed:
                session.rollback()
                if staging_directory is not None:
                    shutil.rmtree(staging_directory, ignore_errors=True)
                if final_directory is not None:
                    shutil.rmtree(final_directory, ignore_errors=True)
            raise
        finally:
            timings.setdefault(
                "total_ms",
                (perf_counter() - generation_started) * 1000,
            )
            logger.info(
                "docsync.generation_timing operation_id=%s document_set_id=%s %s",
                operation_id,
                document_set_id,
                " ".join(
                    f"{name}={value:.2f}"
                    for name, value in timings.items()
                ),
            )


def restore_document_version(
    session: Session,
    document_id: str,
    target_version_id: str,
    request: VersionRestoreRequest,
) -> dict:
    """Copy a historical immutable version into a new current version."""

    with EDITOR_GENERATION_LOCK:
        committed = False
        staging_directory: Path | None = None
        final_directory: Path | None = None
        try:
            document = session.get(DocumentRecord, document_id)
            if document is None:
                raise HTTPException(status_code=404, detail="Document not found.")

            target_version = get_version_or_404(session, target_version_id)
            if target_version.document_id != document.id:
                raise HTTPException(
                    status_code=422,
                    detail="The version to restore does not belong to this document.",
                )

            current_version = current_version_for_document(session, document)
            head = session.get(DocumentHead, document.id)
            if head is None:
                raise HTTPException(
                    status_code=500,
                    detail="Document version head is missing.",
                )
            if head.current_version_id != request.expected_current_version_id:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail=(
                        f"{document.original_name} changed after version history was "
                        "opened. Reload the current version and try again."
                    ),
                )
            if target_version.id == current_version.id:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail="The selected version is already current.",
                )

            source_path = document_version_path(target_version)
            operation_id = new_id()
            generated_root = (
                settings.data_dir / "generated" / document.document_set_id
            )
            generated_root.mkdir(parents=True, exist_ok=True)
            candidate_staging_directory = (
                generated_root / f".{operation_id}.staging"
            )
            candidate_final_directory = generated_root / operation_id
            if (
                candidate_staging_directory.exists()
                or candidate_final_directory.exists()
            ):
                raise HTTPException(
                    status_code=409,
                    detail="A version restoration storage collision occurred. Try again.",
                )
            candidate_staging_directory.mkdir(parents=False, exist_ok=False)
            staging_directory = candidate_staging_directory

            highest_version_number = int(
                session.scalar(
                    select(func.max(DocumentVersion.version_number)).where(
                        DocumentVersion.document_id == document.id
                    )
                )
                or 0
            )
            next_version_number = highest_version_number + 1

            from .document_service import safe_download_name

            output_name = safe_download_name(
                f"{Path(document.original_name).stem}-v{next_version_number}.docx"
            )
            staging_path = staging_directory / output_name
            shutil.copyfile(source_path, staging_path)
            _load_docx(staging_path)
            checksum = _sha256(staging_path.read_bytes())
            result_version_id = new_id()

            zip_path = staging_directory / "current-documents.zip"
            with zipfile.ZipFile(
                zip_path,
                mode="w",
                compression=zipfile.ZIP_DEFLATED,
            ) as archive:
                documents = list(
                    session.scalars(
                        select(DocumentRecord)
                        .where(
                            DocumentRecord.document_set_id
                            == document.document_set_id
                        )
                        .order_by(DocumentRecord.original_name)
                    )
                )
                for current_document in documents:
                    archive_path = (
                        staging_path
                        if current_document.id == document.id
                        else document_version_path(
                            current_version_for_document(session, current_document)
                        )
                    )
                    archive.write(
                        archive_path,
                        arcname=safe_download_name(current_document.original_name),
                    )

            staging_directory.replace(candidate_final_directory)
            staging_directory = None
            final_directory = candidate_final_directory

            operation = EditorOperation(
                id=operation_id,
                document_set_id=document.document_set_id,
                operation_type="version_restore",
                status="completed",
                stage="completed",
                expected_head_revision=head.revision,
                preview_json={
                    "document_id": document.id,
                    "document_name": document.original_name,
                    "restored_from_version_id": target_version.id,
                    "restored_from_version_number": target_version.version_number,
                    "previous_current_version_id": current_version.id,
                    "result_version_id": result_version_id,
                    "expected_current_version_id": request.expected_current_version_id,
                    "writes_performed": True,
                    "status": "completed",
                },
                completed_at=utc_now(),
            )
            session.add(operation)
            session.flush()

            version = DocumentVersion(
                id=result_version_id,
                document_id=document.id,
                parent_version_id=current_version.id,
                generation_id=None,
                editor_operation_id=operation.id,
                version_number=next_version_number,
                storage_area="generated",
                storage_name=(
                    f"{document.document_set_id}/{operation.id}/{output_name}"
                ),
                download_name=output_name,
                checksum_sha256=checksum,
            )
            session.add(version)
            session.flush()

            location_to_element = _replace_current_elements_and_create_revisions(
                session,
                document=document,
                version=version,
                source_path=final_directory / output_name,
                base_version_id=target_version.id,
            )

            head.current_version_id = version.id
            head.revision += 1
            head.updated_at = utc_now()
            session.flush()

            from .document_service import (
                _rebuild_exact_link_groups,
                get_document_set_or_404,
                rendered_pdf_path,
                serialize_document_set,
            )

            _rebuild_exact_link_groups(session, document.document_set_id)
            session.flush()
            refreshed = serialize_document_set(
                get_document_set_or_404(session, document.document_set_id)
            )
            rendered_pdf_path(document).unlink(missing_ok=True)
            session.commit()
            committed = True

            preview_jobs = _queue_previews_safely(session, [version.id])

            return {
                "operation_id": operation.id,
                "generation_id": operation.id,
                "operation_type": operation.operation_type,
                "status": operation.status,
                "document_id": document.id,
                "document_name": document.original_name,
                "restored_from_version_id": target_version.id,
                "restored_from_version_number": target_version.version_number,
                "previous_current_version_id": current_version.id,
                "version": {
                    "id": version.id,
                    "document_id": document.id,
                    "document_name": document.original_name,
                    "version_id": version.id,
                    "version_number": version.version_number,
                    "parent_version_id": version.parent_version_id,
                    "restored_from_version_id": target_version.id,
                    "restored_from_version_number": target_version.version_number,
                    "checksum_sha256": version.checksum_sha256,
                    "created_at": utc_isoformat(version.created_at),
                    "status": "completed",
                    "is_current": True,
                    "generation_id": operation.id,
                    "operation_type": operation.operation_type,
                    "download_url": (
                        f"/api/document-versions/{version.id}/download"
                    ),
                    "editor_content_url": (
                        f"/api/document-versions/{version.id}/editor-content"
                    ),
                    "element_ids_by_location": location_to_element,
                },
                "download_url": f"/api/editor-operations/{operation.id}/download",
                "preview_jobs": preview_jobs,
                "document_set": refreshed,
            }
        except Exception:
            if not committed:
                session.rollback()
                if staging_directory is not None:
                    shutil.rmtree(staging_directory, ignore_errors=True)
                if final_directory is not None:
                    shutil.rmtree(final_directory, ignore_errors=True)
            raise


def editor_operation_download_path(
    session: Session,
    operation_id: str,
) -> Path:
    operation = session.get(EditorOperation, operation_id)
    if operation is None or operation.status != "completed":
        raise HTTPException(status_code=404, detail="Editor generation not found.")
    path = (
        settings.data_dir
        / "generated"
        / operation.document_set_id
        / operation.id
        / "current-documents.zip"
    )
    if not path.is_file():
        raise HTTPException(status_code=404, detail="Generated ZIP file is missing.")
    return path


def register_legacy_generation_versions(
    session: Session,
    *,
    generation_id: str,
    generated_files: list[tuple[DocumentRecord, str, Path]],
) -> None:
    """Bridge the legacy exact-edit path into true version lineage immediately."""

    generated_rows = {
        row.source_document_id: row
        for row in session.scalars(
            select(GeneratedVersion).where(
                GeneratedVersion.generation_id == generation_id
            )
        )
    }
    for document, _output_name, output_path in generated_files:
        generated = generated_rows.get(document.id)
        if generated is None:
            raise HTTPException(
                status_code=500,
                detail="Generated version metadata could not be registered.",
            )
        existing = session.get(DocumentVersion, generated.id)
        if existing is not None:
            continue
        base_version = current_version_for_document(session, document)
        version = DocumentVersion(
            id=generated.id,
            document_id=document.id,
            parent_version_id=base_version.id,
            generation_id=generation_id,
            editor_operation_id=None,
            version_number=base_version.version_number + 1,
            storage_area="generated",
            storage_name=generated.storage_name,
            download_name=generated.download_name,
            checksum_sha256=hashlib.sha256(output_path.read_bytes()).hexdigest(),
        )
        session.add(version)
        session.flush()
        _replace_current_elements_and_create_revisions(
            session,
            document=document,
            version=version,
            source_path=output_path,
            base_version_id=base_version.id,
        )
        head = session.get(DocumentHead, document.id)
        head.current_version_id = version.id
        head.revision += 1
        head.updated_at = utc_now()


def serialize_editor_history(
    session: Session,
    document_set_id: str,
) -> list[dict]:
    operations = list(
        session.scalars(
            select(EditorOperation)
            .where(EditorOperation.document_set_id == document_set_id)
            .options(
                selectinload(EditorOperation.targets),
                selectinload(EditorOperation.versions),
                selectinload(EditorOperation.batch_operations).selectinload(
                    EditBatchOperation.occurrences
                ),
            )
            .order_by(EditorOperation.created_at.desc())
        )
    )
    return [
        {
            "operation_id": operation.id,
            "generation_id": operation.id,
            "event_type": (
                "version_restore"
                if operation.operation_type == "version_restore"
                else (
                    "batch_edit"
                    if operation.operation_type == "batch"
                    else "editor_edit"
                )
            ),
            "edit_mode": (
                None
                if operation.operation_type == "version_restore"
                else operation.operation_type
            ),
            "operation_type": operation.operation_type,
            "restored_from_version_id": (
                (operation.preview_json or {}).get("restored_from_version_id")
                if isinstance(operation.preview_json, dict)
                else None
            ),
            "previous_current_version_id": (
                (operation.preview_json or {}).get("previous_current_version_id")
                if isinstance(operation.preview_json, dict)
                else None
            ),
            "restored_from_version_number": (
                (operation.preview_json or {}).get("restored_from_version_number")
                if isinstance(operation.preview_json, dict)
                else None
            ),
            "status": operation.status,
            "created_at": utc_isoformat(operation.created_at),
            "completed_at": (
                utc_isoformat(operation.completed_at)
                if operation.completed_at is not None
                else None
            ),
            "target_count": len(operation.targets) + sum(
                sum(
                    1
                    for occurrence in item.occurrences
                    if (
                        occurrence.result_version_id is not None
                        if operation.status == "completed"
                        else occurrence.selected
                    )
                )
                for item in operation.batch_operations
            ),
            "version_count": len(operation.versions),
            "title": (
                (operation.preview_json or {}).get("title")
                if isinstance(operation.preview_json, dict)
                else None
            ),
            "batch_operations": [
                {
                    "operation_id": item.id,
                    "operation_index": item.operation_index,
                    "operation_type": item.operation_type,
                    "label": item.label,
                    "replacement_text": item.replacement_text,
                    "enabled": item.enabled,
                    "occurrence_count": sum(
                        1
                        for occurrence in item.occurrences
                        if (
                            occurrence.result_version_id is not None
                            if operation.status == "completed"
                            else occurrence.selected
                        )
                    ),
                    "document_count": len(
                        {occurrence.document_id for occurrence in item.occurrences}
                    ),
                }
                for item in sorted(
                    operation.batch_operations,
                    key=lambda item: item.operation_index,
                )
            ],
            "versions": [
                {
                    "document_id": version.document_id,
                    "version_id": version.id,
                    "version_number": version.version_number,
                    "parent_version_id": version.parent_version_id,
                    "download_url": (
                        f"/api/document-versions/{version.id}/download"
                    ),
                }
                for version in sorted(
                    operation.versions,
                    key=lambda item: (item.document_id, item.version_number),
                )
            ],
            "targets": [
                {
                    "element_id": target.element_id,
                    "document_id": target.document_id,
                    "base_version_id": target.base_version_id,
                    "result_version_id": target.result_version_id,
                    "before": target.before_text,
                    "after": target.after_text,
                }
                for target in operation.targets
            ],
            "download_url": (
                f"/api/editor-operations/{operation.id}/download"
            ),
        }
        for operation in operations
    ]
