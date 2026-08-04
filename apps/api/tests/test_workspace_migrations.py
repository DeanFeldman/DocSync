from __future__ import annotations

import importlib
import io
import sys
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine


API_DIR = Path(__file__).resolve().parents[1]
if str(API_DIR) not in sys.path:
    sys.path.insert(0, str(API_DIR))

from app.migration_service import (  # noqa: E402
    WorkspaceMigration,
    WorkspaceMigrationError,
    create_verified_backup,
    detect_schema_version,
    run_workspace_migrations,
)


def database_url(path: Path) -> str:
    return f"sqlite:///{path.as_posix()}"


def test_ordered_migrations_run_once_and_record_current_version(
    tmp_path: Path,
) -> None:
    database_path = tmp_path / "workspace.db"
    engine = create_engine(database_url(database_path))
    calls: list[int] = []

    def migration(version: int):
        def apply(session) -> None:
            calls.append(version)
            session.connection().exec_driver_sql(
                "CREATE TABLE IF NOT EXISTS migration_probe "
                "(version INTEGER PRIMARY KEY)"
            )
            session.connection().exec_driver_sql(
                "INSERT INTO migration_probe (version) VALUES (?)",
                (version,),
            )

        return apply

    migrations = (
        WorkspaceMigration(1, "first", migration(1)),
        WorkspaceMigration(2, "second", migration(2)),
    )
    first = run_workspace_migrations(
        engine=engine,
        database_url=database_url(database_path),
        backup_directory=tmp_path / "backups",
        create_schema=lambda _connection: None,
        migrations=migrations,
    )
    second = run_workspace_migrations(
        engine=engine,
        database_url=database_url(database_path),
        backup_directory=tmp_path / "backups",
        create_schema=lambda _connection: None,
        migrations=migrations,
    )

    assert first.applied_versions == (1, 2)
    assert second.applied_versions == ()
    assert calls == [1, 2]
    assert detect_schema_version(database_path) == 2


def test_legacy_database_is_backed_up_before_migration(tmp_path: Path) -> None:
    database_path = tmp_path / "workspace.db"
    engine = create_engine(database_url(database_path))
    with engine.begin() as connection:
        connection.exec_driver_sql(
            "CREATE TABLE legacy_documents (id INTEGER PRIMARY KEY)"
        )
        connection.exec_driver_sql("INSERT INTO legacy_documents (id) VALUES (1)")

    backup_directory = tmp_path / "migration-backups"

    def apply(session) -> None:
        backups = list(backup_directory.glob("workspace-before-schema-1-*.db"))
        assert len(backups) == 1
        assert backups[0].stat().st_size > 0
        session.connection().exec_driver_sql(
            "CREATE TABLE migrated_value (id INTEGER PRIMARY KEY)"
        )

    result = run_workspace_migrations(
        engine=engine,
        database_url=database_url(database_path),
        backup_directory=backup_directory,
        create_schema=lambda _connection: None,
        migrations=(WorkspaceMigration(1, "legacy", apply),),
    )

    assert result.backup_path is not None
    assert result.backup_path.is_file()
    assert result.backup_path.parent != database_path.parent


def test_failed_migration_restores_original_database(tmp_path: Path) -> None:
    database_path = tmp_path / "workspace.db"
    engine = create_engine(database_url(database_path))
    with engine.begin() as connection:
        connection.exec_driver_sql(
            "CREATE TABLE legacy_documents (id INTEGER PRIMARY KEY, value TEXT)"
        )
        connection.exec_driver_sql(
            "INSERT INTO legacy_documents (id, value) VALUES (1, 'original')"
        )

    def fail_after_write(session) -> None:
        session.connection().exec_driver_sql(
            "UPDATE legacy_documents SET value = 'changed' WHERE id = 1"
        )
        raise RuntimeError("forced migration failure")

    with pytest.raises(WorkspaceMigrationError, match="restored"):
        run_workspace_migrations(
            engine=engine,
            database_url=database_url(database_path),
            backup_directory=tmp_path / "backups",
            create_schema=lambda _connection: None,
            migrations=(WorkspaceMigration(1, "failing", fail_after_write),),
        )

    with create_engine(database_url(database_path)).connect() as connection:
        value = connection.exec_driver_sql(
            "SELECT value FROM legacy_documents WHERE id = 1"
        ).scalar_one()
    assert value == "original"
    assert detect_schema_version(database_path) == 0


def test_backup_failure_prevents_migration_from_running(tmp_path: Path) -> None:
    database_path = tmp_path / "workspace.db"
    engine = create_engine(database_url(database_path))
    with engine.begin() as connection:
        connection.exec_driver_sql(
            "CREATE TABLE legacy_documents (id INTEGER PRIMARY KEY)"
        )
    invalid_backup_directory = tmp_path / "not-a-directory"
    invalid_backup_directory.write_text("blocked", encoding="utf-8")
    called = False

    def must_not_run(_session) -> None:
        nonlocal called
        called = True

    with pytest.raises(WorkspaceMigrationError, match="No migration was run"):
        run_workspace_migrations(
            engine=engine,
            database_url=database_url(database_path),
            backup_directory=invalid_backup_directory,
            create_schema=lambda _connection: None,
            migrations=(WorkspaceMigration(1, "blocked", must_not_run),),
        )
    assert called is False
    assert detect_schema_version(database_path) == 0


def test_only_five_newest_migration_backups_are_retained(tmp_path: Path) -> None:
    database_path = tmp_path / "workspace.db"
    engine = create_engine(database_url(database_path))
    with engine.begin() as connection:
        connection.exec_driver_sql(
            "CREATE TABLE legacy_documents (id INTEGER PRIMARY KEY)"
        )
    backup_directory = tmp_path / "backups"

    for target_version in range(1, 8):
        create_verified_backup(
            database_path,
            backup_directory,
            target_version,
        )

    assert database_path.is_file()
    assert len(list(backup_directory.glob("workspace-before-schema-*.db"))) == 5


def test_migration_identifiers_must_be_monotonic(tmp_path: Path) -> None:
    database_path = tmp_path / "workspace.db"
    engine = create_engine(database_url(database_path))
    with pytest.raises(WorkspaceMigrationError, match="monotonically"):
        run_workspace_migrations(
            engine=engine,
            database_url=database_url(database_path),
            backup_directory=tmp_path / "backups",
            create_schema=lambda _connection: None,
            migrations=(
                WorkspaceMigration(2, "second", lambda _session: None),
                WorkspaceMigration(1, "first", lambda _session: None),
            ),
        )


def _docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_repeated_real_startup_does_not_duplicate_version_foundation(
    tmp_path: Path,
    monkeypatch,
) -> None:
    data_directory = tmp_path / "data"
    database_path = tmp_path / "v141-workspace.db"
    monkeypatch.setenv("DOCUMENTSYNC_DATA_DIR", str(data_directory))
    monkeypatch.setenv("DOCUMENTSYNC_DATABASE_URL", database_url(database_path))
    monkeypatch.setenv("DOCUMENTSYNC_SESSION_TOKEN", "")
    monkeypatch.delenv("DOCUMENTSYNC_WEB_DIST", raising=False)
    for module_name in list(sys.modules):
        if module_name == "app" or module_name.startswith("app."):
            del sys.modules[module_name]

    main = importlib.import_module("app.main")
    with TestClient(main.app) as client:
        response = client.post(
            "/api/document-sets",
            data={"name": "Representative v1.4.1 workspace"},
            files=[
                ("files", ("one.docx", _docx_bytes("Shared"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
                ("files", ("two.docx", _docx_bytes("Shared"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
            ],
        )
        assert response.status_code == 201

    database = importlib.import_module("app.database")
    with database.engine.begin() as connection:
        baseline = {
            table: connection.exec_driver_sql(f"SELECT COUNT(*) FROM {table}").scalar_one()
            for table in (
                "document_versions",
                "document_heads",
                "document_block_revisions",
                "editor_operations",
                "generated_versions",
            )
        }
        connection.exec_driver_sql("DROP TABLE docsync_schema_migrations")

    database.init_db()
    database.init_db()

    with database.engine.connect() as connection:
        after = {
            table: connection.exec_driver_sql(f"SELECT COUNT(*) FROM {table}").scalar_one()
            for table in baseline
        }
    assert after == baseline
    assert detect_schema_version(database_path) == 5
    backups = list((data_directory / "migration-backups").glob("*.db"))
    assert len(backups) == 1


def _table_docx(text: str) -> bytes:
    document = Document()
    document.add_heading("Legacy table", level=1)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = text
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _linked_header_docx(text: str) -> bytes:
    document = Document()
    document.sections[0].header.paragraphs[0].text = text
    document.add_paragraph("First section body")
    second = document.add_section()
    document.add_paragraph("Second section body")
    assert second.header.is_linked_to_previous
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_v15_migration_regenerates_legacy_table_cells_after_verified_backup(
    tmp_path: Path,
    monkeypatch,
) -> None:
    data_directory = tmp_path / "data"
    database_path = tmp_path / "v142-workspace.db"
    monkeypatch.setenv("DOCUMENTSYNC_DATA_DIR", str(data_directory))
    monkeypatch.setenv("DOCUMENTSYNC_DATABASE_URL", database_url(database_path))
    monkeypatch.setenv("DOCUMENTSYNC_SESSION_TOKEN", "")
    monkeypatch.delenv("DOCUMENTSYNC_WEB_DIST", raising=False)
    for module_name in list(sys.modules):
        if module_name == "app" or module_name.startswith("app."):
            del sys.modules[module_name]

    main = importlib.import_module("app.main")
    with TestClient(main.app) as client:
        response = client.post(
            "/api/document-sets",
            data={"name": "Legacy table workspace"},
            files=[
                (
                    "files",
                    (
                        "one.docx",
                        _table_docx("Shared legacy cell"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                (
                    "files",
                    (
                        "two.docx",
                        _table_docx("Shared legacy cell"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        assert response.status_code == 201, response.text
        workspace = response.json()

    database = importlib.import_module("app.database")
    editor_service = importlib.import_module("app.editor_service")
    with database.engine.begin() as connection:
        connection.exec_driver_sql(
            """
            UPDATE document_elements
            SET style_name = 'table_cell_order:1:0:0:0'
            WHERE style_name LIKE 'table_paragraph_order:%'
            """
        )
        connection.exec_driver_sql(
            """
            UPDATE document_block_revisions
            SET element_type = 'table_cell',
                location_json = '{"kind":"table_cell","document_order":1,"paragraph_index":1,"table_index":0,"row_index":0,"column_index":0}'
            WHERE element_type = 'table_paragraph'
            """
        )
        connection.exec_driver_sql(
            "DELETE FROM docsync_schema_migrations WHERE version >= 2"
        )

    original_revision_values = editor_service._revision_values
    cached_revision_calls = 0

    def require_cached_document_structure(*args, **kwargs):
        nonlocal cached_revision_calls
        assert kwargs.get("paragraphs") is not None
        assert kwargs.get("tables") is not None
        assert kwargs.get("style_name_cache") is not None
        assert kwargs.get("header_footer_parts") is not None
        cached_revision_calls += 1
        return original_revision_values(*args, **kwargs)

    monkeypatch.setattr(
        editor_service,
        "_revision_values",
        require_cached_document_structure,
    )
    database.init_db()

    assert cached_revision_calls > 0
    assert detect_schema_version(database_path) == 5
    backups = list((data_directory / "migration-backups").glob("*.db"))
    assert len(backups) == 1
    assert backups[0].stat().st_size > 0

    with TestClient(main.app) as client:
        first = workspace["documents"][0]
        content = client.get(
            f"/api/document-versions/{first['version_id']}/editor-content"
        )
        assert content.status_code == 200, content.text
        table_blocks = [
            block
            for block in content.json()["blocks"]
            if block["element_type"] == "table_paragraph"
        ]
        assert len(table_blocks) == 1
        assert table_blocks[0]["text"] == "Shared legacy cell"
        assert table_blocks[0]["paragraph_index"] == 0
        assert table_blocks[0]["supported"] is True


def test_v17_migration_backfills_linked_headers_from_immutable_docx(
    tmp_path: Path,
    monkeypatch,
) -> None:
    data_directory = tmp_path / "data"
    database_path = tmp_path / "v160-workspace.db"
    monkeypatch.setenv("DOCUMENTSYNC_DATA_DIR", str(data_directory))
    monkeypatch.setenv("DOCUMENTSYNC_DATABASE_URL", database_url(database_path))
    monkeypatch.setenv("DOCUMENTSYNC_SESSION_TOKEN", "")
    monkeypatch.delenv("DOCUMENTSYNC_WEB_DIST", raising=False)
    for module_name in list(sys.modules):
        if module_name == "app" or module_name.startswith("app."):
            del sys.modules[module_name]

    main = importlib.import_module("app.main")
    with TestClient(main.app) as client:
        response = client.post(
            "/api/document-sets",
            data={"name": "Legacy header workspace"},
            files=[
                (
                    "files",
                    (
                        name,
                        _linked_header_docx("Shared linked header"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
                for name in ("one.docx", "two.docx")
            ],
        )
        assert response.status_code == 201, response.text
        workspace = response.json()

    database = importlib.import_module("app.database")
    with database.engine.begin() as connection:
        connection.exec_driver_sql(
            "DELETE FROM link_members WHERE element_id IN "
            "(SELECT id FROM document_elements "
            "WHERE style_name LIKE 'header_footer_order:%')"
        )
        connection.exec_driver_sql(
            "DELETE FROM document_block_revisions "
            "WHERE element_type IN ('header_paragraph', 'footer_paragraph')"
        )
        connection.exec_driver_sql(
            "DELETE FROM document_elements "
            "WHERE style_name LIKE 'header_footer_order:%'"
        )
        connection.exec_driver_sql(
            "DELETE FROM docsync_schema_migrations WHERE version >= 4"
        )

    database.init_db()

    assert detect_schema_version(database_path) == 5
    backups = list((data_directory / "migration-backups").glob("*.db"))
    assert len(backups) == 1
    with TestClient(main.app) as client:
        first = workspace["documents"][0]
        content = client.get(
            f"/api/document-versions/{first['version_id']}/editor-content"
        )
        assert content.status_code == 200, content.text
        headers = [
            block
            for block in content.json()["blocks"]
            if block["element_type"] == "header_paragraph"
        ]
        assert len(headers) == 1
        assert headers[0]["text"] == "Shared linked header"
        assert headers[0]["section_indexes"] == [0, 1]
        assert headers[0]["linked_section_indexes"] == [1]
