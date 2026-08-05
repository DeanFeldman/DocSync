from __future__ import annotations

import importlib
import io
from pathlib import Path
import sys
import time

from docx import Document
import pymupdf
from fastapi.testclient import TestClient


API_DIR = Path(__file__).resolve().parents[1]
if str(API_DIR) not in sys.path:
    sys.path.insert(0, str(API_DIR))

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def load_test_app(tmp_path: Path, monkeypatch, *, threshold: str = "0.90"):
    monkeypatch.setenv("DOCUMENTSYNC_DATA_DIR", str(tmp_path / "data"))
    monkeypatch.setenv(
        "DOCUMENTSYNC_DATABASE_URL",
        f"sqlite:///{tmp_path / 'render-map.db'}",
    )
    monkeypatch.setenv("DOCUMENTSYNC_SESSION_TOKEN", "")
    monkeypatch.setenv("DOCUMENTSYNC_RENDER_MAP_CONFIDENCE_THRESHOLD", threshold)
    monkeypatch.delenv("DOCUMENTSYNC_WEB_DIST", raising=False)
    for module_name in list(sys.modules):
        if module_name == "app" or module_name.startswith("app."):
            del sys.modules[module_name]
    return importlib.import_module("app.main").app


def save_docx(document: Document) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def upload(client: TestClient, payload: bytes) -> tuple[dict, dict]:
    response = client.post(
        "/api/document-sets",
        data={"name": "Render map test"},
        files=[
            ("files", ("Mapped.docx", io.BytesIO(payload), DOCX_MEDIA_TYPE)),
            ("files", ("Companion.docx", io.BytesIO(payload), DOCX_MEDIA_TYPE)),
        ],
    )
    assert response.status_code == 201, response.text
    workspace = response.json()
    record = next(item for item in workspace["documents"] if item["name"] == "Mapped.docx")
    return workspace, record


def write_pdf(path: Path, pages: list[list[tuple[float, float, str, float | None]]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.unlink(missing_ok=True)
    pdf = pymupdf.open()
    for lines in pages:
        page = pdf.new_page(width=612, height=792)
        for x, y, text, box_width in lines:
            if box_width:
                page.insert_textbox(
                    pymupdf.Rect(x, y, x + box_width, y + 180),
                    text,
                    fontsize=11,
                )
            else:
                page.insert_text((x, y), text, fontsize=11)
    pdf.set_metadata({"producer": f"DocSync test {time.time_ns()}"})
    pdf.save(path)
    pdf.close()


def poll_map(client: TestClient, version_id: str, timeout: float = 8) -> dict:
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        response = client.get(f"/api/document-versions/{version_id}/render-map")
        assert response.status_code == 200, response.text
        payload = response.json()
        if payload["status"] in {"completed", "partial", "failed"}:
            return payload
        time.sleep(0.03)
    raise AssertionError("Render map did not reach a terminal state.")


def test_render_map_is_version_bound_cached_and_invalidated(
    tmp_path: Path,
    monkeypatch,
) -> None:
    long_text = (
        "This intentionally long paragraph wraps across several visual lines while "
        "remaining one immutable structured Word block for direct selection."
    )
    duplicate = "Repeated resident contact details"
    cross_page = "Cross-page paragraph begins here and continues on the next page"
    document = Document()
    document.add_heading("Mapped document", level=1)
    document.add_paragraph(long_text)
    document.add_paragraph(duplicate)
    document.add_paragraph(duplicate)
    document.add_paragraph(cross_page)

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, record = upload(client, save_docx(document))
        version_id = record["version_id"]
        pdf_path = (
            tmp_path
            / "data"
            / "renders"
            / workspace["id"]
            / f"{version_id}.pdf"
        )
        write_pdf(
            pdf_path,
            [
                [
                    (72, 90, "Mapped document", None),
                    (72, 125, long_text, 250),
                    (72, 250, duplicate, None),
                    (72, 285, duplicate, None),
                    (72, 740, "Cross-page paragraph begins here", None),
                ],
                [(72, 70, "and continues on the next page", None)],
            ],
        )

        queued = client.get(f"/api/document-versions/{version_id}/render-map")
        assert queued.status_code == 200
        assert queued.json()["status"] == "queued"
        mapped = poll_map(client, version_id)
        assert mapped["status"] == "completed"
        assert mapped["version_id"] == version_id
        assert mapped["source_sha256"]
        assert mapped["pdf_sha256"]
        assert mapped["mapper"] == "PyMuPDF"
        assert mapped["pdf_engine"].startswith("Microsoft Word")
        assert mapped["interactive_threshold"] == 0.90
        assert mapped["coordinate_unit"] == "normalised"
        assert mapped["mapped_element_count"] == 5
        assert mapped["interactive_element_count"] == 5
        assert len({region["element_id"] for region in mapped["regions"]}) == 5
        assert all(
            0 <= region[key] <= 1
            for region in mapped["regions"]
            for key in ("x", "y", "width", "height")
        )
        long_regions = [
            region
            for region in mapped["regions"]
            if region["text_preview"] == long_text
        ]
        assert len(long_regions) >= 2
        duplicate_regions = [
            region
            for region in mapped["regions"]
            if region["text_preview"] == duplicate
        ]
        assert len({region["element_id"] for region in duplicate_regions}) == 2
        assert all(region["confidence"] == 0.95 for region in duplicate_regions)
        cross_regions = [
            region
            for region in mapped["regions"]
            if region["text_preview"] == cross_page
        ]
        assert {region["page_number"] for region in cross_regions} == {1, 2}
        assert all(
            region["mapping_method"] == "word_pdf_text_context_order"
            for region in cross_regions
        )
        assert all(page["render_version"] == mapped["render_id"] for page in mapped["pages"])

        cached = client.get(f"/api/document-versions/{version_id}/render-map").json()
        assert cached["render_id"] == mapped["render_id"]
        assert cached["generated_at"] == mapped["generated_at"]

        page = client.get(mapped["pages"][0]["image_url"])
        assert page.status_code == 200
        assert page.headers["content-type"] == "image/png"
        assert page.content.startswith(b"\x89PNG")
        denied = client.get(
            f"/api/document-versions/{version_id}/render-pages/"
            f"{'0' * 24}/1.png"
        )
        assert denied.status_code == 404
        companion = next(
            item for item in workspace["documents"] if item["name"] == "Companion.docx"
        )
        mismatch = client.get(
            f"/api/document-versions/{companion['version_id']}/render-pages/"
            f"{mapped['render_id']}/1.png"
        )
        assert mismatch.status_code == 404

        old_render_id = mapped["render_id"]
        write_pdf(
            pdf_path,
            [
                [
                    (72, 90, "Mapped document", None),
                    (72, 125, long_text, 250),
                    (72, 250, duplicate, None),
                    (72, 285, duplicate, None),
                    (72, 740, "Cross-page paragraph begins here", None),
                ],
                [(72, 70, "and continues on the next page", None)],
                [(72, 90, "A new PDF page invalidates the coordinate cache", None)],
            ],
        )
        requeued = client.get(f"/api/document-versions/{version_id}/render-map")
        assert requeued.json()["status"] == "queued"
        regenerated = poll_map(client, version_id)
        assert regenerated["render_id"] != old_render_id
        assert regenerated["page_count"] == 3


def test_ambiguous_and_read_only_regions_never_become_editable(
    tmp_path: Path,
    monkeypatch,
) -> None:
    duplicate = "Ambiguous duplicate content"
    document = Document()
    document.add_heading("Safety boundary", level=1)
    document.add_paragraph(duplicate)
    document.add_paragraph(duplicate)
    table = document.add_table(rows=1, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Merged read-only table content"

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, record = upload(client, save_docx(document))
        version_id = record["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        # There is only one PDF occurrence for two identical body revisions.
        # The mapper must reject both rather than guess.
        write_pdf(
            pdf_path,
            [
                [
                    (72, 90, "Safety boundary", None),
                    (72, 130, duplicate, None),
                    (72, 180, "Merged read-only table content", None),
                ]
            ],
        )
        mapped = poll_map(client, version_id)
        assert mapped["status"] == "partial"
        ambiguous = [item for item in mapped["unmapped"] if item["element_type"] == "paragraph"]
        assert len(ambiguous) == 2
        assert all("could not be resolved safely" in item["reason"] for item in ambiguous)
        assert not any(
            region["text_preview"] == duplicate for region in mapped["regions"]
        )
        read_only = next(
            region
            for region in mapped["regions"]
            if region["text_preview"] == "Merged read-only table content"
        )
        assert read_only["supported"] is False
        assert read_only["interactive"] is False
        assert read_only["editable"] is False
        assert "merged structure" in read_only["read_only_reason"]


def test_header_footer_occurrences_map_across_pages_and_corrupt_pdf_fails_safely(
    tmp_path: Path,
    monkeypatch,
) -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Resident services header"
    document.sections[0].footer.paragraphs[0].text = "Private and confidential"
    document.add_paragraph("Page body content")

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, record = upload(client, save_docx(document))
        version_id = record["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        page_lines = [
            (72, 38, "Resident services header", None),
            (72, 120, "Page body content", None),
            (72, 760, "Private and confidential", None),
        ]
        write_pdf(pdf_path, [page_lines, page_lines])
        mapped = poll_map(client, version_id)
        assert mapped["status"] == "partial"  # Body occurs twice but only one source block.
        headers = [
            region for region in mapped["regions"] if region["element_type"] == "header_paragraph"
        ]
        footers = [
            region for region in mapped["regions"] if region["element_type"] == "footer_paragraph"
        ]
        assert {region["page_number"] for region in headers} == {1, 2}
        assert {region["page_number"] for region in footers} == {1, 2}
        assert all(region["interactive"] for region in headers + footers)

        pdf_path.write_bytes(b"not a PDF")
        failed = poll_map(client, version_id)
        assert failed["status"] == "failed"
        assert failed["regions"] == []
        assert "error" not in failed
        assert "PDF remains available" in failed["status_detail"]


def test_render_map_requires_an_existing_word_pdf(tmp_path: Path, monkeypatch) -> None:
    document = Document()
    document.add_paragraph("Not rendered yet")
    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        _workspace, record = upload(client, save_docx(document))
        response = client.get(
            f"/api/document-versions/{record['version_id']}/render-map"
        )
        assert response.status_code == 200
        assert response.json()["status"] == "not_requested"
        missing = client.get("/api/document-versions/not-a-version/render-map")
        assert missing.status_code == 404


def test_lists_and_multiple_table_paragraphs_receive_exact_regions(
    tmp_path: Path,
    monkeypatch,
) -> None:
    document = Document()
    document.add_heading("Structured locations", level=1)
    document.add_paragraph("First numbered procedure", style="List Number")
    document.add_paragraph("Nested bullet procedure", style="List Bullet 2")
    first_table = document.add_table(rows=1, cols=1)
    first_table.cell(0, 0).paragraphs[0].text = "First table paragraph"
    first_table.cell(0, 0).add_paragraph("Second paragraph in the same cell")
    second_table = document.add_table(rows=1, cols=1)
    second_table.cell(0, 0).paragraphs[0].text = "Second table paragraph"

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, record = upload(client, save_docx(document))
        version_id = record["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        lines = [
            (72, 80, "Structured locations", None),
            (82, 120, "First numbered procedure", None),
            (94, 150, "Nested bullet procedure", None),
            (90, 210, "First table paragraph", None),
            (90, 240, "Second paragraph in the same cell", None),
            (90, 310, "Second table paragraph", None),
        ]
        write_pdf(pdf_path, [lines])
        mapped = poll_map(client, version_id)
        assert mapped["status"] == "completed"
        element_types = [region["element_type"] for region in mapped["regions"]]
        assert element_types.count("list_item") == 2
        assert element_types.count("table_paragraph") == 3
        table_locations = {
            (
                region["location"]["table_index"],
                region["location"]["row_index"],
                region["location"]["column_index"],
                region["location"]["paragraph_index"],
            )
            for region in mapped["regions"]
            if region["element_type"] == "table_paragraph"
        }
        assert table_locations == {(0, 0, 0, 0), (0, 0, 0, 1), (1, 0, 0, 0)}


def test_configured_confidence_threshold_disables_lower_confidence_duplicates(
    tmp_path: Path,
    monkeypatch,
) -> None:
    duplicate = "Ordered duplicate mapping"
    document = Document()
    document.add_heading("Threshold boundary", level=1)
    document.add_paragraph(duplicate)
    document.add_paragraph(duplicate)
    app = load_test_app(tmp_path, monkeypatch, threshold="0.98")
    with TestClient(app) as client:
        workspace, record = upload(client, save_docx(document))
        version_id = record["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        write_pdf(
            pdf_path,
            [[
                (72, 80, "Threshold boundary", None),
                (72, 130, duplicate, None),
                (72, 165, duplicate, None),
            ]],
        )
        mapped = poll_map(client, version_id)
        duplicate_regions = [
            region for region in mapped["regions"] if region["text_preview"] == duplicate
        ]
        assert mapped["interactive_threshold"] == 0.98
        assert len(duplicate_regions) == 2
        assert all(region["confidence"] == 0.95 for region in duplicate_regions)
        assert all(region["interactive"] is False for region in duplicate_regions)
        assert all(region["read_only"] is True for region in duplicate_regions)
        assert all("below the interactive" in region["reason"] for region in duplicate_regions)
