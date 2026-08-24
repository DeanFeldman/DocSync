from __future__ import annotations

from io import BytesIO
from pathlib import Path
import zipfile

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from lxml import etree


UNIQUE_PHRASE = "DOCSYNC_UNIQUE_SEARCH_PHRASE"
CROSS_RUN_ONLY_PHRASE = "DOCSYNC_CROSS_RUN_ONLY"
EXPECTED_EXHAUSTIVE_OCCURRENCES = 28
EXPECTED_DEFAULT_OCCURRENCES = 26

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPE_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
CHART_NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"


def _text_run(text: str):
    run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    return run


def _append_hyperlink(document: Document, paragraph, text: str) -> None:
    relationship_id = document.part.relate_to(
        "https://example.com/docsync-inventory",
        RT.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = _text_run(text)
    properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    properties.append(style)
    run.insert(0, properties)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_content_control(document: Document, text: str) -> None:
    control = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), "DocSync inventory control")
    properties.append(alias)
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    paragraph.append(_text_run(text))
    content.append(paragraph)
    control.extend((properties, content))
    document.element.body.insert(len(document.element.body) - 1, control)


def _append_text_box(document: Document, text: str) -> None:
    paragraph = parse_xml(
        f"""
        <w:p xmlns:w="{WORD_NS}" xmlns:v="urn:schemas-microsoft-com:vml">
          <w:r>
            <w:pict>
              <v:shape id="DocSyncTextBox" style="width:180pt;height:30pt" type="#_x0000_t202">
                <v:textbox>
                  <w:txbxContent>
                    <w:p><w:r><w:t>{text}</w:t></w:r></w:p>
                  </w:txbxContent>
                </v:textbox>
              </v:shape>
            </w:pict>
          </w:r>
        </w:p>
        """
    )
    document.element.body.insert(len(document.element.body) - 1, paragraph)


def _append_field_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(f"Field context {text} before ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' DATE \\@ "d MMMM yyyy" '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1 January 2026"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, result, end):
        run = OxmlElement("w:r")
        run.append(node)
        paragraph._p.append(run)
    paragraph.add_run(" payable immediately")


def _append_revision_paragraphs(document: Document, text: str) -> None:
    insertion = OxmlElement("w:p")
    inserted = OxmlElement("w:ins")
    inserted.set(qn("w:id"), "42")
    inserted.set(qn("w:author"), "DocSync")
    inserted.append(_text_run(text))
    insertion.append(inserted)
    document.element.body.insert(len(document.element.body) - 1, insertion)

    deletion = OxmlElement("w:p")
    deleted = OxmlElement("w:del")
    deleted.set(qn("w:id"), "43")
    deleted.set(qn("w:author"), "DocSync")
    run = OxmlElement("w:r")
    value = OxmlElement("w:delText")
    value.text = text
    run.append(value)
    deleted.append(run)
    deletion.append(deleted)
    document.element.body.insert(len(document.element.body) - 1, deletion)


def _append_drawing_text(document: Document, text: str) -> None:
    paragraph = parse_xml(
        f"""
        <w:p xmlns:w="{WORD_NS}"
             xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
             xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
             xmlns:c="{CHART_NS}"
             xmlns:r="{REL_NS}">
          <w:r>
            <w:drawing>
              <wp:inline distT="0" distB="0" distL="0" distR="0">
                <wp:extent cx="4572000" cy="2286000"/>
                <wp:effectExtent l="0" t="0" r="0" b="0"/>
                <wp:docPr id="1000" name="DocSync inventory chart"/>
                <wp:cNvGraphicFramePr/>
                <a:graphic>
                  <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/chart">
                    <c:chart r:id="rIdDocSyncChart"/>
                  </a:graphicData>
                </a:graphic>
              </wp:inline>
            </w:drawing>
          </w:r>
        </w:p>
        """
    )
    document.element.body.insert(len(document.element.body) - 1, paragraph)


def _add_relationship(root: etree._Element, relationship_id: str, kind: str, target: str) -> None:
    node = etree.SubElement(root, f"{{{PACKAGE_REL_NS}}}Relationship")
    node.set("Id", relationship_id)
    node.set(
        "Type",
        f"http://schemas.openxmlformats.org/officeDocument/2006/relationships/{kind}",
    )
    node.set("Target", target)


def _add_content_type(root: etree._Element, part_name: str, kind: str) -> None:
    node = etree.SubElement(root, f"{{{CONTENT_TYPE_NS}}}Override")
    node.set("PartName", part_name)
    content_type = (
        "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"
        if kind == "chart"
        else f"application/vnd.openxmlformats-officedocument.wordprocessingml.{kind}+xml"
    )
    node.set("ContentType", content_type)


def _word_part_xml(kind: str, text: str) -> bytes:
    container = {"footnotes": "footnote", "endnotes": "endnote"}[kind]
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
    <w:{kind} xmlns:w="{WORD_NS}">
      <w:{container} w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:{container}>
      <w:{container} w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:{container}>
      <w:{container} w:id="1"><w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:{container}>
    </w:{kind}>""".encode("utf-8")


def _comments_xml(text: str) -> bytes:
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
    <w:comments xmlns:w="{WORD_NS}">
      <w:comment w:id="0" w:author="DocSync"><w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:comment>
    </w:comments>""".encode("utf-8")


def _chart_xml(text: str) -> bytes:
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
    <c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DRAWING_NS}">
      <c:date1904 val="0"/><c:lang val="en-US"/><c:roundedCorners val="0"/>
      <c:chart>
        <c:title>
          <c:tx><c:rich><a:bodyPr/><a:lstStyle/><a:p><a:r><a:rPr lang="en-US"/><a:t>{text}</a:t></a:r></a:p></c:rich></c:tx>
          <c:layout/><c:overlay val="0"/>
        </c:title>
        <c:autoTitleDeleted val="0"/>
        <c:plotArea>
          <c:layout/>
          <c:barChart><c:barDir val="col"/><c:grouping val="clustered"/><c:varyColors val="0"/><c:gapWidth val="150"/><c:axId val="164527536"/><c:axId val="164527752"/></c:barChart>
          <c:catAx><c:axId val="164527536"/><c:scaling><c:orientation val="minMax"/></c:scaling><c:delete val="0"/><c:axPos val="b"/><c:tickLblPos val="nextTo"/><c:crossAx val="164527752"/><c:crosses val="autoZero"/><c:auto val="1"/><c:lblAlgn val="ctr"/><c:lblOffset val="100"/><c:noMultiLvlLbl val="0"/></c:catAx>
          <c:valAx><c:axId val="164527752"/><c:scaling><c:orientation val="minMax"/></c:scaling><c:delete val="0"/><c:axPos val="l"/><c:majorGridlines/><c:numFmt formatCode="General" sourceLinked="1"/><c:tickLblPos val="nextTo"/><c:crossAx val="164527536"/><c:crosses val="autoZero"/><c:crossBetween val="between"/></c:valAx>
        </c:plotArea>
        <c:plotVisOnly val="1"/><c:dispBlanksAs val="gap"/><c:showDLblsOverMax val="0"/>
      </c:chart>
    </c:chartSpace>""".encode("utf-8")


def _inject_note_and_comment_parts(payload: bytes, text: str) -> bytes:
    with zipfile.ZipFile(BytesIO(payload)) as archive:
        infos = archive.infolist()
        files = {info.filename: archive.read(info) for info in infos if not info.is_dir()}

    parser = etree.XMLParser(remove_blank_text=False)
    relationships = etree.fromstring(files["word/_rels/document.xml.rels"], parser)
    content_types = etree.fromstring(files["[Content_Types].xml"], parser)
    document = etree.fromstring(files["word/document.xml"], parser)
    body = document.find(f"{{{WORD_NS}}}body")
    assert body is not None
    section_properties = body.find(f"{{{WORD_NS}}}sectPr")
    insertion_index = body.index(section_properties) if section_properties is not None else len(body)

    for relationship_id, kind, target in (
        ("rIdDocSyncFootnotes", "footnotes", "footnotes.xml"),
        ("rIdDocSyncEndnotes", "endnotes", "endnotes.xml"),
        ("rIdDocSyncComments", "comments", "comments.xml"),
        ("rIdDocSyncChart", "chart", "charts/chart1.xml"),
    ):
        _add_relationship(relationships, relationship_id, kind, target)

    for part_name, kind in (
        ("/word/footnotes.xml", "footnotes"),
        ("/word/endnotes.xml", "endnotes"),
        ("/word/comments.xml", "comments"),
        ("/word/charts/chart1.xml", "chart"),
    ):
        _add_content_type(content_types, part_name, kind)

    reference_paragraph = etree.fromstring(
        f"""
        <w:p xmlns:w="{WORD_NS}">
          <w:commentRangeStart w:id="0"/>
          <w:r><w:t>Inventory references</w:t></w:r>
          <w:commentRangeEnd w:id="0"/>
          <w:r><w:commentReference w:id="0"/></w:r>
          <w:r><w:footnoteReference w:id="1"/></w:r>
          <w:r><w:endnoteReference w:id="1"/></w:r>
        </w:p>
        """,
        parser,
    )
    body.insert(insertion_index, reference_paragraph)

    files["word/document.xml"] = etree.tostring(
        document, encoding="UTF-8", xml_declaration=True, standalone=True
    )
    files["word/_rels/document.xml.rels"] = etree.tostring(
        relationships, encoding="UTF-8", xml_declaration=True, standalone=True
    )
    files["[Content_Types].xml"] = etree.tostring(
        content_types, encoding="UTF-8", xml_declaration=True, standalone=True
    )
    files["word/footnotes.xml"] = _word_part_xml("footnotes", text)
    files["word/endnotes.xml"] = _word_part_xml("endnotes", text)
    files["word/comments.xml"] = _comments_xml(text)
    files["word/charts/chart1.xml"] = _chart_xml(text)

    output = BytesIO()
    seen: set[str] = set()
    with zipfile.ZipFile(output, "w") as archive:
        for info in infos:
            if info.is_dir():
                archive.writestr(info, b"")
                continue
            archive.writestr(info, files[info.filename])
            seen.add(info.filename)
        for name in (
            "word/footnotes.xml",
            "word/endnotes.xml",
            "word/comments.xml",
            "word/charts/chart1.xml",
        ):
            if name not in seen:
                archive.writestr(name, files[name], compress_type=zipfile.ZIP_DEFLATED)
    return output.getvalue()


def make_exhaustive_text_inventory_docx() -> bytes:
    document = Document()
    document.add_paragraph(
        f"{UNIQUE_PHRASE} then {UNIQUE_PHRASE} and finally {UNIQUE_PHRASE}"
    )
    document.add_heading(UNIQUE_PHRASE, level=1)
    document.add_paragraph(UNIQUE_PHRASE, style="List Number")
    document.add_paragraph(UNIQUE_PHRASE, style="List Bullet")

    bold = document.add_paragraph()
    bold.add_run(UNIQUE_PHRASE).bold = True
    italic = document.add_paragraph()
    italic.add_run(UNIQUE_PHRASE).italic = True
    split = document.add_paragraph()
    for index, value in enumerate(("DOCSYNC_", "UNIQUE_", "SEARCH_", "PHRASE")):
        run = split.add_run(value)
        run.bold = index == 1
        run.italic = index == 2
    cross_run_only = document.add_paragraph()
    for value in ("DOCSYNC_", "CROSS_", "RUN_", "ONLY"):
        cross_run_only.add_run(value)

    hyperlink = document.add_paragraph("Hyperlink: ")
    _append_hyperlink(document, hyperlink, UNIQUE_PHRASE)

    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).text = UNIQUE_PHRASE
    table.cell(0, 1).text = "First cell paragraph"
    table.cell(0, 1).add_paragraph(UNIQUE_PHRASE)
    nested = table.cell(0, 2).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = UNIQUE_PHRASE

    section = document.sections[0]
    document.settings.odd_and_even_pages_header_footer = True
    section.different_first_page_header_footer = True
    section.header.paragraphs[0].text = UNIQUE_PHRASE
    section.first_page_header.paragraphs[0].text = UNIQUE_PHRASE
    section.even_page_header.paragraphs[0].text = UNIQUE_PHRASE
    section.footer.paragraphs[0].text = UNIQUE_PHRASE
    section.first_page_footer.paragraphs[0].text = UNIQUE_PHRASE
    section.even_page_footer.paragraphs[0].text = UNIQUE_PHRASE

    _append_content_control(document, UNIQUE_PHRASE)
    _append_text_box(document, UNIQUE_PHRASE)
    _append_field_paragraph(document, UNIQUE_PHRASE)
    _append_revision_paragraphs(document, UNIQUE_PHRASE)
    _append_drawing_text(document, UNIQUE_PHRASE)

    output = BytesIO()
    document.save(output)
    return _inject_note_and_comment_parts(output.getvalue(), UNIQUE_PHRASE)


def write_fixture(path: Path) -> Path:
    path.write_bytes(make_exhaustive_text_inventory_docx())
    return path
