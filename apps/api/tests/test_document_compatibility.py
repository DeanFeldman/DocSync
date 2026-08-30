from __future__ import annotations

from io import BytesIO
from pathlib import Path
import sys
import time
import zipfile

from docx import Document

API_DIR = Path(__file__).resolve().parents[1]
if str(API_DIR) not in sys.path:
    sys.path.insert(0, str(API_DIR))

from app.text_inventory_service import build_text_inventory
from docx_inventory_fixtures import make_exhaustive_text_inventory_docx


def _bytes(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _core_document() -> bytes:
    document = Document()
    document.add_heading("Compatibility heading", 1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Run-split ").bold = True
    paragraph.add_run("business text").italic = True
    document.add_paragraph("Soft line one\nSoft line two")
    document.add_paragraph("Indented bullet", style="List Bullet")
    document.add_paragraph("Ordered item", style="List Number")
    return _bytes(document)


def _table_document() -> bytes:
    document = Document()
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Merged agreement field"
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 0).text = "Repeated wording"
    table.cell(1, 1).text = "Repeated wording"
    table.cell(2, 0).text = "Table owner"
    table.cell(2, 1).text = ""
    return _bytes(document)


def _header_footer_document() -> bytes:
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Compatibility header"
    section.footer.paragraphs[0].text = "Compatibility footer"
    document.add_paragraph("Section body text")
    return _bytes(document)


def evaluate_compatibility(name: str, payload: bytes) -> dict:
    started = time.perf_counter()
    assert zipfile.is_zipfile(BytesIO(payload))
    Document(BytesIO(payload))
    inventory = build_text_inventory(payload, document_id=name, version_id="compatibility")
    editable = 0
    readonly = 0
    for segment in inventory.segments:
        can_edit, _reason = segment.editability_for_range(0, len(segment.text))
        editable += int(can_edit)
        readonly += int(not can_edit)
    return {
        "fixture": name,
        "parse_success": True,
        "generated_docx_valid": True,
        "mapped_blocks": len(inventory.segments),
        "editable_blocks": editable,
        "readonly_blocks": readonly,
        "duration_ms": round((time.perf_counter() - started) * 1000, 2),
    }


def test_compatibility_corpus_preserves_safe_editability_boundaries():
    corpus = {
        "core-paragraphs-lists.docx": _core_document(),
        "tables-merged-repeated.docx": _table_document(),
        "header-footer.docx": _header_footer_document(),
        "exhaustive-unsupported-structures.docx": make_exhaustive_text_inventory_docx(),
    }
    results = [evaluate_compatibility(name, payload) for name, payload in corpus.items()]
    assert len(results) == 4
    assert all(item["parse_success"] and item["generated_docx_valid"] for item in results)
    assert all(item["mapped_blocks"] >= item["editable_blocks"] for item in results)
    exhaustive = next(item for item in results if item["fixture"].startswith("exhaustive"))
    assert exhaustive["readonly_blocks"] > 0
    assert sum(item["editable_blocks"] for item in results) > 0
