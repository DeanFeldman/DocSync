from __future__ import annotations

from copy import deepcopy
import hashlib
import importlib
import io
import os
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.testclient import TestClient


API_DIR = Path(__file__).resolve().parents[1]
if str(API_DIR) not in sys.path:
    sys.path.insert(0, str(API_DIR))

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def load_test_app(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("DOCUMENTSYNC_DATA_DIR", str(tmp_path / "data"))
    monkeypatch.setenv(
        "DOCUMENTSYNC_DATABASE_URL",
        f"sqlite:///{tmp_path / 'editor-workflow.db'}",
    )
    monkeypatch.setenv("DOCUMENTSYNC_SESSION_TOKEN", "")
    monkeypatch.setenv("DOCUMENTSYNC_NEAR_MATCH_THRESHOLD", "0.82")
    monkeypatch.delenv("DOCUMENTSYNC_WEB_DIST", raising=False)

    for module_name in list(sys.modules):
        if module_name == "app" or module_name.startswith("app."):
            del sys.modules[module_name]

    return importlib.import_module("app.main").app


def save_docx(document: Document) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def make_paragraph_docx(title: str, paragraphs: list[str]) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    for text in paragraphs:
        document.add_paragraph(text)
    return save_docx(document)


def make_custom_style_docx(
    title: str,
    shared_text: str,
    *,
    numbered: bool,
) -> bytes:
    document = Document()
    document.styles.add_style("Clause2Sub", WD_STYLE_TYPE.PARAGRAPH)
    document.add_heading(title, level=1)
    paragraph = document.add_paragraph(shared_text, style="Clause2Sub")
    if numbered:
        numbering = deepcopy(
            document.styles["List Number"].element.pPr.numPr
        )
        paragraph._p.get_or_add_pPr().append(numbering)
    return save_docx(document)


def make_rich_docx(exact_text: str, near_text: str) -> bytes:
    document = Document()
    document.add_heading("Rich editor contract", level=1)

    rich = document.add_paragraph()
    bold = rich.add_run("Bold")
    bold.bold = True
    italic = rich.add_run(" Italic")
    italic.italic = True
    underlined = rich.add_run(" Underlined")
    underlined.underline = True

    document.add_paragraph("First numbered item", style="List Number")
    document.add_paragraph("Indented bullet item", style="List Bullet 2")
    document.add_paragraph(exact_text)
    document.add_paragraph(near_text)
    return save_docx(document)


def make_read_only_table_docx(shared_text: str) -> bytes:
    document = Document()
    document.add_heading("Read-only structures", level=1)
    document.add_paragraph(shared_text)
    table = document.add_table(rows=1, cols=2)
    cell = table.cell(0, 0).merge(table.cell(0, 1))
    cell.paragraphs[0].text = "Read-only merged content"
    return save_docx(document)


def make_expanded_table_docx(
    title: str,
    shared_text: str,
    secondary_text: str,
) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph("Body content before the table.")
    # The same words in body context must not become a table match target.
    document.add_paragraph(shared_text)
    table = document.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = ""
    primary_run = cell.paragraphs[0].add_run(shared_text)
    primary_run.bold = True
    secondary = cell.add_paragraph(style="List Bullet")
    secondary_run = secondary.add_run(secondary_text)
    secondary_run.italic = True
    secondary.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.cell(0, 1).text = f"{title} untouched value"
    document.add_paragraph("Body content after the table.")
    return save_docx(document)


def make_unsupported_table_structures_docx() -> bytes:
    document = Document()
    document.add_heading("Unsupported table structures", level=1)
    table = document.add_table(rows=2, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Merged content appears once"

    nested_cell = table.cell(1, 0)
    nested_cell.text = "Outer text beside a nested table"
    nested = nested_cell.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Nested content"

    object_cell = table.cell(1, 1)
    object_cell.text = "Caption with an unsupported image"
    object_cell.paragraphs[0].add_run().add_picture(
        str(API_DIR.parents[1] / "build" / "icon.png")
    )
    object_cell.add_paragraph("")
    return save_docx(document)


def upload_set(
    client: TestClient,
    originals: dict[str, bytes],
    *,
    name: str = "Editor workflow",
) -> dict:
    response = client.post(
        "/api/document-sets",
        data={"name": name},
        files=[
            (
                "files",
                (filename, io.BytesIO(payload), DOCX_MEDIA_TYPE),
            )
            for filename, payload in originals.items()
        ],
    )
    assert response.status_code == 201, response.text
    return response.json()


def documents_by_name(workspace: dict) -> dict[str, dict]:
    return {document["name"]: document for document in workspace["documents"]}


def read_editor_content(client: TestClient, version_id: str) -> dict:
    response = client.get(f"/api/document-versions/{version_id}/editor-content")
    assert response.status_code == 200, response.text
    return response.json()


def block_with_text(content: dict, text: str) -> dict:
    return next(block for block in content["blocks"] if block["text"] == text)


def current_versions(workspace: dict, document_names: list[str]) -> dict[str, str]:
    by_name = documents_by_name(workspace)
    return {
        by_name[name]["id"]: by_name[name]["version_id"]
        for name in document_names
    }


def version_count(client: TestClient, document_id: str) -> int:
    response = client.get(f"/api/documents/{document_id}/versions")
    assert response.status_code == 200, response.text
    return len(response.json()["versions"])


def test_table_paragraphs_are_individual_safe_rich_versioned_blocks(
    tmp_path: Path,
    monkeypatch,
) -> None:
    shared = "Annual resident safety assessment"
    alpha_secondary = "Review emergency contacts every quarter"
    beta_secondary = "Review emergency contact details every quarter"
    originals = {
        "Alpha.docx": make_expanded_table_docx(
            "Alpha",
            shared,
            alpha_secondary,
        ),
        "Beta.docx": make_expanded_table_docx(
            "Beta",
            shared,
            beta_secondary,
        ),
    }

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace = upload_set(client, originals, name="Expanded table editing")
        documents = documents_by_name(workspace)
        alpha_content = read_editor_content(
            client,
            documents["Alpha.docx"]["version_id"],
        )
        beta_content = read_editor_content(
            client,
            documents["Beta.docx"]["version_id"],
        )

        alpha_table = [
            block
            for block in alpha_content["blocks"]
            if block["element_type"] == "table_paragraph"
            and block["column_index"] == 0
        ]
        assert [block["paragraph_index"] for block in alpha_table] == [0, 1]
        assert [block["text"] for block in alpha_table] == [
            shared,
            alpha_secondary,
        ]
        assert all(block["supported"] for block in alpha_table)
        assert alpha_table[0]["delta"]["ops"][0]["attributes"] == {"bold": True}
        assert alpha_table[1]["list_type"] == "bullet"
        assert alpha_table[1]["alignment"] == "right"
        assert alpha_table[1]["delta"]["ops"][0]["attributes"] == {
            "italic": True
        }
        assert (
            block_with_text(alpha_content, "Body content before the table.")["order"]
            < alpha_table[0]["order"]
            < alpha_table[1]["order"]
            < block_with_text(alpha_content, "Body content after the table.")["order"]
        )

        source = alpha_table[0]
        beta_source = next(
            block
            for block in beta_content["blocks"]
            if block["element_type"] == "table_paragraph"
            and block["text"] == shared
        )
        exact = client.get(
            f"/api/document-elements/{source['element_id']}/matches"
        )
        assert exact.status_code == 200, exact.text
        assert [
            item["element_id"] for item in exact.json()["exact_matches"]
        ] == [beta_source["element_id"]]
        assert all(
            item["element_type"] == "table_paragraph"
            for item in exact.json()["exact_matches"]
        )

        near_source = alpha_table[1]
        near = client.get(
            f"/api/document-elements/{near_source['element_id']}/similar-matches",
            params={"threshold": 0.5},
        )
        assert near.status_code == 200, near.text
        assert any(
            item["text"] == beta_secondary
            and item["element_type"] == "table_paragraph"
            and 0.5 <= item["similarity_score"] < 1
            and item["difference_spans"]
            for item in near.json()["matches"]
        )

        heading_attempt = client.post(
            f"/api/document-sets/{workspace['id']}/editor-preview",
            json={
                "base_versions": current_versions(workspace, ["Alpha.docx"]),
                "source_element_id": source["element_id"],
                "edit_mode": "override",
                "targets": [
                    {
                        "element_id": source["element_id"],
                        "replacement_text": shared,
                        "delta": {
                            "ops": [
                                {"insert": shared},
                                {"insert": "\n", "attributes": {"header": 2}},
                            ]
                        },
                    }
                ],
            },
        )
        assert heading_attempt.status_code == 422
        assert "Heading levels cannot be applied" in heading_attempt.json()["detail"]

        payload = {
            "base_versions": current_versions(
                workspace,
                ["Alpha.docx", "Beta.docx"],
            ),
            "source_element_id": source["element_id"],
            "edit_mode": "per_document",
            "targets": [
                {
                    "element_id": source["element_id"],
                    "replacement_text": "Alpha annual safety review",
                },
                {
                    "element_id": beta_source["element_id"],
                    "replacement_text": "Beta annual safety review",
                },
            ],
        }
        preview = client.post(
            f"/api/document-sets/{workspace['id']}/editor-preview",
            json=payload,
        )
        assert preview.status_code == 200, preview.text
        assert preview.json()["writes_performed"] is False
        assert preview.json()["edit_mode"] == "per_document"
        assert all(
            change["element_type"] == "table_paragraph"
            and change["table_index"] == 0
            and change["row_index"] == 0
            and change["column_index"] == 0
            and change["paragraph_index"] == 0
            for document_preview in preview.json()["documents"]
            for change in document_preview["changes"]
        )

        generated = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate",
            json=payload,
        )
        assert generated.status_code == 201, generated.text
        versions = {
            item["document_name"]: item
            for item in generated.json()["versions"]
        }
        for document_name, replacement in (
            ("Alpha.docx", "Alpha annual safety review"),
            ("Beta.docx", "Beta annual safety review"),
        ):
            downloaded = client.get(versions[document_name]["download_url"])
            assert downloaded.status_code == 200
            result = Document(io.BytesIO(downloaded.content))
            cell = result.tables[0].cell(0, 0)
            assert [paragraph.text for paragraph in cell.paragraphs] == [
                replacement,
                alpha_secondary
                if document_name == "Alpha.docx"
                else beta_secondary,
            ]
            assert cell.paragraphs[0].runs[0].bold is True
            assert cell.paragraphs[1].runs[0].italic is True
            assert cell.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.RIGHT
            assert result.tables[0].cell(0, 1).text == (
                f"{Path(document_name).stem} untouched value"
            )
            current_content = read_editor_content(
                client,
                versions[document_name]["version_id"],
            )
            current_cell = [
                block
                for block in current_content["blocks"]
                if block["element_type"] == "table_paragraph"
                and block.get("table_index") == 0
                and block.get("row_index") == 0
                and block.get("column_index") == 0
            ]
            assert [block["text"] for block in current_cell] == [
                replacement,
                alpha_secondary
                if document_name == "Alpha.docx"
                else beta_secondary,
            ]

        alpha_current = read_editor_content(
            client,
            versions["Alpha.docx"]["version_id"],
        )
        alpha_list = next(
            block
            for block in alpha_current["blocks"]
            if block["element_type"] == "table_paragraph"
            and block.get("table_index") == 0
            and block.get("row_index") == 0
            and block.get("column_index") == 0
            and block["paragraph_index"] == 1
        )
        list_replacement = "Updated ordered contact checklist"
        list_update = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate",
            json={
                "base_versions": {
                    documents["Alpha.docx"]["id"]: versions["Alpha.docx"][
                        "version_id"
                    ]
                },
                "source_element_id": alpha_list["element_id"],
                "edit_mode": "override",
                "targets": [
                    {
                        "element_id": alpha_list["element_id"],
                        "replacement_text": list_replacement,
                        "delta": {
                            "ops": [
                                {
                                    "insert": list_replacement,
                                    "attributes": {"italic": True},
                                },
                                {
                                    "insert": "\n",
                                    "attributes": {
                                        "list": "ordered",
                                        "indent": 1,
                                        "align": "center",
                                    },
                                },
                            ]
                        },
                    }
                ],
            },
        )
        assert list_update.status_code == 201, list_update.text
        updated_version = list_update.json()["versions"][0]
        updated_download = client.get(updated_version["download_url"])
        assert updated_download.status_code == 200
        updated_docx = Document(io.BytesIO(updated_download.content))
        updated_cell = updated_docx.tables[0].cell(0, 0)
        assert [paragraph.text for paragraph in updated_cell.paragraphs] == [
            "Alpha annual safety review",
            list_replacement,
        ]
        assert updated_cell.paragraphs[1].style.name.startswith("List Number")
        assert updated_cell.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert updated_cell.paragraphs[1].runs[0].italic is True


def test_unsafe_table_structures_are_visible_once_and_read_only(
    tmp_path: Path,
    monkeypatch,
) -> None:
    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace = upload_set(
            client,
            {
                "Unsafe.docx": make_unsupported_table_structures_docx(),
                "Control.docx": make_paragraph_docx("Control", ["Safe body text"]),
            },
        )
        unsafe = documents_by_name(workspace)["Unsafe.docx"]
        content = read_editor_content(client, unsafe["version_id"])
        table_blocks = [
            block
            for block in content["blocks"]
            if block["element_type"] == "table_paragraph"
        ]

        assert [
            block["text"]
            for block in table_blocks
            if block["text"] == "Merged content appears once"
        ] == ["Merged content appears once"]
        by_text = {block["text"]: block for block in table_blocks}
        assert "merged structure" in by_text[
            "Merged content appears once"
        ]["unsupported_reason"]
        assert "nested structure" in by_text[
            "Outer text beside a nested table"
        ]["unsupported_reason"]
        assert "Drawing or floating object" in by_text[
            "Caption with an unsupported image"
        ]["unsupported_reason"]
        assert all(
            block["read_only"] and not block["supported"]
            for block in table_blocks
        )
        assert all(block["text"] for block in table_blocks)
        assert content["unsupported_count"] >= 3


def test_document_view_exposes_safe_optional_layout_region_contract(
    tmp_path: Path,
    monkeypatch,
) -> None:
    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace = upload_set(
            client,
            {
                "Layout.docx": make_paragraph_docx(
                    "Selectable layout",
                    ["Choose this paragraph from the structured layout."],
                ),
                "Companion.docx": make_paragraph_docx(
                    "Companion layout",
                    ["A second document keeps this a valid document set."],
                ),
            },
        )
        document = documents_by_name(workspace)["Layout.docx"]

        response = client.get(
            f"/api/document-versions/{document['version_id']}/pages"
        )

        assert response.status_code == 200, response.text
        payload = response.json()
        assert payload["document_id"] == document["id"]
        assert payload["version_id"] == document["version_id"]
        assert payload["layout_regions"] == []
        assert payload["pages"][0]["elements"]


def test_editor_content_preserves_delta_structure_and_normalized_matching(
    tmp_path: Path,
    monkeypatch,
) -> None:
    exact_a = "Ｔhe RESIDENT   submits the safety report every month."
    exact_b = "the resident submits the safety report every month."
    near_a = "The caretaker completes the inspection every Friday."
    near_c = "The caretaker completes the inspection every Monday."
    originals = {
        "Alpha.docx": make_rich_docx(exact_a, near_a),
        "Beta.docx": make_paragraph_docx("Beta contract", [exact_b]),
        "Gamma.docx": make_paragraph_docx("Gamma contract", [near_c]),
    }

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace = upload_set(client, originals)
        documents = documents_by_name(workspace)
        alpha = read_editor_content(client, documents["Alpha.docx"]["version_id"])
        beta = read_editor_content(client, documents["Beta.docx"]["version_id"])
        gamma = read_editor_content(client, documents["Gamma.docx"]["version_id"])

        rich = block_with_text(alpha, "Bold Italic Underlined")
        assert rich["delta"] == {
            "ops": [
                {"insert": "Bold", "attributes": {"bold": True}},
                {"insert": " Italic", "attributes": {"italic": True}},
                {"insert": " Underlined", "attributes": {"underline": True}},
                {"insert": "\n"},
            ]
        }
        assert rich["formatting"]["runs"] == [
            {"text": "Bold", "bold": True},
            {"text": " Italic", "italic": True},
            {"text": " Underlined", "underline": True},
        ]

        heading = block_with_text(alpha, "Rich editor contract")
        assert heading["element_type"] == "heading"
        assert heading["delta"]["ops"][-1] == {
            "insert": "\n",
            "attributes": {"header": 1},
        }

        numbered = block_with_text(alpha, "First numbered item")
        assert numbered["element_type"] == "list_item"
        assert numbered["list_type"] == "ordered"
        assert numbered["list_level"] == 0
        assert numbered["delta"]["ops"][-1]["attributes"] == {"list": "ordered"}

        bullet = block_with_text(alpha, "Indented bullet item")
        assert bullet["element_type"] == "list_item"
        assert bullet["list_type"] == "bullet"
        assert bullet["list_level"] == 1
        assert bullet["delta"]["ops"][-1]["attributes"] == {
            "list": "bullet",
            "indent": 1,
        }

        source = block_with_text(alpha, exact_a)
        exact_candidate = block_with_text(beta, exact_b)
        assert source["normalized_text"] == exact_candidate["normalized_text"]
        assert source["normalized_text"] == (
            "the resident submits the safety report every month."
        )
        assert source["exact_match_hash"] == exact_candidate["exact_match_hash"]

        matches = client.get(
            f"/api/document-elements/{source['element_id']}/matches"
        )
        assert matches.status_code == 200, matches.text
        assert matches.json()["exact_match_count"] == 1
        assert matches.json()["exact_matches"][0]["element_id"] == (
            exact_candidate["element_id"]
        )
        assert matches.json()["exact_matches"][0]["match_type"] == "exact"

        near_source = block_with_text(alpha, near_a)
        near_candidate = block_with_text(gamma, near_c)
        compared = client.post(
            f"/api/document-elements/{near_source['element_id']}/compare",
            json={"candidate_element_ids": [near_candidate["element_id"]]},
        )
        assert compared.status_code == 200, compared.text
        comparison = compared.json()["items"][0]
        assert comparison["match_type"] == "near"
        assert 0 < comparison["similarity_score"] < 1
        assert comparison["diff_spans"] == comparison["difference_spans"]
        assert any(
            span["kind"] == "equal" and span["text"]
            for span in comparison["difference_spans"]
        )
        assert any(
            span["kind"] == "changed"
            and span["source_text"] == "Friday"
            and span["candidate_text"] == "Monday"
            for span in comparison["difference_spans"]
        )

        discovered = client.get(
            f"/api/document-elements/{near_source['element_id']}/similar-matches",
            params={"threshold": 0, "limit": 20},
        )
        assert discovered.status_code == 200, discovered.text
        discovered_candidate = next(
            item
            for item in discovered.json()["matches"]
            if item["element_id"] == near_candidate["element_id"]
        )
        assert discovered_candidate["match_type"] == "near"
        assert discovered_candidate["algorithm_version"] if (
            "algorithm_version" in discovered_candidate
        ) else discovered.json()["algorithm_version"] == "nfkc-sequence-v1"


def test_exact_matches_exclude_custom_styled_paragraph_with_word_numbering(
    tmp_path: Path,
    monkeypatch,
) -> None:
    shared = (
        "The Occupant understands the financial implications, fees and charges."
    )
    app = load_test_app(tmp_path, monkeypatch)

    with TestClient(app) as client:
        workspace = upload_set(
            client,
            {
                "One.docx": make_custom_style_docx(
                    "One",
                    shared,
                    numbered=False,
                ),
                "Two-numbered.docx": make_custom_style_docx(
                    "Two",
                    shared,
                    numbered=True,
                ),
                "Three.docx": make_custom_style_docx(
                    "Three",
                    shared,
                    numbered=False,
                ),
            },
            name="Mixed Word numbering",
        )
        documents = documents_by_name(workspace)
        one = read_editor_content(client, documents["One.docx"]["version_id"])
        two = read_editor_content(
            client,
            documents["Two-numbered.docx"]["version_id"],
        )
        three = read_editor_content(
            client,
            documents["Three.docx"]["version_id"],
        )
        source = block_with_text(one, shared)
        numbered = block_with_text(two, shared)
        compatible = block_with_text(three, shared)

        assert source["element_type"] == "paragraph"
        assert compatible["element_type"] == "paragraph"
        assert numbered["element_type"] == "list_item"

        groups = [
            group
            for group in workspace["link_groups"]
            if group["representative_text"] == shared
        ]
        assert len(groups) == 1
        assert {
            member["document_name"] for member in groups[0]["members"]
        } == {"One.docx", "Three.docx"}

        matches = client.get(
            f"/api/document-elements/{source['element_id']}/matches"
        )
        assert matches.status_code == 200, matches.text
        assert matches.json()["exact_match_count"] == 1
        assert [
            item["element_id"] for item in matches.json()["exact_matches"]
        ] == [compatible["element_id"]]
        assert {
            member["element_type"]
            for member in matches.json()["link_group"]["members"]
        } == {"paragraph"}

        preview = client.post(
            f"/api/document-sets/{workspace['id']}/editor-preview",
            json={
                "base_versions": current_versions(
                    workspace,
                    ["One.docx", "Two-numbered.docx", "Three.docx"],
                ),
                "source_element_id": source["element_id"],
                "edit_mode": "shared",
                "targets": [
                    {
                        "element_id": source["element_id"],
                        "replacement_text": f"{shared} Updated",
                    },
                    {
                        "element_id": compatible["element_id"],
                        "replacement_text": f"{shared} Updated",
                    },
                ],
            },
        )
        assert preview.status_code == 200, preview.text


def test_editor_generation_can_be_queued_and_reconciled_by_status(
    tmp_path: Path,
    monkeypatch,
) -> None:
    shared = "The resident submits the signed schedule every month."
    app = load_test_app(tmp_path, monkeypatch)

    with TestClient(app) as client:
        workspace = upload_set(
            client,
            {
                "Alpha.docx": make_paragraph_docx("Alpha", [shared]),
                "Beta.docx": make_paragraph_docx("Beta", [shared]),
            },
            name="Queued generation",
        )
        documents = documents_by_name(workspace)
        alpha = read_editor_content(
            client,
            documents["Alpha.docx"]["version_id"],
        )
        beta = read_editor_content(
            client,
            documents["Beta.docx"]["version_id"],
        )
        source = block_with_text(alpha, shared)
        candidate = block_with_text(beta, shared)
        request = {
            "base_versions": current_versions(
                workspace,
                ["Alpha.docx", "Beta.docx"],
            ),
            "source_element_id": source["element_id"],
            "edit_mode": "shared",
            "targets": [
                {
                    "element_id": source["element_id"],
                    "replacement_text": f"{shared} Updated",
                },
                {
                    "element_id": candidate["element_id"],
                    "replacement_text": f"{shared} Updated",
                },
            ],
        }

        queued = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate-async",
            json=request,
        )
        assert queued.status_code == 202, queued.text
        assert queued.json()["status"] == "queued"
        operation_id = queued.json()["operation_id"]

        completed = client.get(f"/api/editor-operations/{operation_id}")
        assert completed.status_code == 200, completed.text
        assert completed.json()["status"] == "completed"
        assert len(completed.json()["versions"]) == 2
        assert completed.json()["document_set"]["id"] == workspace["id"]
        assert completed.json()["download_url"].endswith(
            f"/{operation_id}/download"
        )

        stale_duplicate = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate-async",
            json=request,
        )
        assert stale_duplicate.status_code == 409


def test_match_decisions_persist_and_control_near_match_targets(
    tmp_path: Path,
    monkeypatch,
) -> None:
    texts = {
        "Alpha.docx": "The resident submits the safety report every month.",
        "Beta.docx": "The resident submits the safety report each month.",
        "Gamma.docx": "The resident submits the safety report every calendar month.",
        "Delta.docx": "The resident submits the safety report every single month.",
    }
    originals = {
        name: make_paragraph_docx("Common heading", [text])
        for name, text in texts.items()
    }

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace = upload_set(client, originals)
        documents = documents_by_name(workspace)
        contents = {
            name: read_editor_content(client, document["version_id"])
            for name, document in documents.items()
        }
        blocks = {
            name: block_with_text(contents[name], text)
            for name, text in texts.items()
        }
        source = blocks["Alpha.docx"]

        saved = client.post(
            f"/api/document-elements/{source['element_id']}/match-decisions",
            json={
                "decisions": [
                    {
                        "candidate_element_id": blocks["Beta.docx"]["element_id"],
                        "status": "confirmed",
                    },
                    {
                        "candidate_element_id": blocks["Gamma.docx"]["element_id"],
                        "status": "ignored",
                    },
                    {
                        "candidate_element_id": blocks["Delta.docx"]["element_id"],
                        "status": "removed",
                    },
                ]
            },
        )
        assert saved.status_code == 200, saved.text
        assert {
            item["candidate_element_id"]: item["status"]
            for item in saved.json()["decisions"]
        } == {
            blocks["Beta.docx"]["element_id"]: "confirmed",
            blocks["Gamma.docx"]["element_id"]: "ignored",
            blocks["Delta.docx"]["element_id"]: "removed",
        }

        reloaded = client.get(
            f"/api/document-elements/{source['element_id']}/similar-matches",
            params={"threshold": 0},
        )
        assert reloaded.status_code == 200, reloaded.text
        persisted = {
            item["element_id"]: item["decision"]
            for item in reloaded.json()["matches"]
        }
        assert persisted[blocks["Beta.docx"]["element_id"]] == "confirmed"
        assert persisted[blocks["Gamma.docx"]["element_id"]] == "ignored"
        assert persisted[blocks["Delta.docx"]["element_id"]] == "removed"

        confirmed_request = {
            "base_versions": current_versions(
                workspace, ["Alpha.docx", "Beta.docx"]
            ),
            "source_element_id": source["element_id"],
            "edit_mode": "shared",
            "targets": [
                {
                    "element_id": source["element_id"],
                    "replacement_text": "The revised safety report is due monthly.",
                },
                {
                    "element_id": blocks["Beta.docx"]["element_id"],
                    "replacement_text": "The revised safety report is due monthly.",
                },
            ],
        }
        accepted = client.post(
            f"/api/document-sets/{workspace['id']}/editor-preview",
            json=confirmed_request,
        )
        assert accepted.status_code == 200, accepted.text
        assert accepted.json()["affected_document_count"] == 2

        for name in ("Gamma.docx", "Delta.docx"):
            rejected = client.post(
                f"/api/document-sets/{workspace['id']}/editor-preview",
                json={
                    "base_versions": current_versions(
                        workspace, ["Alpha.docx", name]
                    ),
                    "source_element_id": source["element_id"],
                    "edit_mode": "shared",
                    "targets": [
                        {
                            "element_id": source["element_id"],
                            "replacement_text": "Revised source",
                        },
                        {
                            "element_id": blocks[name]["element_id"],
                            "replacement_text": "Revised candidate",
                        },
                    ],
                },
            )
            assert rejected.status_code == 422
            assert "confirmed" in rejected.json()["detail"].casefold()


def test_preview_is_side_effect_free_and_generation_is_target_specific_and_versioned(
    tmp_path: Path,
    monkeypatch,
) -> None:
    shared = "The manager submits the compliance report every month."
    originals = {
        "Alpha.docx": make_paragraph_docx("Alpha agreement", [shared]),
        "Beta.docx": make_paragraph_docx("Beta agreement", [shared]),
    }

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace = upload_set(client, originals)
        documents = documents_by_name(workspace)
        contents = {
            name: read_editor_content(client, document["version_id"])
            for name, document in documents.items()
        }
        blocks = {
            name: block_with_text(contents[name], shared)
            for name in documents
        }
        original_version_ids = {
            name: document["version_id"] for name, document in documents.items()
        }
        original_checksums = {
            name: hashlib.sha256(payload).hexdigest()
            for name, payload in originals.items()
        }
        request = {
            "base_versions": current_versions(
                workspace, ["Alpha.docx", "Beta.docx"]
            ),
            "source_element_id": blocks["Alpha.docx"]["element_id"],
            "edit_mode": "per_document",
            "targets": [
                {
                    "element_id": blocks["Alpha.docx"]["element_id"],
                    "replacement_text": "Alpha submits its report every Friday.",
                },
                {
                    "element_id": blocks["Beta.docx"]["element_id"],
                    "replacement_text": "Beta submits its report every Monday.",
                },
            ],
        }

        preview = client.post(
            f"/api/document-sets/{workspace['id']}/editor-preview",
            json=request,
        )
        assert preview.status_code == 200, preview.text
        assert preview.json()["status"] == "previewed"
        assert preview.json()["writes_performed"] is False
        assert preview.json()["affected_document_count"] == 2
        assert preview.json()["affected_location_count"] == 2
        assert {
            (item["document_name"], item["changes"][0]["after"])
            for item in preview.json()["documents"]
        } == {
            ("Alpha.docx", "Alpha submits its report every Friday."),
            ("Beta.docx", "Beta submits its report every Monday."),
        }
        for name, document in documents.items():
            assert version_count(client, document["id"]) == 1
            original_download = client.get(
                f"/api/document-versions/{original_version_ids[name]}/download"
            )
            assert original_download.status_code == 200
            assert original_download.content == originals[name]
        history_before = client.get(
            f"/api/document-sets/{workspace['id']}/history"
        )
        assert history_before.status_code == 200
        assert history_before.json()["events"] == []

        generated = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate",
            json=request,
        )
        assert generated.status_code == 201, generated.text
        result = generated.json()
        assert result["status"] == "completed"
        assert result["edit_mode"] == "per_document"
        assert len(result["versions"]) == 2

        generated_by_name = {
            item["document_name"]: item for item in result["versions"]
        }
        expected_replacements = {
            "Alpha.docx": "Alpha submits its report every Friday.",
            "Beta.docx": "Beta submits its report every Monday.",
        }
        for name, document in documents.items():
            version = generated_by_name[name]
            assert version["version_number"] == 2
            assert version["parent_version_id"] == original_version_ids[name]
            assert version["checksum_sha256"] != original_checksums[name]

            versions_response = client.get(
                f"/api/documents/{document['id']}/versions"
            )
            assert versions_response.status_code == 200
            versions_payload = versions_response.json()
            assert versions_payload["current_version_id"] == version["version_id"]
            assert [
                item["version_number"] for item in versions_payload["versions"]
            ] == [2, 1]
            assert versions_payload["versions"][0]["parent_version_id"] == (
                original_version_ids[name]
            )
            assert versions_payload["versions"][1]["parent_version_id"] is None
            assert versions_payload["versions"][1]["checksum_sha256"] == (
                original_checksums[name]
            )

            original_download = client.get(
                f"/api/document-versions/{original_version_ids[name]}/download"
            )
            assert original_download.status_code == 200
            assert original_download.content == originals[name]
            assert hashlib.sha256(original_download.content).hexdigest() == (
                original_checksums[name]
            )

            generated_download = client.get(version["download_url"])
            assert generated_download.status_code == 200
            assert hashlib.sha256(generated_download.content).hexdigest() == (
                version["checksum_sha256"]
            )
            edited_docx = Document(io.BytesIO(generated_download.content))
            assert expected_replacements[name] in [
                paragraph.text for paragraph in edited_docx.paragraphs
            ]
            unexpected = (
                expected_replacements["Beta.docx"]
                if name == "Alpha.docx"
                else expected_replacements["Alpha.docx"]
            )
            assert unexpected not in [
                paragraph.text for paragraph in edited_docx.paragraphs
            ]

            current_download = client.get(
                f"/api/documents/{document['id']}/download"
            )
            assert current_download.status_code == 200
            assert current_download.content == generated_download.content

        archive = client.get(result["download_url"])
        assert archive.status_code == 200
        with zipfile.ZipFile(io.BytesIO(archive.content)) as bundle:
            assert set(bundle.namelist()) == set(originals)
            for name, replacement in expected_replacements.items():
                bundled = Document(io.BytesIO(bundle.read(name)))
                assert replacement in [
                    paragraph.text for paragraph in bundled.paragraphs
                ]

        stale = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate",
            json=request,
        )
        assert stale.status_code == 409
        assert "changed after this edit was opened" in stale.json()["detail"]
        assert all(
            version_count(client, document["id"]) == 2
            for document in documents.values()
        )


def test_full_override_detaches_block_and_excludes_it_from_shared_matches(
    tmp_path: Path,
    monkeypatch,
) -> None:
    shared = "This clause remains deliberately identical."
    originals = {
        "Alpha.docx": make_paragraph_docx("Alpha", [shared]),
        "Beta.docx": make_paragraph_docx("Beta", [shared]),
    }

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace = upload_set(client, originals)
        documents = documents_by_name(workspace)
        alpha_content = read_editor_content(
            client, documents["Alpha.docx"]["version_id"]
        )
        beta_content = read_editor_content(
            client, documents["Beta.docx"]["version_id"]
        )
        alpha_source = block_with_text(alpha_content, shared)
        beta_source = block_with_text(beta_content, shared)

        generated = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate",
            json={
                "base_versions": current_versions(workspace, ["Alpha.docx"]),
                "source_element_id": alpha_source["element_id"],
                "edit_mode": "full_override",
                "targets": [
                    {
                        "element_id": alpha_source["element_id"],
                        "replacement_text": shared,
                    }
                ],
            },
        )
        assert generated.status_code == 201, generated.text
        assert generated.json()["edit_mode"] == "override"
        assert len(generated.json()["versions"]) == 1
        alpha_version = generated.json()["versions"][0]
        assert alpha_version["document_name"] == "Alpha.docx"

        detached_content = read_editor_content(
            client, alpha_version["version_id"]
        )
        detached = block_with_text(detached_content, shared)
        assert detached["shared_state"] == "detached"
        assert detached["detached_from_shared"] is True

        beta_matches = client.get(
            f"/api/document-elements/{beta_source['element_id']}/matches"
        )
        assert beta_matches.status_code == 200, beta_matches.text
        assert beta_matches.json()["exact_match_count"] == 0
        assert beta_matches.json()["exact_matches"] == []

        detached_matches = client.get(
            f"/api/document-elements/{detached['element_id']}/matches"
        )
        assert detached_matches.status_code == 200, detached_matches.text
        assert detached_matches.json()["exact_match_count"] == 0
        assert detached_matches.json()["exact_matches"] == []

        beta_versions = client.get(
            f"/api/documents/{documents['Beta.docx']['id']}/versions"
        )
        assert beta_versions.status_code == 200
        assert len(beta_versions.json()["versions"]) == 1


def test_invalid_delta_and_unsupported_block_leave_no_version_or_operation(
    tmp_path: Path,
    monkeypatch,
) -> None:
    shared = "A normal editable paragraph."
    originals = {
        "Alpha.docx": make_read_only_table_docx(shared),
        "Beta.docx": make_paragraph_docx("Beta", [shared]),
    }

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace = upload_set(client, originals)
        documents = documents_by_name(workspace)
        alpha = read_editor_content(client, documents["Alpha.docx"]["version_id"])
        editable = block_with_text(alpha, shared)
        read_only = next(
            block
            for block in alpha["blocks"]
            if block["element_type"] == "table_paragraph"
        )
        assert read_only["supported"] is False
        assert read_only["read_only"] is True
        assert "merged structure" in read_only["unsupported_reason"]

        invalid_delta = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate",
            json={
                "base_versions": current_versions(workspace, ["Alpha.docx"]),
                "source_element_id": editable["element_id"],
                "edit_mode": "per_document",
                "targets": [
                    {
                        "element_id": editable["element_id"],
                        "replacement_text": shared,
                        "delta": {"ops": [{"retain": 1}]},
                    }
                ],
            },
        )
        assert invalid_delta.status_code == 422
        assert "retain/delete operations are unsupported" in (
            invalid_delta.json()["detail"]
        )

        unsupported = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate",
            json={
                "base_versions": current_versions(workspace, ["Alpha.docx"]),
                "source_element_id": read_only["element_id"],
                "edit_mode": "override",
                "targets": [
                    {
                        "element_id": read_only["element_id"],
                        "replacement_text": "Attempted replacement",
                    }
                ],
            },
        )
        assert unsupported.status_code == 422
        assert "merged structure" in unsupported.json()["detail"]

        for document in documents.values():
            versions = client.get(
                f"/api/documents/{document['id']}/versions"
            )
            assert versions.status_code == 200
            assert len(versions.json()["versions"]) == 1
            assert versions.json()["current_version_id"] == document["version_id"]

        history = client.get(f"/api/document-sets/{workspace['id']}/history")
        assert history.status_code == 200
        assert history.json()["events"] == []
        generated_root = tmp_path / "data" / "generated"
        assert list(generated_root.rglob("*")) == []


def test_restore_historical_version_creates_new_current_version_and_audit(
    tmp_path: Path,
    monkeypatch,
) -> None:
    original_text = "The original wording remains available."
    replacement_text = "The revised wording is now current."
    originals = {
        "Alpha.docx": make_paragraph_docx("Alpha agreement", [original_text]),
        "Beta.docx": make_paragraph_docx(
            "Beta agreement",
            ["This document remains unchanged."],
        ),
    }

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace = upload_set(client, originals)
        documents = documents_by_name(workspace)
        alpha = documents["Alpha.docx"]
        alpha_v1_id = alpha["version_id"]
        source = block_with_text(
            read_editor_content(client, alpha_v1_id),
            original_text,
        )

        generated = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate",
            json={
                "base_versions": {alpha["id"]: alpha_v1_id},
                "source_element_id": source["element_id"],
                "edit_mode": "per_document",
                "targets": [
                    {
                        "element_id": source["element_id"],
                        "replacement_text": replacement_text,
                    }
                ],
            },
        )
        assert generated.status_code == 201, generated.text
        alpha_v2 = generated.json()["versions"][0]
        alpha_v2_id = alpha_v2["version_id"]

        restored = client.post(
            f"/api/documents/{alpha['id']}/versions/{alpha_v1_id}/restore",
            json={"expected_current_version_id": alpha_v2_id},
        )
        assert restored.status_code == 201, restored.text
        payload = restored.json()
        alpha_v3 = payload["version"]

        assert payload["operation_type"] == "version_restore"
        assert payload["restored_from_version_id"] == alpha_v1_id
        assert payload["restored_from_version_number"] == 1
        assert payload["previous_current_version_id"] == alpha_v2_id
        assert alpha_v3["version_number"] == 3
        assert alpha_v3["parent_version_id"] == alpha_v2_id
        assert alpha_v3["restored_from_version_id"] == alpha_v1_id
        assert alpha_v3["restored_from_version_number"] == 1
        assert alpha_v3["version_id"] not in {alpha_v1_id, alpha_v2_id}
        assert alpha_v3["checksum_sha256"] == hashlib.sha256(
            originals["Alpha.docx"]
        ).hexdigest()

        versions_response = client.get(f"/api/documents/{alpha['id']}/versions")
        assert versions_response.status_code == 200, versions_response.text
        versions_payload = versions_response.json()
        assert versions_payload["current_version_id"] == alpha_v3["version_id"]
        assert [
            version["version_number"] for version in versions_payload["versions"]
        ] == [3, 2, 1]
        assert versions_payload["versions"][0]["operation_type"] == "version_restore"
        assert (
            versions_payload["versions"][0]["restored_from_version_id"]
            == alpha_v1_id
        )
        assert versions_payload["versions"][0]["restored_from_version_number"] == 1

        original_download = client.get(
            f"/api/document-versions/{alpha_v1_id}/download"
        )
        revised_download = client.get(
            f"/api/document-versions/{alpha_v2_id}/download"
        )
        restored_download = client.get(alpha_v3["download_url"])
        current_download = client.get(f"/api/documents/{alpha['id']}/download")
        assert original_download.content == originals["Alpha.docx"]
        assert replacement_text in "\n".join(
            paragraph.text
            for paragraph in Document(io.BytesIO(revised_download.content)).paragraphs
        )
        assert restored_download.content == originals["Alpha.docx"]
        assert current_download.content == originals["Alpha.docx"]

        restored_content = read_editor_content(client, alpha_v3["version_id"])
        assert restored_content["current_version"] is True
        assert block_with_text(restored_content, original_text)

        archive = client.get(payload["download_url"])
        assert archive.status_code == 200, archive.text
        with zipfile.ZipFile(io.BytesIO(archive.content)) as bundle:
            assert set(bundle.namelist()) == set(originals)
            assert bundle.read("Alpha.docx") == originals["Alpha.docx"]
            assert bundle.read("Beta.docx") == originals["Beta.docx"]

        history = client.get(f"/api/document-sets/{workspace['id']}/history")
        assert history.status_code == 200, history.text
        restore_event = history.json()["events"][0]
        assert restore_event["event_type"] == "version_restore"
        assert restore_event["operation_type"] == "version_restore"
        assert restore_event["restored_from_version_id"] == alpha_v1_id
        assert restore_event["restored_from_version_number"] == 1
        assert restore_event["previous_current_version_id"] == alpha_v2_id
        assert restore_event["version_count"] == 1
        assert restore_event["target_count"] == 0


def test_restore_rejects_stale_mismatched_and_current_versions_without_writes(
    tmp_path: Path,
    monkeypatch,
) -> None:
    original_text = "The first version of Alpha."
    originals = {
        "Alpha.docx": make_paragraph_docx("Alpha", [original_text]),
        "Beta.docx": make_paragraph_docx("Beta", ["The first version of Beta."]),
    }

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace = upload_set(client, originals)
        documents = documents_by_name(workspace)
        alpha = documents["Alpha.docx"]
        beta = documents["Beta.docx"]
        source = block_with_text(
            read_editor_content(client, alpha["version_id"]),
            original_text,
        )
        generated = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate",
            json={
                "base_versions": {alpha["id"]: alpha["version_id"]},
                "source_element_id": source["element_id"],
                "edit_mode": "per_document",
                "targets": [
                    {
                        "element_id": source["element_id"],
                        "replacement_text": "The second version of Alpha.",
                    }
                ],
            },
        )
        assert generated.status_code == 201, generated.text
        alpha_v2_id = generated.json()["versions"][0]["version_id"]

        generated_root = tmp_path / "data" / "generated"
        files_before = {
            path.relative_to(generated_root)
            for path in generated_root.rglob("*")
            if path.is_file()
        }

        stale = client.post(
            (
                f"/api/documents/{alpha['id']}/versions/"
                f"{alpha['version_id']}/restore"
            ),
            json={"expected_current_version_id": alpha["version_id"]},
        )
        assert stale.status_code == 409
        assert "changed after version history was opened" in stale.json()["detail"]

        mismatched = client.post(
            (
                f"/api/documents/{alpha['id']}/versions/"
                f"{beta['version_id']}/restore"
            ),
            json={"expected_current_version_id": alpha_v2_id},
        )
        assert mismatched.status_code == 422
        assert "does not belong to this document" in mismatched.json()["detail"]

        already_current = client.post(
            f"/api/documents/{alpha['id']}/versions/{alpha_v2_id}/restore",
            json={"expected_current_version_id": alpha_v2_id},
        )
        assert already_current.status_code == 409
        assert "already current" in already_current.json()["detail"]

        versions = client.get(f"/api/documents/{alpha['id']}/versions")
        assert versions.status_code == 200, versions.text
        assert versions.json()["current_version_id"] == alpha_v2_id
        assert len(versions.json()["versions"]) == 2

        history = client.get(f"/api/document-sets/{workspace['id']}/history")
        assert history.status_code == 200, history.text
        assert len(history.json()["events"]) == 1
        assert history.json()["events"][0]["event_type"] == "editor_edit"

        files_after = {
            path.relative_to(generated_root)
            for path in generated_root.rglob("*")
            if path.is_file()
        }
        assert files_after == files_before


def test_restore_storage_collision_preserves_unowned_directory_and_database(
    tmp_path: Path,
    monkeypatch,
) -> None:
    original_text = "The original collision-test wording."
    originals = {
        "Alpha.docx": make_paragraph_docx("Alpha", [original_text]),
        "Beta.docx": make_paragraph_docx(
            "Beta",
            ["The unchanged collision-test companion."],
        ),
    }

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace = upload_set(client, originals)
        alpha = documents_by_name(workspace)["Alpha.docx"]
        source = block_with_text(
            read_editor_content(client, alpha["version_id"]),
            original_text,
        )
        generated = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate",
            json={
                "base_versions": {alpha["id"]: alpha["version_id"]},
                "source_element_id": source["element_id"],
                "edit_mode": "per_document",
                "targets": [
                    {
                        "element_id": source["element_id"],
                        "replacement_text": "The second collision-test wording.",
                    }
                ],
            },
        )
        assert generated.status_code == 201, generated.text
        alpha_v2_id = generated.json()["versions"][0]["version_id"]

        collision_id = "11111111-1111-4111-8111-111111111111"
        collision_directory = (
            tmp_path
            / "data"
            / "generated"
            / workspace["id"]
            / collision_id
        )
        collision_directory.mkdir(parents=True)
        sentinel = collision_directory / "do-not-delete.txt"
        sentinel.write_text("owned by another operation", encoding="utf-8")

        editor_service = sys.modules["app.editor_service"]
        monkeypatch.setattr(editor_service, "new_id", lambda: collision_id)

        collided = client.post(
            (
                f"/api/documents/{alpha['id']}/versions/"
                f"{alpha['version_id']}/restore"
            ),
            json={"expected_current_version_id": alpha_v2_id},
        )
        assert collided.status_code == 409, collided.text
        assert "storage collision" in collided.json()["detail"]
        assert sentinel.read_text(encoding="utf-8") == "owned by another operation"

        versions = client.get(f"/api/documents/{alpha['id']}/versions")
        assert versions.status_code == 200, versions.text
        assert versions.json()["current_version_id"] == alpha_v2_id
        assert len(versions.json()["versions"]) == 2

        history = client.get(f"/api/document-sets/{workspace['id']}/history")
        assert history.status_code == 200, history.text
        assert len(history.json()["events"]) == 1


def test_restore_serialization_failure_rolls_back_before_commit(
    tmp_path: Path,
    monkeypatch,
) -> None:
    original_text = "The original serialization-test wording."
    originals = {
        "Alpha.docx": make_paragraph_docx("Alpha", [original_text]),
        "Beta.docx": make_paragraph_docx(
            "Beta",
            ["The unchanged serialization-test companion."],
        ),
    }

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app, raise_server_exceptions=False) as client:
        workspace = upload_set(client, originals)
        alpha = documents_by_name(workspace)["Alpha.docx"]
        source = block_with_text(
            read_editor_content(client, alpha["version_id"]),
            original_text,
        )
        generated = client.post(
            f"/api/document-sets/{workspace['id']}/editor-generate",
            json={
                "base_versions": {alpha["id"]: alpha["version_id"]},
                "source_element_id": source["element_id"],
                "edit_mode": "per_document",
                "targets": [
                    {
                        "element_id": source["element_id"],
                        "replacement_text": "The second serialization-test wording.",
                    }
                ],
            },
        )
        assert generated.status_code == 201, generated.text
        alpha_v2_id = generated.json()["versions"][0]["version_id"]

        generated_root = tmp_path / "data" / "generated"
        files_before = {
            path.relative_to(generated_root)
            for path in generated_root.rglob("*")
            if path.is_file()
        }

        document_service = sys.modules["app.document_service"]

        def fail_serialization(_document_set) -> dict:
            raise RuntimeError("Injected refreshed-set serialization failure.")

        monkeypatch.setattr(
            document_service,
            "serialize_document_set",
            fail_serialization,
        )

        failed = client.post(
            (
                f"/api/documents/{alpha['id']}/versions/"
                f"{alpha['version_id']}/restore"
            ),
            json={"expected_current_version_id": alpha_v2_id},
        )
        assert failed.status_code == 500

        versions = client.get(f"/api/documents/{alpha['id']}/versions")
        assert versions.status_code == 200, versions.text
        assert versions.json()["current_version_id"] == alpha_v2_id
        assert len(versions.json()["versions"]) == 2

        history = client.get(f"/api/document-sets/{workspace['id']}/history")
        assert history.status_code == 200, history.text
        assert len(history.json()["events"]) == 1
        assert history.json()["events"][0]["event_type"] == "editor_edit"

        files_after = {
            path.relative_to(generated_root)
            for path in generated_root.rglob("*")
            if path.is_file()
        }
        assert files_after == files_before
