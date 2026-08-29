from __future__ import annotations

import importlib
import io
import os
from pathlib import Path
import sys
import time

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastapi import HTTPException
from fastapi.testclient import TestClient
import pymupdf


API_DIR = Path(__file__).resolve().parents[1]
if str(API_DIR) not in sys.path:
    sys.path.insert(0, str(API_DIR))

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def load_test_app(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("DOCUMENTSYNC_DATA_DIR", str(tmp_path / "data"))
    monkeypatch.setenv(
        "DOCUMENTSYNC_DATABASE_URL",
        f"sqlite:///{tmp_path / 'inline-preview.db'}",
    )
    monkeypatch.setenv("DOCUMENTSYNC_SESSION_TOKEN", "")
    monkeypatch.delenv("DOCUMENTSYNC_WEB_DIST", raising=False)
    for module_name in list(sys.modules):
        if module_name == "app" or module_name.startswith("app."):
            del sys.modules[module_name]
    return importlib.import_module("app.main").app


def docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def upload(client: TestClient, payload: bytes) -> tuple[dict, dict]:
    response = client.post(
        "/api/document-sets",
        data={"name": "Inline preview"},
        files=[
            ("files", ("Primary.docx", io.BytesIO(payload), DOCX_MEDIA_TYPE)),
            ("files", ("Related.docx", io.BytesIO(payload), DOCX_MEDIA_TYPE)),
        ],
    )
    assert response.status_code == 201, response.text
    workspace = response.json()
    primary = next(
        document
        for document in workspace["documents"]
        if document["name"] == "Primary.docx"
    )
    return workspace, primary


def write_pdf(path: Path, lines: list[tuple[float, float, str, float | None]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = pymupdf.open()
    page = pdf.new_page(width=612, height=792)
    for x, y, text, width in lines:
        if width:
            page.insert_textbox(
                pymupdf.Rect(x, y, x + width, y + 180),
                text,
                fontsize=11,
            )
        else:
            page.insert_text((x, y), text, fontsize=11)
    pdf.save(path)
    pdf.close()


def write_multipage_pdf(path: Path, page_texts: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = pymupdf.open()
    for text in page_texts:
        page = pdf.new_page(width=612, height=792)
        page.insert_text((72, 100), text, fontsize=11)
    pdf.save(path)
    pdf.close()


def write_split_paragraph_pdf(path: Path, before: str, after: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = pymupdf.open()
    first_page = pdf.new_page(width=612, height=792)
    first_page.insert_text((72, 750), before, fontsize=11)
    second_page = pdf.new_page(width=612, height=792)
    second_page.insert_text((540, 40), "2", fontsize=11)
    second_page.insert_text((72, 100), after, fontsize=11)
    pdf.save(path)
    pdf.close()


def write_empty_table_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = pymupdf.open()
    page = pdf.new_page(width=612, height=792)
    left, middle, right = 72, 300, 540
    top, middle_y, bottom = 100, 155, 210
    page.draw_rect(pymupdf.Rect(left, top, right, bottom), width=0.8)
    page.draw_line((middle, top), (middle, bottom), width=0.8)
    page.draw_line((left, middle_y), (right, middle_y), width=0.8)
    page.insert_text((82, 132), "Name", fontsize=11)
    page.insert_text((82, 187), "Identity number", fontsize=11)
    pdf.save(path)
    pdf.close()


def write_signature_line_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = pymupdf.open()
    page = pdf.new_page(width=612, height=792)
    page.insert_text((72, 120), "Signature:", fontsize=11)
    page.draw_line((160, 126), (460, 126), width=0.8)
    pdf.save(path)
    pdf.close()


def write_stacked_body_form_lines_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = pymupdf.open()
    page = pdf.new_page(width=612, height=792)
    page.insert_text((72, 120), "and", fontsize=11)
    page.draw_line((72, 145), (420, 145), width=0.8)
    page.draw_line((72, 170), (190, 170), width=0.8)
    page.insert_text((72, 205), "on 2026", fontsize=11)
    pdf.save(path)
    pdf.close()


def add_bottom_border(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    properties.append(borders)


def poll_job(client: TestClient, job_id: str, timeout: float = 8) -> dict:
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        response = client.get(f"/api/preview-jobs/{job_id}")
        assert response.status_code == 200, response.text
        payload = response.json()
        if payload["status"] in {"completed", "failed"}:
            return payload
        time.sleep(0.03)
    raise AssertionError("Preview job did not reach a terminal state.")


def poll_map(client: TestClient, version_id: str, timeout: float = 8) -> dict:
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        response = client.get(f"/api/document-versions/{version_id}/render-map")
        assert response.status_code == 200, response.text
        payload = response.json()
        if payload["status"] in {"completed", "partial", "failed"}:
            return payload
        time.sleep(0.03)
    raise AssertionError("Render map did not reach a terminal state.")


def test_preview_job_returns_before_word_finishes_and_reuses_cache(
    tmp_path: Path,
    monkeypatch,
) -> None:
    text = "A supported paragraph for inline editing"
    app = load_test_app(tmp_path, monkeypatch)
    document_service = importlib.import_module("app.document_service")
    render_calls = 0

    def fake_word_render(session, document, version):
        nonlocal render_calls
        render_calls += 1
        time.sleep(0.25)
        output_path = document_service.rendered_pdf_path(document, version.id)
        temporary_path = output_path.with_name(f"{version.id}-test-refresh.pdf")
        write_pdf(
            temporary_path,
            [(72, 100, text, None)],
        )
        temporary_path.replace(output_path)
        return document_service.serialize_cached_word_preview(
            session,
            document,
            version,
        )

    monkeypatch.setattr(
        document_service,
        "render_document_with_word",
        fake_word_render,
    )

    with TestClient(app) as client:
        _workspace, primary = upload(client, docx_bytes(text))
        version_id = primary["version_id"]
        started = time.monotonic()
        accepted = client.post(
            f"/api/document-versions/{version_id}/preview-jobs"
        )
        elapsed = time.monotonic() - started

        assert accepted.status_code == 202, accepted.text
        assert elapsed < 0.2
        initial = accepted.json()
        assert initial["status"] in {"queued", "processing"}
        assert initial["version_id"] == version_id
        assert initial["pdf_ready"] is False

        completed = poll_job(client, initial["job_id"])
        assert completed["status"] == "completed"
        assert completed["stage"] == "ready_to_edit"
        assert completed["pdf_ready"] is True
        assert completed["render_map_ready"] is True
        assert completed["cache_hit"] is False
        preview_response = client.get(completed["preview_url"])
        assert preview_response.status_code == 200
        assert preview_response.json()["preview_cache_status"] == "fresh"
        assert render_calls == 1

        first_pages = client.get(f"/api/document-versions/{version_id}/pages")
        second_pages = client.get(f"/api/document-versions/{version_id}/pages")
        assert first_pages.status_code == 200
        assert second_pages.status_code == 200
        assert second_pages.headers["x-docsync-preview-cache"] == "hit"

        cached = client.post(
            f"/api/document-versions/{version_id}/preview-jobs"
        )
        assert cached.status_code == 202
        assert cached.json()["cached_preview"]["preview_cache_status"] == "fresh"
        cached_result = poll_job(client, cached.json()["job_id"])
        assert cached_result["status"] == "completed"
        assert cached_result["pdf_ready"] is True
        assert cached_result["cache_hit"] is True
        assert render_calls == 1

        database = importlib.import_module("app.database")
        models = importlib.import_module("app.models")
        editor_service = importlib.import_module("app.editor_service")
        with database.SessionLocal() as session:
            version = session.get(models.DocumentVersion, version_id)
            assert version is not None
            cache = session.get(models.DocumentPreviewCache, version_id)
            assert cache is not None
            assert cache.word_preview_json["version_id"] == version_id
            source_path = editor_service.document_version_path(version)
            source_stat = source_path.stat()
            os.utime(
                source_path,
                ns=(source_stat.st_atime_ns, source_stat.st_mtime_ns + 1_000_000),
            )

        stale = client.post(f"/api/document-versions/{version_id}/preview-jobs")
        assert stale.status_code == 202
        stale_payload = stale.json()
        assert stale_payload["stale_preview_available"] is True
        assert stale_payload["cached_preview"]["preview_cache_status"] == "stale"
        refreshed = poll_job(client, stale_payload["job_id"])
        assert refreshed["status"] == "completed"
        assert refreshed["cache_hit"] is False
        assert refreshed["stale_preview_available"] is False
        assert render_calls == 2

        with database.SessionLocal() as session:
            version = session.get(models.DocumentVersion, version_id)
            assert version is not None
            source_path = editor_service.document_version_path(version)
            source_stat = source_path.stat()
            os.utime(
                source_path,
                ns=(source_stat.st_atime_ns, source_stat.st_mtime_ns + 1_000_000),
            )

        def fail_refresh(*_args, **_kwargs):
            raise HTTPException(status_code=422, detail="Forced cached refresh failure")

        monkeypatch.setattr(
            document_service,
            "render_document_with_word",
            fail_refresh,
        )
        retained = client.post(
            f"/api/document-versions/{version_id}/preview-jobs"
        )
        assert retained.status_code == 202
        retained_payload = retained.json()
        assert retained_payload["stale_preview_available"] is True
        assert retained_payload["cached_preview"]["preview_cache_status"] == "stale"
        failed_refresh = poll_job(client, retained_payload["job_id"])
        assert failed_refresh["status"] == "failed"
        assert failed_refresh["stale_preview_available"] is True
        still_visible = client.get(f"/api/document-versions/{version_id}/preview")
        assert still_visible.status_code == 200
        assert still_visible.json()["preview_cache_status"] == "stale"
        assert render_calls == 2


def test_controlled_pages_are_published_before_coordinate_matching(
    tmp_path: Path,
    monkeypatch,
) -> None:
    first = "Repeated wording"
    wrapped = "A long supported paragraph that wraps over several visual lines in the controlled preview"
    app = load_test_app(tmp_path, monkeypatch)
    render_map_service = importlib.import_module("app.render_map_service")
    original_match = render_map_service._match_blocks

    def slow_match(*args, **kwargs):
        time.sleep(0.35)
        return original_match(*args, **kwargs)

    monkeypatch.setattr(render_map_service, "_match_blocks", slow_match)

    with TestClient(app) as client:
        workspace, primary = upload(client, docx_bytes(first, first, wrapped))
        version_id = primary["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        write_pdf(
            pdf_path,
            [
                (72, 100, first, None),
                (72, 135, first, None),
                (72, 180, wrapped, 250),
            ],
        )

        queued = client.get(f"/api/document-versions/{version_id}/render-map")
        assert queued.status_code == 200
        progressive = None
        deadline = time.monotonic() + 4
        while time.monotonic() < deadline:
            candidate = client.get(
                f"/api/document-versions/{version_id}/render-map"
            ).json()
            if candidate["status"] == "processing" and candidate["pages"]:
                progressive = candidate
                break
            time.sleep(0.02)
        assert progressive is not None
        assert progressive["page_count"] == 1
        assert progressive["regions"] == []
        page_response = client.get(progressive["pages"][0]["image_url"])
        assert page_response.status_code == 200
        assert page_response.content.startswith(b"\x89PNG")

        mapped = poll_map(client, version_id)
        assert mapped["status"] == "completed"
        duplicate_regions = [
            region for region in mapped["regions"] if region["text_preview"] == first
        ]
        assert len({region["element_id"] for region in duplicate_regions}) == 2
        assert all(region["confidence"] == 0.95 for region in duplicate_regions)
        wrapped_regions = [
            region
            for region in mapped["regions"]
            if region["text_preview"] == wrapped
        ]
        assert len(wrapped_regions) >= 2
        assert all(0 <= region[axis] <= 1 for region in mapped["regions"] for axis in ("x", "y", "width", "height"))


def test_render_map_covers_numbered_toc_and_repeated_legal_paragraphs(
    tmp_path: Path,
    monkeypatch,
) -> None:
    legal_clause = (
        "The Occupant was the previous holder of a Life Right in terms of "
        "a Life Right Agreement entered into between the Occupant and Anson."
    )
    document = Document()
    for text in ("PARTIES", "AGREEMENT", "and", legal_clause, "AGREEMENT"):
        document.add_paragraph(text)
    stream = io.BytesIO()
    document.save(stream)

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, primary = upload(client, stream.getvalue())
        version_id = primary["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        write_pdf(
            pdf_path,
            [
                (72, 90, "1. PARTIES 4", None),
                (72, 125, "2. AGREEMENT 4", None),
                (72, 160, "and", None),
                (72, 195, "A normal sentence with and inside it", None),
                (72, 230, f"2.1 {legal_clause}", 460),
                (72, 315, "AGREEMENT", None),
            ],
        )

        mapped = poll_map(client, version_id)
        assert mapped["status"] == "completed"
        assert mapped["total_element_count"] == 5
        assert mapped["mapped_element_count"] == 5
        assert mapped["interactive_element_count"] == 5
        assert mapped["unmapped"] == []

        regions_by_text: dict[str, set[str]] = {}
        for region in mapped["regions"]:
            regions_by_text.setdefault(region["text_preview"], set()).add(
                region["element_id"]
            )
        assert len(regions_by_text["AGREEMENT"]) == 2
        assert len(regions_by_text["and"]) == 1


def test_render_map_does_not_treat_an_arbitrary_line_substring_as_a_paragraph(
    tmp_path: Path,
    monkeypatch,
) -> None:
    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, primary = upload(client, docx_bytes("Agreement"))
        version_id = primary["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        write_pdf(pdf_path, [(72, 100, "The Agreement applies", None)])

        mapped = poll_map(client, version_id)
        assert mapped["status"] == "failed"
        assert mapped["mapped_element_count"] == 0
        assert len(mapped["unmapped"]) == 1


def test_render_map_uses_neighbouring_blocks_to_resolve_repeated_text(
    tmp_path: Path,
    monkeypatch,
) -> None:
    document = Document()
    for text in ("Before anchor", "AGREEMENT", "After anchor"):
        document.add_paragraph(text)
    stream = io.BytesIO()
    document.save(stream)

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, primary = upload(client, stream.getvalue())
        version_id = primary["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        write_pdf(
            pdf_path,
            [
                (72, 80, "AGREEMENT", None),
                (72, 120, "Before anchor", None),
                (72, 160, "AGREEMENT", None),
                (72, 200, "After anchor", None),
                (72, 240, "AGREEMENT", None),
            ],
        )

        mapped = poll_map(client, version_id)
        assert mapped["status"] == "completed"
        assert mapped["mapped_element_count"] == 3
        agreement_regions = [
            region
            for region in mapped["regions"]
            if region["text_preview"] == "AGREEMENT"
        ]
        assert len(agreement_regions) == 1
        assert agreement_regions[0]["confidence"] == 0.93


def test_render_map_ignores_automatic_page_number_inside_split_paragraph(
    tmp_path: Path,
    monkeypatch,
) -> None:
    before = "This long paragraph begins near the bottom of one page and"
    after = "continues at the top of the following page without interruption."
    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, primary = upload(client, docx_bytes(f"{before} {after}"))
        version_id = primary["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        write_split_paragraph_pdf(pdf_path, before, after)

        mapped = poll_map(client, version_id)
        assert mapped["status"] == "completed"
        assert mapped["mapped_element_count"] == 1
        assert {region["page_number"] for region in mapped["regions"]} == {1, 2}


def test_body_paragraph_wins_when_header_text_overlaps_its_pdf_range(
    tmp_path: Path,
    monkeypatch,
) -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "AGREEMENT"
    document.add_paragraph("2\tAGREEMENT\t4")
    stream = io.BytesIO()
    document.save(stream)

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, primary = upload(client, stream.getvalue())
        version_id = primary["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        write_pdf(
            pdf_path,
            [
                (500, 40, "AGREEMENT", None),
                (72, 160, "2 AGREEMENT 4", None),
            ],
        )

        mapped = poll_map(client, version_id)
        assert mapped["status"] == "completed"
        assert mapped["mapped_element_count"] == 2
        assert {
            region["text_preview"] for region in mapped["regions"]
        } == {"AGREEMENT", "2\tAGREEMENT\t4"}


def test_render_map_makes_empty_table_cells_editable(
    tmp_path: Path,
    monkeypatch,
) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = ""
    table.cell(1, 0).text = "Identity number"
    table.cell(1, 1).text = ""
    stream = io.BytesIO()
    document.save(stream)

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, primary = upload(client, stream.getvalue())
        version_id = primary["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        write_empty_table_pdf(pdf_path)

        mapped = poll_map(client, version_id)
        assert mapped["status"] == "completed"
        assert mapped["total_element_count"] == 4
        assert mapped["mapped_element_count"] == 4
        empty_regions = [
            region for region in mapped["regions"] if region["text_preview"] == ""
        ]
        assert len(empty_regions) == 2
        assert all(region["interactive"] for region in empty_regions)
        assert all(
            region["mapping_method"] == "word_pdf_empty_table_cell"
            for region in empty_regions
        )


def test_render_map_makes_blank_signature_line_editable(
    tmp_path: Path,
    monkeypatch,
) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Signature:"
    table.cell(0, 1).merge(table.cell(0, 2)).text = ""
    stream = io.BytesIO()
    document.save(stream)

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, primary = upload(client, stream.getvalue())
        version_id = primary["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        write_signature_line_pdf(pdf_path)

        mapped = poll_map(client, version_id)
        assert mapped["status"] == "completed"
        assert mapped["total_element_count"] == 2
        assert mapped["mapped_element_count"] == 2
        empty_region = next(
            region for region in mapped["regions"] if region["text_preview"] == ""
        )
        assert empty_region["interactive"] is True
        assert (
            empty_region["mapping_method"]
            == "word_pdf_empty_signature_line"
        )


def test_signature_line_maps_once_to_the_rightmost_empty_word_cell(
    tmp_path: Path,
    monkeypatch,
) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Signature:"
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = ""
    stream = io.BytesIO()
    document.save(stream)

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, primary = upload(client, stream.getvalue())
        version_id = primary["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        write_signature_line_pdf(pdf_path)

        mapped = poll_map(client, version_id)
        empty_regions = [
            region for region in mapped["regions"] if region["text_preview"] == ""
        ]
        assert mapped["total_element_count"] == 3
        assert mapped["mapped_element_count"] == 2
        assert len(empty_regions) == 1
        assert empty_regions[0]["location"]["column_index"] == 2


def test_render_map_maps_distinct_bordered_body_form_lines_below_a_label(
    tmp_path: Path,
    monkeypatch,
) -> None:
    document = Document()
    document.add_paragraph("and")
    first = document.add_paragraph("")
    second = document.add_paragraph("")
    add_bottom_border(first)
    add_bottom_border(second)
    document.add_paragraph("on 2026")
    stream = io.BytesIO()
    document.save(stream)

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, primary = upload(client, stream.getvalue())
        version_id = primary["version_id"]
        pdf_path = tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        write_stacked_body_form_lines_pdf(pdf_path)

        mapped = poll_map(client, version_id)
        empty_regions = [
            region
            for region in mapped["regions"]
            if region["text_preview"] == ""
        ]
        assert len(empty_regions) == 2
        assert {region["mapping_method"] for region in empty_regions} == {
            "word_pdf_empty_body_form_line"
        }
        assert len({region["element_id"] for region in empty_regions}) == 2
        assert len({(region["x"], region["y"], region["width"]) for region in empty_regions}) == 2
        assert all(region["interactive"] for region in empty_regions)
        assert mapped["mapping_diagnostics"] == {
            "detected_horizontal_line_count": 2,
            "blank_form_candidate_count": 2,
            "matched_form_line_count": 2,
        }


def test_render_map_does_not_create_a_target_for_an_unformatted_blank_body_paragraph(
    tmp_path: Path,
    monkeypatch,
) -> None:
    document = Document()
    document.add_paragraph("and")
    document.add_paragraph("")
    document.add_paragraph("on 2026")
    stream = io.BytesIO()
    document.save(stream)

    app = load_test_app(tmp_path, monkeypatch)
    with TestClient(app) as client:
        workspace, primary = upload(client, stream.getvalue())
        version_id = primary["version_id"]
        pdf_path = tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        write_stacked_body_form_lines_pdf(pdf_path)

        mapped = poll_map(client, version_id)
        assert all(region["text_preview"] for region in mapped["regions"])


def test_later_pdf_pages_render_only_on_demand_and_reuse_cache(
    tmp_path: Path,
    monkeypatch,
) -> None:
    page_texts = [f"Unique selectable content for page {page}" for page in range(1, 7)]
    app = load_test_app(tmp_path, monkeypatch)

    with TestClient(app) as client:
        workspace, primary = upload(client, docx_bytes(*page_texts))
        version_id = primary["version_id"]
        pdf_path = (
            tmp_path / "data" / "renders" / workspace["id"] / f"{version_id}.pdf"
        )
        write_multipage_pdf(pdf_path, page_texts)

        mapped = poll_map(client, version_id)
        assert mapped["page_count"] == 6
        render_id = mapped["render_id"]
        page_directory = pdf_path.parent / f"{version_id}.pages" / render_id
        assert not (page_directory / "page-5.png").exists()
        assert not (page_directory / "page-6.png").exists()

        page_url = mapped["pages"][4]["image_url"]
        first = client.get(page_url)
        assert first.status_code == 200
        assert first.content.startswith(b"\x89PNG")
        page_five = page_directory / "page-5.png"
        first_mtime = page_five.stat().st_mtime_ns

        second = client.get(page_url)
        assert second.status_code == 200
        assert page_five.stat().st_mtime_ns == first_mtime
        assert client.get(
            f"/api/document-versions/{version_id}/render-pages/{'0' * 24}/5.png"
        ).status_code == 404
        assert client.get(
            f"/api/document-versions/{version_id}/render-pages/{render_id}/7.png"
        ).status_code == 404


def test_failed_preview_can_be_retried_without_blocking_the_workspace(
    tmp_path: Path,
    monkeypatch,
) -> None:
    text = "Retry this preview"
    app = load_test_app(tmp_path, monkeypatch)
    document_service = importlib.import_module("app.document_service")

    def fail_render(*_args, **_kwargs):
        raise HTTPException(status_code=422, detail="Forced Word render failure")

    monkeypatch.setattr(document_service, "render_document_with_word", fail_render)
    with TestClient(app) as client:
        workspace, primary = upload(client, docx_bytes(text))
        version_id = primary["version_id"]
        first = client.post(f"/api/document-versions/{version_id}/preview-jobs")
        failed = poll_job(client, first.json()["job_id"])
        assert failed["status"] == "failed"
        assert failed["stage"] == "failed"
        assert failed["retry_allowed"] is True
        assert "Forced Word render failure" in failed["error"]
        assert client.get(f"/api/document-versions/{version_id}/pages").status_code == 200

        def recover_render(session, document, version):
            write_pdf(
                tmp_path
                / "data"
                / "renders"
                / workspace["id"]
                / f"{version.id}.pdf",
                [(72, 100, text, None)],
            )
            return document_service.serialize_cached_word_preview(
                session,
                document,
                version,
            )

        monkeypatch.setattr(
            document_service,
            "render_document_with_word",
            recover_render,
        )
        retried = client.post(f"/api/document-versions/{version_id}/preview-jobs")
        assert retried.json()["job_id"] != failed["job_id"]
        recovered = poll_job(client, retried.json()["job_id"])
        assert recovered["status"] == "completed"
        assert recovered["pdf_ready"] is True
